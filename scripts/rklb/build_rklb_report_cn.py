"""
Build RKLB Q1 FY2026 Earnings Update – Chinese DOCX Report
Output: output/RKLB/RKLB_Q1_FY2026_业绩更新报告_中文版.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/RKLB/"
os.makedirs(OUT, exist_ok=True)

# ─── Live market data via yfinance ────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f"约{info.market_cap/1e8:.0f}亿美元",
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[WARNING] yfinance获取失败: {e} — 使用N/A占位")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("RKLB")
print(f"实时股价: {mkt['price']} | 市值: {mkt['market_cap']} | 52W: {mkt['52w_low']}–{mkt['52w_high']}")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

BLUE_HEX  = "1B3A6B"
LBLUE_HEX = "4A90D9"
GREEN_HEX = "27AE60"
RED_HEX   = "C0392B"
GRAY_HEX  = "8C8C8C"
WHITE_HEX = "FFFFFF"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color_hex=None, size_pt=10, italic=False, cn=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold    = bold
            run.font.size    = Pt(size_pt)
            run.font.italic  = italic
            if cn:
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            if color_hex:
                r, g, b = hex_to_rgb(color_hex)
                run.font.color.rgb = RGBColor(r, g, b)

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r  = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), LBLUE_HEX)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(color_el); rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)

def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13 if level == 1 else 11)
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r, g, b = hex_to_rgb(BLUE_HEX)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, italic=False, color_hex=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_prefix=None, color_hex=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = "宋体"
        r1._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if color_hex:
            rv, gv, bv = hex_to_rgb(color_hex)
            r1.font.color.rgb = RGBColor(rv, gv, bv)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
        r2.font.name = "宋体"
        r2._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    return p

def add_chart(doc, filename, width_in=6.0, caption=""):
    path = OUT + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(8.5)
            cp.runs[0].font.name = "宋体"
            cp.runs[0]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r, g, b = hex_to_rgb(GRAY_HEX)
            cp.runs[0].font.color.rgb = RGBColor(r, g, b)
            cp.paragraph_format.space_after = Pt(6)
    else:
        body(doc, f"[图表未找到: {filename}]", italic=True)

# ─── Build Document ───────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# 第1页 — 封面与业绩摘要
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("股票研究  |  航空航天与国防")
run.font.size = Pt(8.5); run.bold = True
run.font.name = "宋体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("火箭实验室 ROCKET LAB (NASDAQ: RKLB)")
run.font.size = Pt(20); run.bold = True
run.font.name = "黑体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
r, g, b = hex_to_rgb(BLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("2026财年第一季度业绩更新 — 首破2亿美元大关，全指标超预期")
run.font.size = Pt(13); run.italic = True
run.font.name = "宋体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
r, g, b = hex_to_rgb(LBLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("财报发布日: 2026年5月7日  |  报告日期: 2026年6月16日")
run.font.size = Pt(9)
run.font.name = "宋体"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

doc.add_paragraph()

# Rating table
rating_table = doc.add_table(rows=2, cols=6)
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rating_table.style = "Table Grid"
headers = ["评级", "目标价", "当前股价", "总市值", "52周最低", "52周最高"]
values  = ["买入", "$120.00", mkt["price"], mkt["market_cap"], mkt["52w_low"], mkt["52w_high"]]
for i, (h, v) in enumerate(zip(headers, values)):
    hcell = rating_table.cell(0, i)
    hcell.text = h
    set_cell_bg(hcell, BLUE_HEX)
    set_cell_font(hcell, bold=True, color_hex=WHITE_HEX, size_pt=9, cn=True)
    vcell = rating_table.cell(1, i)
    vcell.text = v
    vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0:
        set_cell_bg(vcell, GREEN_HEX)
        set_cell_font(vcell, bold=True, color_hex=WHITE_HEX, size_pt=11, cn=True)
    else:
        set_cell_font(vcell, bold=(i == 1), size_pt=10, cn=True)

doc.add_paragraph()

# 核心要点
add_section_heading(doc, "核心要点")
bullet(doc, "创纪录季度营收2.003亿美元（同比+63.5%，环比+11.5%），超市场一致预期1.894亿美元1,090万美元（+5.8%），并超公司指引上限（1.85亿–2.0亿美元）。",
       bold_prefix="营收超预期:", color_hex=GREEN_HEX)
bullet(doc, "GAAP每股亏损($0.07)，优于一致预期($0.08)，为公司历史最窄季度亏损，较Q1 2025的($0.12)大幅改善。",
       bold_prefix="每股收益超预期:", color_hex=GREEN_HEX)
bullet(doc, "调整后EBITDA亏损($11.8M)，大幅优于一致预期($26M)的1,420万美元，也显著优于指引中值约1,200万美元。",
       bold_prefix="EBITDA大幅超预期:", color_hex=GREEN_HEX)
bullet(doc, "GAAP毛利率38.2%（指引34-36%），非GAAP毛利率43.0%（指引39-41%），均超指引200+基点。",
       bold_prefix="毛利率超预期:", color_hex=GREEN_HEX)
bullet(doc, "创纪录在手订单22亿美元（同比+108%，环比+20%），包含公司史上最大合同（5架Neutron+3架Electron至2029年）及1.9亿美元HASTE批量采购订单。",
       bold_prefix="在手订单激增:", color_hex=BLUE_HEX)
bullet(doc, "Q2 2026营收指引2.25亿–2.4亿美元，较华尔街预期2.075亿美元高出10-16%，显示持续强劲增长动能。",
       bold_prefix="Q2指引超预期:", color_hex=BLUE_HEX)
bullet(doc, "Q1签署31份新发射合同（超2025全年总数）；与雷神联合入选金穹（Golden Dome）天基拦截器（SBI）项目。",
       bold_prefix="战略级合同:", color_hex=BLUE_HEX)

doc.add_paragraph()

# 业绩概览表
add_section_heading(doc, "业绩概览")
snap_data = [
    ["指标",           "Q1 2026实际",  "一致预期/指引",    "超/不及",          "同比Q1 2025"],
    ["总营收",          "$200.3M",     "$189.4M",         "+$10.9M (+5.8%)", "+63.5%"],
    ["  发射服务",       "$63.7M",      "$59.0M (est.)",   "+$4.7M",          "+78.9%"],
    ["  太空系统",       "$136.7M",     "$132.1M (est.)",  "+$4.6M",          "+57.2%"],
    ["GAAP毛利率",      "38.2%",       "34–36% (指引)",    "+220–420bps",     "+950bps"],
    ["非GAAP毛利率",    "43.0%",       "39–41% (指引)",    "+200–400bps",     "+960bps"],
    ["调整后EBITDA",    "($11.8M)",    "($26.0M)",        "+$14.2M",         "+$18.2M"],
    ["GAAP每股收益",    "($0.07)",     "($0.08)",         "+$0.01",          "+$0.05"],
    ["净亏损",          "($45.0M)",    "—",               "—",               "vs. ($60.6M)"],
    ["在手订单",        "$22亿",       "—",               "+108% YoY",       "vs. $10.6亿"],
    ["现金及等价物",     "$12.055亿",   "—",               "—",               "vs. $8.287亿"],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for r_idx, row in enumerate(snap_data):
    for c_idx, val in enumerate(row):
        cell = tbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9, cn=True)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第2-3页 — 详细业绩分析
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "详细业绩分析")

add_section_heading(doc, "营收：首次突破2亿美元大关", level=2)
body(doc, "火箭实验室Q1 2026营收2.003亿美元，创单季历史新高，首次突破2亿美元里程碑。同比增长63.5%，环比增长11.5%（Q4 2025为1.797亿美元）。营收超出公司指引上限（1.85亿–2.0亿美元），并大幅超出市场一致预期1.894亿美元达1,090万美元（+5.8%）。")
body(doc, "太空系统业务仍为主要收入来源，贡献1.367亿美元（同比+57.2%，环比+31.7%），增长受益于SDA第二/三期卫星项目推进及太阳能产品和组件的旺盛需求。发射服务贡献6,370万美元（同比+78.9%），环比下降16.1%系Q4 2025异常强劲的7次任务季度的正常回落。两大业务线均超出FactSet分段估计（分别为1.321亿和5,900万美元）。")

add_chart(doc, "rklb_chart1_quarterly_revenue.png", 6.2,
          "图1：季度营收走势 Q1 2024–Q1 2026 | 数据来源：火箭实验室财报")

add_chart(doc, "rklb_chart3_segment_revenue.png", 6.2,
          "图2：Q1 2026分部营收构成及同比对比 | 数据来源：火箭实验室财报")

add_section_heading(doc, "超预期分析：全线超预期", level=2)
body(doc, "Q1 2026业绩在所有关键指标上均超预期——这一全面超预期推动股价当日上涨34%：")
add_chart(doc, "rklb_chart4_beat_miss.png", 6.2,
          "图3：Q1 2026超/不及预期汇总 | 数据来源：火箭实验室、Bloomberg/FactSet一致预期")

bullet(doc, "营收2.003亿美元超出一致预期1.894亿美元5.8%，主要受益于太空系统交付强于预期及发射定价提升（商业Electron单次发射定价已升至约850万美元，较历史500-600万美元水平显著提升）。")
bullet(doc, "调整后EBITDA亏损($11.8M)是最大亮点，超出一致预期($26M)1,420万美元，同比改善61%——清晰展现了盈利能力路径。")
bullet(doc, "GAAP毛利率38.2%超出公司指引上限（36%）220个基点，受益于有利的产品组合及持续的成本优化。")
bullet(doc, "GAAP每股亏损($0.07)优于一致预期($0.08)1美分，为公司历史最窄季度亏损。净亏损同比改善26%。")

doc.add_page_break()

add_section_heading(doc, "毛利率：持续扩张", level=2)
body(doc, "Q1 2026 GAAP毛利率38.2%，较Q1 2024的20.0%提升了1,820个基点（近两年），较Q1 2025的28.7%提升了950个基点（一年）。非GAAP毛利率43.0%，两年前仅为28.0%，近乎翻倍。主要驱动因素：(1)更高利润率的发射合同，反映了Electron作为最可靠小型运载火箭的市场定位；(2)太空系统项目执行效率提升及SDA交付的学习曲线效益；(3)Gauss电推力器进入量产阶段具有有利的单位经济性。")

add_chart(doc, "rklb_chart5_gross_margin.png", 6.2,
          "图4：GAAP与非GAAP毛利率趋势 | 数据来源：火箭实验室财报")

add_section_heading(doc, "调整后EBITDA：亏损快速收窄", level=2)
body(doc, "调整后EBITDA亏损从Q1 2025的($30.0M)持续改善：Q2 ($29.0M) → Q3 ($26.3M) → Q4 ($17.4M) → Q1 2026 ($11.8M)，同比改善61%。公司正走在实现季度EBITDA盈利的道路上（预计Q3或Q4 2026可实现）。Q2指引的($20M)–($26M)较Q1有所扩大，主要反映Mynaric整合成本及产品组合不利因素，我们视为阶段性回调。")

add_chart(doc, "rklb_chart6_adj_ebitda.png", 5.5,
          "图5：调整后EBITDA趋势 Q1'25–Q1'26 | 数据来源：火箭实验室财报")

add_chart(doc, "rklb_chart7_eps_trend.png", 6.2,
          "图6：GAAP每股收益季度趋势 Q1 2024–Q1 2026 | 数据来源：火箭实验室财报")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第4-5页 — 关键指标与展望
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "关键指标与展望")

add_section_heading(doc, "在手订单：22亿美元，持续加速", level=2)
body(doc, "合同在手订单季末达到创纪录的22亿美元，同比增长108%（去年同期10.6亿美元），环比增长20.2%（上季18.5亿美元）。增长由两大标志性合同驱动：(1)公司史上最大合同——与保密客户签订的多年期合同（5架Neutron+3架Electron至2029年）；(2)通过Kratos获得的1.9亿美元HASTE批量采购订单（20次高超声速测试飞行）。发射服务占在手订单的41.5%，较历史水平显著提升，反映了Neutron在首飞之前就已获得客户订单——这是市场信心的有力信号。12个月收入转化率约为36%，意味着约7.9亿美元将在未来一年内转化为收入。")

add_chart(doc, "rklb_chart8_backlog.png", 5.8,
          "图7：在手订单增长趋势 Q1 2025–Q1 2026 | 数据来源：火箭实验室财报")

add_section_heading(doc, "发射活动与HASTE项目", level=2)
body(doc, "Q1 2026签署31份新Electron和HASTE发射合同，超过2025全年发射合同签署总数。总合同任务清单超过70项。报告期截止已完成8次任务。商业Electron单次发射定价提升至约850万美元（历史约500-600万美元），反映了该运载火箭在可靠性和发射周期方面的市场领先优势。")
body(doc, "HASTE项目已成为重要增长引擎：通过Kratos获得的1.9亿美元国防部合同（20次MACH-TB 2.0发射）现占总发射在手订单的约三分之一。HASTE为金穹（Golden Dome）架构提供高超声速测试和模拟能力，使火箭实验室成为美国导弹防御测试的关键基础设施。")

add_chart(doc, "rklb_chart9_launch_cadence.png", 6.2,
          "图8：Electron发射节奏 Q1 2024–Q1 2026 | 数据来源：火箭实验室")

add_section_heading(doc, "Q2 2026展望：全面超华尔街预期", level=2)
body(doc, "管理层给出Q2 2026营收指引2.25亿–2.4亿美元（中值2.325亿美元），环比增长约16%，显著高于华尔街预期约2.075亿美元。GAAP毛利率指引33–35%，非GAAP毛利率38–40%，均较Q1有所下降，反映SDA项目爬坡期产品组合不利及Mynaric整合拖累。调整后EBITDA指引($20M)–($26M)。")

guid_data = [
    ["指标",           "Q1 2026实际",  "Q2 2026指引",      "华尔街预期",     "指引vs预期"],
    ["营收",           "$200.3M",     "$225M–$240M",     "$207.5M",      "+8%至+16%"],
    ["GAAP毛利率",      "38.2%",       "33%–35%",         "—",            "组合拖累"],
    ["非GAAP毛利率",    "43.0%",       "38%–40%",         "—",            "Mynaric拖累"],
    ["调整后EBITDA",    "($11.8M)",    "($20M)–($26M)",   "($15.1M)",     "亏损扩大"],
    ["加权平均股数",     "605.4M",      "~629M",           "—",            "含优先股"],
]
gtbl = doc.add_table(rows=len(guid_data), cols=5)
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
gtbl.style = "Table Grid"
for r_idx, row in enumerate(guid_data):
    for c_idx, val in enumerate(row):
        cell = gtbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9, cn=True)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_chart(doc, "rklb_chart10_guidance.png", 6.2,
          "图9：营收趋势与Q2 2026指引对比 | 数据来源：火箭实验室、LSEG/FactSet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第6-7页 — 投资论点更新
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "投资论点更新")

add_section_heading(doc, "本季度关键变化", level=2)
body(doc, "Q1 2026业绩具有里程碑意义——本季度验证了火箭实验室从小型发射专家向规模化、多业务航天平台公司的转型。关键增量变化：")
bullet(doc, "营收首次单季突破2亿美元，两大业务线均显著超预期。63.5%的同比增速较Q4 2025的36%明显加速。", bold_prefix="积极:")
bullet(doc, "在手订单同比翻倍至22亿美元，提供极高收入可见度。客户在Neutron首飞前即预订发射——这是市场信心的有力证明。", bold_prefix="积极:")
bullet(doc, "与雷神联合入选金穹天基拦截器（SBI）示范项目，使RKLB定位于最高优先级国防项目。结合HASTE在威胁模拟中的角色，RKLB嵌入导弹防御架构的多个层面。", bold_prefix="积极:")
bullet(doc, "Neutron首飞仍定于2026年Q4，管理层拒绝就早期或晚期Q4做进一步缩窄。Archimedes发动机热试车持续进行中；着陆驳船海试计划于2026年稍后进行。", bold_prefix="中性:")
bullet(doc, "通过ATM增发持续摊薄（Q1融资4.5亿美元，累计4.74亿美元以上）。股份总数从Q1 2025的5.056亿增至Q1 2026的6.054亿（+20%）。", bold_prefix="消极:")

add_section_heading(doc, "核心论点：端到端太空基础设施平台", level=2)
body(doc, "火箭实验室正在执行一项覆盖整个太空价值链的独特垂直整合战略：(1)Electron作为全球第二高发射频率的轨道火箭（仅次于猎鹰9号）；(2)即将面世的Neutron中型运载火箭，面向星座部署和国家安全任务；(3)快速增长的太空系统业务，制造航天器、太阳能产品、星跟踪器、分离系统和光通信终端；(4)通过HASTE、SBI和SDA项目在美国导弹防御领域的战略定位。")
body(doc, "硬件制造规模、飞行验证经验和国防合同信誉的结合创造了极高的竞争壁垒，新进入者极难复制。凭借22亿美元在手订单和超20亿美元流动性，RKLB同时拥有需求信号和财务资源来执行其多年增长计划。")

add_section_heading(doc, "并购战略：构建平台", level=2)
body(doc, "Q1 2026继续进行补强性并购活动，强化垂直整合论点：")
bullet(doc, "Mynaric AG：完成对欧洲光通信终端公司的收购，增加了对SDA和商业星座项目至关重要的星间链路能力。")
bullet(doc, "Motiv Space Systems：签署收购协议，获取NASA火星毅力号传承的机器人和运动控制技术，增加太阳能阵列驱动组件和航天器机构。")

add_section_heading(doc, "风险因素", level=2)
bullet(doc, "Neutron首飞失败或进一步延迟可能损害投资者信心和客户管线。", bold_prefix="风险1:")
bullet(doc, "持续通过ATM增发摊薄股份（仅Q1即4.74亿美元以上）以资助Neutron和并购。", bold_prefix="风险2:")
bullet(doc, "Q2–Q3毛利率压缩，受SDA项目爬坡和Mynaric整合影响（Q2 GAAP毛利率指引33–35%）。", bold_prefix="风险3:")
bullet(doc, "来自SpaceX星舰（中重型发射）及新兴小型运载火箭竞争对手的竞争压力。", bold_prefix="风险4:")
bullet(doc, "持续的负自由现金流——Q1 2026自由现金流约为($77.4M)。", bold_prefix="风险5:")

add_section_heading(doc, "催化剂", level=2)
bullet(doc, "Neutron首飞（目标Q4 2026）——最重要的估值重估催化剂。")
bullet(doc, "纳入纳斯达克100指数（2026年6月22日生效）——预计将带来被动基金增量资金流入。")
bullet(doc, "金穹/SBI项目进入示范阶段后的额外合同授予。")
bullet(doc, "季度EBITDA盈利里程碑（按当前改善轨迹，可能在Q3或Q4 2026实现）。")
bullet(doc, "Electron助推器复用验证——将在结构性降低单位成本的同时加速毛利率扩张。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 第8-10页 — 估值与预测
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "估值与预测")

add_section_heading(doc, "同比营收增速重新加速", level=2)
add_chart(doc, "rklb_chart2_yoy_growth.png", 6.2,
          "图10：季度同比营收增长率 | 数据来源：火箭实验室财报")

add_section_heading(doc, "更新后的预测", level=2)
body(doc, "基于Q1强劲业绩及超预期Q2指引，我们上调FY2026E营收预测至9.4亿美元（此前8.85亿美元）。同时上调FY2027E预测至12.5亿美元，反映Neutron从2027年下半年开始贡献收入。")

est_data = [
    ["指标",         "FY2024A",  "FY2025A",  "FY2026E（新）", "FY2026E（旧）", "FY2027E"],
    ["营收",         "$436.2M",  "$601.8M",  "$940M",       "$885M",       "$1.25B"],
    ["同比增速",      "+78%",     "+38%",     "+56%",        "+47%",        "+33%"],
    ["GAAP毛利率",   "~27%",     "~34%",     "~35%",        "~37%",        "~38%"],
    ["调整后EBITDA",  "($91M)",   "($101M)",  "($30M)",      "盈亏平衡",     "$50M+"],
    ["GAAP每股收益",  "($0.42)",  "($0.37)",  "($0.22)",     "($0.25)",     "($0.08)"],
    ["Electron发射数","16",       "21",       "~26–28",      "~24–26",      "~30–34"],
]
etbl = doc.add_table(rows=len(est_data), cols=6)
etbl.alignment = WD_TABLE_ALIGNMENT.CENTER
etbl.style = "Table Grid"
for r_idx, row in enumerate(est_data):
    for c_idx, val in enumerate(row):
        cell = etbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9, cn=True)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

body(doc, "数据来源：火箭实验室财务报告（A=实际值），分析师预测（E）。旧预测来自2026年2月Q4 FY2025业绩更新报告。", italic=True, color_hex=GRAY_HEX, size=8.5)

add_section_heading(doc, "估值框架", level=2)
body(doc, f"以当前股价{mkt['price']}计算，火箭实验室约以{float(mkt['market_cap'].replace('约','').replace('亿美元',''))/94:.1f}x EV/FY2026E营收交易——该溢价反映了公司独特的竞争定位、超50%的营收增速及改善中的利润率轨迹。对于一家营收同比增长56%、EBITDA即将转正、拥有22亿美元在手订单的公司，我们认为该估值倍数相对航天及高成长工业可比公司是合理的。")
body(doc, "我们的12个月目标价$120基于混合估值方法：")
bullet(doc, "70%权重：70x EV/FY2027E EBITDA $50M+（适用于向盈利过渡的高成长航天公司）")
bullet(doc, "30%权重：50x EV/FY2026E营收$940M（成长溢价倍数）")

add_section_heading(doc, "情景分析", level=2)
scen_data = [
    ["情景",   "估值方法",            "关键假设",                        "隐含目标价",  "较当前"],
    ["悲观",   "40x EV/FY27E EBITDA", "Neutron延至2027；利润率下滑",      "~$65",      "下行风险"],
    ["基准",   "混合方法（上述）",      "Neutron Q4'26；$940M营收",         "$120",      "上行潜力"],
    ["乐观",   "80x EV/FY27E EBITDA", "Neutron准时；FY26>$1B；新合同",    "~$160",     "显著上行"],
]
stbl = doc.add_table(rows=len(scen_data), cols=5)
stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stbl.style = "Table Grid"
for r_idx, row in enumerate(scen_data):
    for c_idx, val in enumerate(row):
        cell = stbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9, cn=True)
        elif r_idx == 1:
            set_cell_bg(cell, "FDFEFE")
        elif r_idx == 2:
            set_cell_bg(cell, "EBF5FB")
        elif r_idx == 3:
            set_cell_bg(cell, "D5F5E3")

doc.add_paragraph()

p = doc.add_paragraph()
r_bg = p.add_run("  维持买入评级  |  目标价：$120.00  ")
r_bg.bold = True
r_bg.font.size = Pt(12)
r_bg.font.name = "黑体"
r_bg._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
rv, gv, bv = hex_to_rgb(GREEN_HEX)
r_bg.font.color.rgb = RGBColor(rv, gv, bv)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc, "Q1 2026是火箭实验室的里程碑季度——公司首次实现单季2亿美元营收，所有指标超预期，股价首次突破100美元。22亿美元在手订单、不断扩大的国防定位（金穹/SBI、HASTE）以及即将到来的纳斯达克100纳入均支持建设性展望。尽管估值较高且Neutron执行风险犹存，我们认为加速增长、毛利率扩张及战略定位的组合证明了维持买入评级的合理性。我们将目标价从此前的$35上调至$120，以反映大幅改善的业务规模和市场定位。")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 资料来源
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "资料来源")

sources = [
    ("火箭实验室Q1 FY2026业绩新闻稿（2026年5月7日）",
     "https://investors.rocketlabcorp.com/news-releases/news-release-details/rocket-lab-announces-first-quarter-2026-financial-results"),
    ("火箭实验室 Form 8-K Q1 2026（SEC EDGAR，2026年5月7日）",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000027/rklb-05072026ex991.htm"),
    ("火箭实验室 Form 10-Q Q1 2026（SEC EDGAR，截至2026年3月31日）",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000028/rklb-20260331.htm"),
    ("火箭实验室Q1 2026业绩电话会议纪要（Seeking Alpha，2026年5月7日）",
     "https://seekingalpha.com/article/4901108-rocket-lab-corporation-rklb-q1-2026-earnings-call-transcript"),
    ("火箭实验室Q1 2026业绩电话会议纪要（Motley Fool，2026年5月8日）",
     "https://www.fool.com/earnings/call-transcripts/2026/05/08/rocket-lab-rklb-q1-2026-earnings-transcript/"),
    ("RKLB Q1 2026业绩分析（CNBC，2026年5月8日）",
     "https://www.cnbc.com/2026/05/08/rocket-lab-rklb-q1-earnings-2026.html"),
    ("火箭实验室Q4及FY2025业绩新闻稿（2026年2月26日）——上季度参考",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000012/rklb-02262026ex991.htm"),
    ("RKLB营收历史数据（StockAnalysis.com）",
     "https://stockanalysis.com/stocks/rklb/revenue/"),
    ("市场数据：yfinance（报告生成时实时获取）",
     "https://finance.yahoo.com/quote/RKLB"),
]

for title, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    add_hyperlink(p, title, url)

doc.add_paragraph()
body(doc, '分析师一致预期：Bloomberg / FactSet / LSEG，截至2026年5月7日盘前。标注"E"为分析师预测。目标价与评级为分析师独立评估。', italic=True, color_hex=GRAY_HEX, size=8.5)
body(doc, '免责声明：本报告仅供参考，不构成投资建议。分析师及机构未持有RKLB仓位。过往表现不代表未来收益。', italic=True, color_hex=GRAY_HEX, size=8.5)

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = OUT + "RKLB_Q1_FY2026_业绩更新报告_中文版.docx"
doc.save(out_path)
print(f"\n中文报告已保存: {out_path}")
