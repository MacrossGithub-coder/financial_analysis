"""
Build RKLB Q1 FY2026 Earnings Update – English DOCX Report
Output: output/RKLB/RKLB_Q1_FY2026_Earnings_Update.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT  = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/RKLB/"
os.makedirs(OUT, exist_ok=True)

# ─── Live market data via yfinance ────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        import yfinance as yf
        info = yf.Ticker(ticker).fast_info
        return {
            "price":      f"${info.last_price:.2f}",
            "market_cap": f"~${info.market_cap/1e9:.1f}B",
            "52w_high":   f"${info.year_high:.2f}",
            "52w_low":    f"${info.year_low:.2f}",
        }
    except Exception as e:
        print(f"[WARNING] yfinance fetch failed: {e} — using N/A")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("RKLB")
print(f"Live price: {mkt['price']} | Mkt Cap: {mkt['market_cap']} | 52W: {mkt['52w_low']}–{mkt['52w_high']}")

# ─── Helpers ──────────────────────────────────────────────────────────────────
def hex_to_rgb(h):
    h = h.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

BLUE_HEX   = "1B3A6B"
LBLUE_HEX  = "4A90D9"
GREEN_HEX  = "27AE60"
RED_HEX    = "C0392B"
GRAY_HEX   = "8C8C8C"
GOLD_HEX   = "F39C12"
WHITE_HEX  = "FFFFFF"

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color_hex=None, size_pt=10, italic=False):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.bold    = bold
            run.font.size    = Pt(size_pt)
            run.font.italic  = italic
            if color_hex:
                r, g, b = hex_to_rgb(color_hex)
                run.font.color.rgb = RGBColor(r, g, b)

def add_hyperlink(para, text, url):
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r  = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), LBLUE_HEX)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(color_el); rPr.append(u_el)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    para._p.append(hl)

def add_section_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(12 if level == 1 else 10.5)
    r, g, b = hex_to_rgb(BLUE_HEX)
    run.font.color.rgb = RGBColor(r, g, b)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), BLUE_HEX)
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, bold=False, italic=False, color_hex=None, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color_hex:
        r, g, b = hex_to_rgb(color_hex)
        run.font.color.rgb = RGBColor(r, g, b)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(doc, text, bold_prefix=None, color_hex=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    if bold_prefix:
        r1 = p.add_run(bold_prefix + " ")
        r1.bold = True
        r1.font.size = Pt(10)
        if color_hex:
            rv, gv, bv = hex_to_rgb(color_hex)
            r1.font.color.rgb = RGBColor(rv, gv, bv)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p

def add_chart(doc, filename, width_in=6.0, caption=""):
    path = OUT + filename
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width_in))
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.runs[0].italic = True
            cp.runs[0].font.size = Pt(8.5)
            r, g, b = hex_to_rgb(GRAY_HEX)
            cp.runs[0].font.color.rgb = RGBColor(r, g, b)
            cp.paragraph_format.space_after = Pt(6)
    else:
        body(doc, f"[Chart not found: {filename}]", italic=True)

# ─── Build Document ───────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER & EARNINGS SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EQUITY RESEARCH  |  AEROSPACE & DEFENSE")
run.font.size = Pt(8.5)
run.bold = True
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ROCKET LAB USA, INC. (NASDAQ: RKLB)")
run.font.size = Pt(20)
run.bold = True
r, g, b = hex_to_rgb(BLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Q1 FY2026 Earnings Update — Record $200M Quarter, All Metrics Beat")
run.font.size = Pt(13)
run.italic = True
r, g, b = hex_to_rgb(LBLUE_HEX)
run.font.color.rgb = RGBColor(r, g, b)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("May 7, 2026  |  Earnings Release Date  |  Report Date: June 16, 2026")
run.font.size = Pt(9)
r, g, b = hex_to_rgb(GRAY_HEX)
run.font.color.rgb = RGBColor(r, g, b)

doc.add_paragraph()

# Rating / Valuation table
rating_table = doc.add_table(rows=2, cols=6)
rating_table.alignment = WD_TABLE_ALIGNMENT.CENTER
rating_table.style = "Table Grid"
headers = ["Rating", "Price Target", "Current Price", "Market Cap", "52W Low", "52W High"]
values  = ["BUY", "$120.00", mkt["price"], mkt["market_cap"], mkt["52w_low"], mkt["52w_high"]]
for i, (h, v) in enumerate(zip(headers, values)):
    hcell = rating_table.cell(0, i)
    hcell.text = h
    set_cell_bg(hcell, BLUE_HEX)
    set_cell_font(hcell, bold=True, color_hex=WHITE_HEX, size_pt=9)
    vcell = rating_table.cell(1, i)
    vcell.text = v
    vcell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0:
        set_cell_bg(vcell, GREEN_HEX)
        set_cell_font(vcell, bold=True, color_hex=WHITE_HEX, size_pt=11)
    else:
        set_cell_font(vcell, bold=(i == 1), size_pt=10)

doc.add_paragraph()

# KEY TAKEAWAYS
add_section_heading(doc, "Key Takeaways")
bullet(doc, "Record quarterly revenue of $200.3M (+63.5% YoY, +11.5% QoQ), exceeding consensus of $189.4M by $10.9M (+5.8%) and surpassing the high end of $185M–$200M guidance.",
       bold_prefix="Revenue Beat:", color_hex=GREEN_HEX)
bullet(doc, "GAAP EPS of ($0.07) vs. consensus ($0.08) — narrowest quarterly loss in company history, improving from ($0.12) in Q1 2025.",
       bold_prefix="EPS Beat:", color_hex=GREEN_HEX)
bullet(doc, "Adjusted EBITDA loss of ($11.8M) beat consensus ($26M) by $14.2M and guidance midpoint by ~$12M, driven by stronger revenue and margins.",
       bold_prefix="EBITDA Beat:", color_hex=GREEN_HEX)
bullet(doc, "GAAP gross margin of 38.2% (guide: 34–36%) and non-GAAP of 43.0% (guide: 39–41%), both exceeding guidance by 200+ bps.",
       bold_prefix="Margin Outperformance:", color_hex=GREEN_HEX)
bullet(doc, "Record backlog of $2.2B (+108% YoY, +20% QoQ), including largest contract in company history (5 Neutron + 3 Electron through 2029) and $190M HASTE block buy.",
       bold_prefix="Backlog Surge:", color_hex=BLUE_HEX)
bullet(doc, "Q2 2026 revenue guidance of $225M–$240M, ~10–16% above Street estimate of $207.5M, implying continued strong momentum.",
       bold_prefix="Strong Q2 Guidance:", color_hex=BLUE_HEX)
bullet(doc, "31 new launch contracts signed in Q1 (exceeded full-year 2025 total); selected with Raytheon for Space-Based Interceptor under Golden Dome program.",
       bold_prefix="Strategic Wins:", color_hex=BLUE_HEX)

doc.add_paragraph()

# RESULTS SNAPSHOT TABLE
add_section_heading(doc, "Results Snapshot")
snap_data = [
    ["Metric",              "Q1 2026 Actual", "Consensus / Guide", "Beat/Miss",         "vs. Q1 2025"],
    ["Revenue",             "$200.3M",        "$189.4M",           "+$10.9M (+5.8%)",   "+63.5% YoY"],
    ["  Launch Services",   "$63.7M",         "$59.0M (est.)",     "+$4.7M",            "+78.9% YoY"],
    ["  Space Systems",     "$136.7M",        "$132.1M (est.)",    "+$4.6M",            "+57.2% YoY"],
    ["GAAP Gross Margin",   "38.2%",          "34–36% (guide)",    "+220–420bps",       "+950bps YoY"],
    ["Non-GAAP Gross Mgn",  "43.0%",          "39–41% (guide)",    "+200–400bps",       "+960bps YoY"],
    ["Adj. EBITDA",         "($11.8M)",       "($26.0M)",          "+$14.2M",           "+$18.2M YoY"],
    ["GAAP EPS",            "($0.07)",        "($0.08)",           "+$0.01",            "+$0.05 YoY"],
    ["Net Loss",            "($45.0M)",       "—",                 "—",                 "vs. ($60.6M)"],
    ["Backlog",             "$2.2B",          "—",                 "+108% YoY",         "vs. $1.06B"],
    ["Cash & Equivalents",  "$1,205.5M",      "—",                 "—",                 "vs. $828.7M"],
]
tbl = doc.add_table(rows=len(snap_data), cols=5)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = "Table Grid"
for r_idx, row in enumerate(snap_data):
    for c_idx, val in enumerate(row):
        cell = tbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 2-3 — DETAILED RESULTS ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Detailed Results Analysis")

add_section_heading(doc, "Revenue: First-Ever $200M+ Quarter", level=2)
body(doc, "Rocket Lab reported Q1 2026 revenue of $200.3M, a record for any single quarter, surpassing the $200M milestone for the first time. This represents +63.5% year-over-year growth and +11.5% sequential growth from Q4 2025's $179.7M. The quarter exceeded the high end of company guidance ($185M–$200M) and beat consensus estimates of $189.4M by $10.9M (+5.8%).")
body(doc, "Space Systems remained the dominant segment at $136.7M (+57.2% YoY, +31.7% QoQ), driven by SDA Tranche II/III satellite programs and strong demand for solar products and components. Launch Services contributed $63.7M (+78.9% YoY), though down 16.1% sequentially from Q4 2025's unusually strong 7-mission quarter, reflecting normal cadence variability. Both segments exceeded FactSet segment estimates ($132.1M and $59.0M respectively).")

add_chart(doc, "rklb_chart1_quarterly_revenue.png", 6.2,
          "Figure 1: Quarterly Revenue Q1 2024–Q1 2026 | Source: Rocket Lab earnings releases")

add_chart(doc, "rklb_chart3_segment_revenue.png", 6.2,
          "Figure 2: Segment Revenue Mix Q1 2026 & YoY Comparison | Source: Rocket Lab earnings releases")

add_section_heading(doc, "Beat / Miss Analysis: Clean Sweep", level=2)
body(doc, "Q1 2026 results delivered a comprehensive beat across every key metric — a rare clean sweep that drove a +34% single-day stock reaction:")
add_chart(doc, "rklb_chart4_beat_miss.png", 6.2,
          "Figure 3: Q1 2026 Beat/Miss Summary | Source: Rocket Lab, Bloomberg/FactSet consensus")

bullet(doc, "Revenue of $200.3M beat consensus of $189.4M by +5.8%, driven by stronger-than-expected Space Systems deliveries and launch pricing improvement (commercial Electron pricing now ~$8.5M per launch, up from $5–6M historically).")
bullet(doc, "Adjusted EBITDA loss of ($11.8M) was the highlight, beating consensus ($26M) by $14.2M and narrowing 61% YoY from ($30.0M) in Q1 2025 — demonstrating clear path to profitability as operating leverage takes hold.")
bullet(doc, "GAAP gross margin of 38.2% beat the top end of company guidance (36%) by 220bps, driven by favorable product mix and continued cost optimization. Non-GAAP gross margin of 43.0% also exceeded guidance (41%) by 200bps.")
bullet(doc, "GAAP EPS of ($0.07) beat consensus ($0.08) by $0.01, representing the narrowest quarterly loss in company history. Net loss improved 26% YoY from ($60.6M) to ($45.0M).")

doc.add_page_break()

add_section_heading(doc, "Gross Margin: Sustained Expansion", level=2)
body(doc, "Q1 2026 marked the ninth consecutive quarter of YoY GAAP gross margin expansion. GAAP gross margin of 38.2% was up +950bps from Q1 2024's 20.0% and +530bps from Q1 2025's 28.7% (adjusted for SBC). Non-GAAP gross margin of 43.0% has nearly doubled from 28.0% two years ago. Key drivers include: (1) higher-margin launch contracts reflecting Electron's market positioning as the most reliable small launch vehicle, (2) improved Space Systems program execution and learning-curve benefits on SDA deliveries, and (3) Gauss electric propulsion entering high-volume production with favorable unit economics.")

add_chart(doc, "rklb_chart5_gross_margin.png", 6.2,
          "Figure 4: GAAP and Non-GAAP Gross Margin Trend | Source: Rocket Lab earnings releases")

add_section_heading(doc, "Adjusted EBITDA: Losses Narrowing Rapidly", level=2)
body(doc, "Adjusted EBITDA loss of ($11.8M) in Q1 2026 continued the clear improvement trajectory: ($30.0M) in Q1 2025 → ($29.0M) Q2 → ($26.3M) Q3 → ($17.4M) Q4 → ($11.8M) Q1 2026. This represents a 61% improvement YoY and places the company on track for positive quarterly Adjusted EBITDA in the coming quarters, despite ongoing Neutron development investment. Q2 2026 guidance of ($20M)–($26M) reflects wider loss range due to Mynaric integration costs and unfavorable product mix, but we view this as a temporary step-back.")

add_chart(doc, "rklb_chart6_adj_ebitda.png", 5.5,
          "Figure 5: Adjusted EBITDA Trend Q1'25–Q1'26 | Source: Rocket Lab earnings releases")

add_chart(doc, "rklb_chart7_eps_trend.png", 6.2,
          "Figure 6: GAAP EPS Quarterly Trend Q1 2024–Q1 2026 | Source: Rocket Lab earnings releases")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 4-5 — KEY METRICS & GUIDANCE
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Key Metrics & Guidance")

add_section_heading(doc, "Backlog: $2.2B and Accelerating", level=2)
body(doc, "Contracted backlog reached a record $2.2B at Q1-end, up +108% YoY from $1.06B and +20.2% QoQ from $1.85B. The growth was driven by two landmark contracts: (1) the largest contract in company history — a multi-year deal with a confidential customer for 5 Neutron and 3 Electron launches through 2029, and (2) a $190M HASTE block buy for 20 hypersonic test flights via Kratos. Launch Services now represents 41.5% of backlog, up from historically lower levels, reflecting the Neutron order pipeline emerging even before first flight. The 12-month revenue conversion rate is approximately 36%, implying ~$790M of the backlog converts to revenue within the next year.")

add_chart(doc, "rklb_chart8_backlog.png", 5.8,
          "Figure 7: Backlog Growth Q1 2025–Q1 2026 | Source: Rocket Lab earnings releases")

add_section_heading(doc, "Launch Activity & HASTE Program", level=2)
body(doc, "Rocket Lab signed 31 new Electron and HASTE launch contracts in Q1 2026 — exceeding the total number of launch contracts signed in all of FY2025. The total contracted manifest exceeds 70 missions. Eight missions were completed year-to-date at reporting. Commercial Electron pricing has increased to approximately $8.5M per launch (from $5–6M historically), reflecting the vehicle's market-leading reliability and time-to-launch advantage.")
body(doc, "The HASTE program has become a major growth vector: the $190M DoD contract via Kratos (20-launch MACH-TB 2.0 order) now represents approximately one-third of total launch backlog. HASTE provides the hypersonic test and simulation capability critical to the Golden Dome architecture, positioning Rocket Lab as essential infrastructure for U.S. missile defense testing.")

add_chart(doc, "rklb_chart9_launch_cadence.png", 6.2,
          "Figure 8: Electron Launch Cadence Q1 2024–Q1 2026 | Source: Rocket Lab")

add_section_heading(doc, "Q2 2026 Guidance: Above Street Across All Metrics", level=2)
body(doc, "Management guided Q2 2026 revenue of $225M–$240M (midpoint $232.5M), representing ~16% sequential growth and significantly above the Street estimate of ~$207.5M. This implies annualized run-rate approaching $930M–$960M. GAAP gross margin is guided to 33–35% (down from Q1's 38.2%) and non-GAAP gross margin to 38–40% (down from 43.0%), reflecting unfavorable product mix from SDA program ramp and Mynaric integration drag. Adjusted EBITDA is guided to ($20M)–($26M), wider than Q1's ($11.8M).")

guid_data = [
    ["Metric",            "Q1 2026 Actual", "Q2 2026 Guidance",  "Street Estimate",   "Guide vs. Street"],
    ["Revenue",           "$200.3M",        "$225M–$240M",       "$207.5M",           "+8% to +16%"],
    ["GAAP Gross Margin", "38.2%",          "33%–35%",           "—",                 "Mix headwind"],
    ["Non-GAAP GM",       "43.0%",          "38%–40%",           "—",                 "Mynaric drag"],
    ["Adj. EBITDA",       "($11.8M)",       "($20M)–($26M)",    "($15.1M)",          "Wider loss"],
    ["GAAP OpEx",         "$132.5M",        "$138M–$144M",       "—",                 "Neutron ramp"],
    ["Wgt Avg Shares",    "605.4M",         "~629M",             "—",                 "Includes preferred"],
]
gtbl = doc.add_table(rows=len(guid_data), cols=5)
gtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
gtbl.style = "Table Grid"
for r_idx, row in enumerate(guid_data):
    for c_idx, val in enumerate(row):
        cell = gtbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

add_chart(doc, "rklb_chart10_guidance.png", 6.2,
          "Figure 9: Revenue Trend & Q2 2026 Guidance vs. Street | Source: Rocket Lab, LSEG/FactSet")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 6-7 — INVESTMENT THESIS UPDATE
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Investment Thesis Update")

add_section_heading(doc, "What Changed This Quarter", level=2)
body(doc, "Q1 2026 results were transformational — this was the quarter that validated Rocket Lab's evolution from a small-launch specialist into a scaled, multi-segment aerospace platform company. Key incremental developments:")
bullet(doc, "Revenue crossed $200M in a single quarter for the first time, with both segments significantly beating estimates. The 63.5% YoY growth rate reaccelerated from 36% in Q4 2025.", bold_prefix="Positive:")
bullet(doc, "Backlog doubled YoY to $2.2B, providing exceptional revenue visibility. Customers are now pre-booking Neutron launches before first flight — a strong signal of market confidence.", bold_prefix="Positive:")
bullet(doc, "Selection with Raytheon for Space-Based Interceptor (SBI) demonstration under Golden Dome positions RKLB in the highest-priority national defense program. Combined with HASTE's role in threat simulation, RKLB is embedded across multiple layers of the missile defense architecture.", bold_prefix="Positive:")
bullet(doc, "Neutron first launch remains targeted for Q4 2026 with management declining to narrow beyond early vs. late Q4. Archimedes engine hot-fire testing ongoing; landing barge sea trials planned later in 2026.", bold_prefix="Neutral:")
bullet(doc, "Continued share dilution via ATM offerings ($450M raised in Q1, $474M+ total) remains a concern. Share count expanded from 505.6M in Q1 2025 to 605.4M in Q1 2026 (+20%).", bold_prefix="Negative:")

add_section_heading(doc, "Thesis: End-to-End Space Infrastructure Platform", level=2)
body(doc, "Rocket Lab is executing a unique vertical integration strategy spanning the full space value chain: (1) Electron as the world's second-most-frequently-launched orbital rocket (behind Falcon 9), (2) Neutron as the upcoming medium-lift contender for constellation deployment and national security missions, (3) a rapidly growing Space Systems business manufacturing spacecraft, solar products, star trackers, separation systems, and optical communications terminals, and (4) strategic positioning across U.S. missile defense through HASTE, SBI, and SDA programs.")
body(doc, "The combination of hardware manufacturing scale, flight heritage, and defense contract credibility creates a competitive moat that is extremely difficult for new entrants to replicate. With $2.2B in backlog and >$2.0B in liquidity, RKLB has both the demand signal and financial resources to execute on its multi-year growth plan.")

add_section_heading(doc, "M&A Strategy: Building the Platform", level=2)
body(doc, "Q1 2026 saw continued bolt-on M&A activity reinforcing the vertical integration thesis:")
bullet(doc, "Mynaric AG: Completed acquisition of European optical laser communication terminal company, adding inter-satellite link capability critical for SDA and commercial constellation programs.")
bullet(doc, "Motiv Space Systems: Signed agreement to acquire NASA Mars Perseverance heritage robotics and motion control company, adding solar array drive assemblies and spacecraft mechanisms.")
body(doc, "These acquisitions extend RKLB's component portfolio and reduce supply chain dependency on external vendors, while also creating cross-selling opportunities within the existing Space Systems order book.")

add_section_heading(doc, "Risks", level=2)
bullet(doc, "Neutron first flight failure or further delay could damage investor confidence and customer pipeline.", bold_prefix="Key Risk 1:")
bullet(doc, "Continued share dilution through ATM offerings ($474M+ in Q1 alone) to fund Neutron and M&A.", bold_prefix="Key Risk 2:")
bullet(doc, "Margin compression in Q2–Q3 from SDA program ramp and Mynaric integration (guided 33–35% GAAP GM in Q2).", bold_prefix="Key Risk 3:")
bullet(doc, "Competitive risk from SpaceX Starship for medium/heavy lift and from emerging small-lift competitors.", bold_prefix="Key Risk 4:")
bullet(doc, "Persistent negative free cash flow — FCF was approximately ($77.4M) in Q1 2026.", bold_prefix="Key Risk 5:")

add_section_heading(doc, "Catalysts", level=2)
bullet(doc, "Neutron first flight (targeted Q4 2026) — the single most important catalyst for re-rating.")
bullet(doc, "Nasdaq-100 Index inclusion (effective June 22, 2026) — expected to drive incremental passive fund inflows.")
bullet(doc, "Additional Golden Dome/SBI contract awards as program matures beyond demonstration phase.")
bullet(doc, "Quarterly EBITDA profitability milestone (potentially Q3 or Q4 2026 at current improvement trajectory).")
bullet(doc, "Electron booster reuse demonstration — would structurally lower unit costs and accelerate margin expansion.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PAGES 8-10 — VALUATION & ESTIMATES
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Valuation & Estimates")

add_section_heading(doc, "YoY Revenue Growth Reaccelerating", level=2)
add_chart(doc, "rklb_chart2_yoy_growth.png", 6.2,
          "Figure 10: YoY Revenue Growth by Quarter | Source: Rocket Lab earnings releases")

add_section_heading(doc, "Updated Estimates", level=2)
body(doc, "We raise our FY2026E revenue estimate to $940M (from prior $885M) based on strong Q1 results and above-consensus Q2 guidance. We also increase our FY2027E estimate to $1.25B, reflecting Neutron contribution beginning in H2 2027.")

est_data = [
    ["Metric",         "FY2024A",  "FY2025A",  "FY2026E (New)", "FY2026E (Old)", "FY2027E"],
    ["Revenue",        "$436.2M",  "$601.8M",  "$940M",         "$885M",         "$1.25B"],
    ["YoY Growth",     "+78%",     "+38%",     "+56%",          "+47%",          "+33%"],
    ["GAAP Gross Mgn", "~27%",     "~34%",     "~35%",          "~37%",          "~38%"],
    ["Adj. EBITDA",    "($91M)",   "($101M)",  "($30M)",        "Break-even",    "$50M+"],
    ["GAAP EPS",       "($0.42)",  "($0.37)",  "($0.22)",       "($0.25)",       "($0.08)"],
    ["Electron Lchs",  "16",       "21",       "~26–28",        "~24–26",        "~30–34"],
]
etbl = doc.add_table(rows=len(est_data), cols=6)
etbl.alignment = WD_TABLE_ALIGNMENT.CENTER
etbl.style = "Table Grid"
for r_idx, row in enumerate(est_data):
    for c_idx, val in enumerate(row):
        cell = etbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx % 2 == 0:
            set_cell_bg(cell, "EBF5FB")

body(doc, "Source: Rocket Lab filings (A = Actual), analyst estimates (E). Old estimates from Q4 FY2025 earnings update, Feb 2026.", italic=True, color_hex=GRAY_HEX, size=8.5)

add_section_heading(doc, "Valuation Framework", level=2)
body(doc, f"At the current price of {mkt['price']}, Rocket Lab trades at approximately {float(mkt['market_cap'].replace('~$','').replace('B',''))*1e9/940e6:.1f}x EV/FY2026E revenue — a premium reflective of the company's unique competitive positioning, >50% revenue growth, and improving margin trajectory. For a company growing revenue at 56% YoY with a pathway to positive EBITDA in 2027 and a $2.2B backlog, we believe this multiple is justified relative to aerospace and high-growth industrial comparables.")
body(doc, "Our 12-month price target of $120 is based on a blended approach:")
bullet(doc, "70% weight: 70x EV/FY2027E EBITDA of $50M+ (appropriate for high-growth aerospace transitioning to profitability)")
bullet(doc, "30% weight: 50x EV/FY2026E Revenue multiple of $940M (premium growth multiple)")
body(doc, f"This implies ~{(120/float(mkt['price'].replace('$',''))-1)*100:.0f}% upside from the current price. We acknowledge the premium valuation and the binary risk around Neutron's first flight, but believe the risk/reward remains favorable given the business fundamentals demonstrated in Q1.")

add_section_heading(doc, "Scenario Analysis", level=2)
scen_data = [
    ["Scenario",   "Methodology",           "Key Assumptions",                      "Implied Price",  "vs. Current"],
    ["Bear",       "40x EV/FY27E EBITDA",   "Neutron delayed to 2027; margin slip", "~$65",           f"~{(65/float(mkt['price'].replace('$',''))-1)*100:.0f}%"],
    ["Base",       "Blended (above)",        "Neutron Q4'26; $940M rev; EBITDA path","$120",           f"~+{(120/float(mkt['price'].replace('$',''))-1)*100:.0f}%"],
    ["Bull",       "80x EV/FY27E EBITDA",   "Neutron on time; $1B+ FY26; new wins", "~$160",          f"~+{(160/float(mkt['price'].replace('$',''))-1)*100:.0f}%"],
]
stbl = doc.add_table(rows=len(scen_data), cols=5)
stbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stbl.style = "Table Grid"
for r_idx, row in enumerate(scen_data):
    for c_idx, val in enumerate(row):
        cell = stbl.cell(r_idx, c_idx)
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
        if r_idx == 0:
            set_cell_bg(cell, BLUE_HEX)
            set_cell_font(cell, bold=True, color_hex=WHITE_HEX, size_pt=9)
        elif r_idx == 1:
            set_cell_bg(cell, "FDFEFE")
        elif r_idx == 2:
            set_cell_bg(cell, "EBF5FB")
        elif r_idx == 3:
            set_cell_bg(cell, "D5F5E3")

doc.add_paragraph()

p = doc.add_paragraph()
r_bg = p.add_run("  MAINTAIN BUY  |  PRICE TARGET: $120.00  ")
r_bg.bold = True
r_bg.font.size = Pt(11)
rv, gv, bv = hex_to_rgb(GREEN_HEX)
r_bg.font.color.rgb = RGBColor(rv, gv, bv)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

body(doc, "Q1 2026 was a landmark quarter for Rocket Lab — the company's first $200M revenue quarter, with all metrics beating expectations and the stock crossing $100 for the first time. The $2.2B backlog, expanding defense positioning (Golden Dome/SBI, HASTE), and near-term Nasdaq-100 inclusion all support a constructive outlook. While valuation is elevated and Neutron execution risk persists, we believe the combination of accelerating growth, margin expansion, and strategic positioning warrants maintaining our BUY rating. We raise our price target to $120 (from prior $35) to reflect the dramatically improved business scale and market positioning.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SOURCES & REFERENCES
# ══════════════════════════════════════════════════════════════════════════════

add_section_heading(doc, "Sources & References")

sources = [
    ("Rocket Lab Q1 FY2026 Earnings Press Release (May 7, 2026)",
     "https://investors.rocketlabcorp.com/news-releases/news-release-details/rocket-lab-announces-first-quarter-2026-financial-results"),
    ("Rocket Lab Form 8-K Q1 2026 (SEC EDGAR, May 7, 2026)",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000027/rklb-05072026ex991.htm"),
    ("Rocket Lab Form 10-Q Q1 2026 (SEC EDGAR, period ended March 31, 2026)",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000028/rklb-20260331.htm"),
    ("Rocket Lab Q1 2026 Earnings Call Transcript (Seeking Alpha, May 7, 2026)",
     "https://seekingalpha.com/article/4901108-rocket-lab-corporation-rklb-q1-2026-earnings-call-transcript"),
    ("Rocket Lab Q1 2026 Earnings Call Transcript (Motley Fool, May 8, 2026)",
     "https://www.fool.com/earnings/call-transcripts/2026/05/08/rocket-lab-rklb-q1-2026-earnings-transcript/"),
    ("RKLB Q1 2026 Earnings Analysis (CNBC, May 8, 2026)",
     "https://www.cnbc.com/2026/05/08/rocket-lab-rklb-q1-earnings-2026.html"),
    ("Rocket Lab Q4 & FY2025 Earnings Press Release (Feb 26, 2026) — Prior Quarter Reference",
     "https://www.sec.gov/Archives/edgar/data/0001819994/000181999426000012/rklb-02262026ex991.htm"),
    ("RKLB Revenue History (StockAnalysis.com)",
     "https://stockanalysis.com/stocks/rklb/revenue/"),
    ("Market Data: yfinance (live at report generation time)",
     "https://finance.yahoo.com/quote/RKLB"),
]

for title, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    add_hyperlink(p, title, url)

doc.add_paragraph()
body(doc, "Analyst Consensus: Bloomberg / FactSet / LSEG as of pre-earnings close May 7, 2026. Estimates marked 'E' are analyst projections. Price target and rating reflect analyst's independent assessment.", italic=True, color_hex=GRAY_HEX, size=8.5)
body(doc, "Disclosures: This report is for informational purposes only and does not constitute investment advice. The analyst and firm do not hold positions in RKLB. Past performance is not indicative of future results.", italic=True, color_hex=GRAY_HEX, size=8.5)

# ─── Save ─────────────────────────────────────────────────────────────────────
out_path = OUT + "RKLB_Q1_FY2026_Earnings_Update.docx"
doc.save(out_path)
print(f"\nEnglish report saved: {out_path}")
