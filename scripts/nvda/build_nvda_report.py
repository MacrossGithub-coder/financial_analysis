#!/usr/bin/env python3
"""Generate NVIDIA Q1 FY2027 Earnings Update Report (English)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "NVDA")
os.makedirs(OUT, exist_ok=True)

# ── Market Data ─────────────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Could not fetch market data: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("NVDA")
price_str = f"${mkt['price']}" if mkt['price'] != 'N/A' else 'N/A'
mcap_str = f"${mkt['market_cap']/1e9:.0f}B" if mkt['market_cap'] != 'N/A' else 'N/A'
hi52 = f"${mkt['52w_high']}" if mkt['52w_high'] != 'N/A' else 'N/A'
lo52 = f"${mkt['52w_low']}" if mkt['52w_low'] != 'N/A' else 'N/A'

# ── Helper Functions ────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_formatted_paragraph(doc, text, style="Normal", bold=False, font_size=11,
                            color=None, alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_bullet(doc, header, body, font_size=10.5):
    p = doc.add_paragraph()
    run_bullet = p.add_run("■ ")
    run_bullet.font.name = "Times New Roman"
    run_bullet.font.size = Pt(font_size)
    run_bullet.bold = True
    run_h = p.add_run(header)
    run_h.font.name = "Times New Roman"
    run_h.font.size = Pt(font_size)
    run_h.bold = True
    p.add_run("\n")
    run_b = p.add_run(body)
    run_b.font.name = "Times New Roman"
    run_b.font.size = Pt(font_size)
    p.paragraph_format.space_after = Pt(8)
    return p

def format_table(table, header_color="1B5E20"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.name = "Times New Roman"
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.name = "Times New Roman"

def add_source_line(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    run.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                          f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
                          f'<w:sz w:val="16"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'
                          f'<w:t>{text}</w:t></w:r></w:hyperlink>')
    paragraph._p.append(hyperlink)

# ── Build Document ──────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════════
# PAGE 1: EARNINGS SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_formatted_paragraph(doc, "NVIDIA CORPORATION (NVDA)", bold=True, font_size=16,
                        color=(0x1B, 0x5E, 0x20), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, "Q1 FY2027 EARNINGS UPDATE", bold=True, font_size=14,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_formatted_paragraph(doc, "Blackwell Ramp Drives Record $81.6B Revenue — Beat & Raise",
                        bold=True, font_size=12, color=(0x1B, 0x5E, 0x20),
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# Header info
header_data = [
    f"Date: June 6, 2026",
    f"Rating: MAINTAIN OUTPERFORM",
    f"Price (as of Jun 5, 2026): {price_str}",
    f"Price Target: $280 (Prior: $260) — RAISED",
    f"Market Cap: {mcap_str}  |  52-Week Range: {lo52} – {hi52}",
]
for h in header_data:
    add_formatted_paragraph(doc, h, font_size=10, space_after=2)

# Earnings Summary Box
add_formatted_paragraph(doc, "", space_after=4)
add_formatted_paragraph(doc, "EARNINGS SUMMARY", bold=True, font_size=12,
                        color=(0x1B, 0x5E, 0x20), space_after=4, space_before=6)

summary_table = doc.add_table(rows=4, cols=4)
summary_table.style = "Light Grid Accent 1"
headers = ["Metric", "Reported", "Consensus", "Variance"]
for i, h in enumerate(headers):
    summary_table.rows[0].cells[i].text = h
data_rows = [
    ["Revenue", "$81.6B", "$78.8B", "+$2.8B (+3.6%)"],
    ["Non-GAAP EPS", "$1.87", "$1.77", "+$0.10 (+5.7%)"],
    ["Non-GAAP Gross Margin", "75.0%", "74.8%", "+20 bps"],
]
for r, row_data in enumerate(data_rows, 1):
    for c, val in enumerate(row_data):
        summary_table.rows[r].cells[c].text = val
format_table(summary_table)
add_source_line(doc, "Source: Company earnings release (May 20, 2026); Bloomberg consensus as of May 19, 2026.")

# Key Takeaways
add_formatted_paragraph(doc, "", space_after=2)

add_bullet(doc,
    "Record revenue of $81.6B beat consensus by 3.6%, driven by Blackwell system demand across hyperscalers",
    "Q1 FY2027 revenue surged 85% YoY and 20% QoQ to $81.6B, exceeding consensus of $78.8B by $2.8B. "
    "Data Center revenue reached $75.2B (+92% YoY), representing 92% of total revenue. Management noted "
    "Blackwell deployment is \"the fastest product ramp in company history,\" with GB300 NVL72 demand particularly "
    "strong. Hyperscale customers contributed $38B and ACIE (AI Clouds, Industrial, Enterprise) added $37B, "
    "up 31% QoQ, demonstrating broadening customer diversification beyond the top cloud providers.")

add_bullet(doc,
    "Non-GAAP gross margin expanded to 75.0%, fully recovering from H20 export-control charges",
    "Non-GAAP gross margin of 75.0% was up 1,400 bps YoY from the H20-impacted Q1 FY2026 (61.0%) and "
    "roughly flat sequentially vs. Q4 FY2026 (75.2%). Margin stability amid the largest product transition "
    "in company history (Hopper to Blackwell) demonstrates strong pricing power and manufacturing yield "
    "improvements. Management guided Q2 FY2027 non-GAAP gross margin at 75.0% (+/-50 bps), suggesting "
    "sustained profitability through the Blackwell ramp.")

add_bullet(doc,
    "Q2 FY2027 guidance of $91.0B significantly above Street expectations, signaling continued acceleration",
    "Management guided Q2 FY2027 revenue to $91.0B (+/-2%), implying ~11.5% QoQ growth and well above "
    "the pre-earnings Street estimate of ~$85B. The guidance excludes any Data Center compute revenue from "
    "China due to ongoing export restrictions. Combined with the $80B new share repurchase authorization and "
    "a 25x dividend increase to $0.25/share quarterly, the capital return framework underscores management's "
    "confidence in sustainable earnings power.")

add_bullet(doc,
    "Maintaining Outperform, raising PT to $280 on higher FY2027E/FY2028E estimates",
    "We raise our FY2027E revenue estimate to $365B (from $340B) and Non-GAAP EPS to $8.80 (from $8.20) "
    "reflecting the Q1 beat and raised guidance. Our $280 price target (from $260) is based on 32x our "
    "FY2028E Non-GAAP EPS of $10.50, supported by NVIDIA's dominant AI infrastructure position, expanding "
    "TAM across inference, sovereign AI, and edge robotics, and sustained gross margin above 74%.")

# Updated Estimates Table
add_formatted_paragraph(doc, "UPDATED FINANCIAL ESTIMATES", bold=True, font_size=11,
                        color=(0x1B, 0x5E, 0x20), space_after=4, space_before=10)

est_table = doc.add_table(rows=8, cols=5)
est_table.style = "Light Grid Accent 1"
est_headers = ["Metric", "FY2027E (Old)", "FY2027E (New)", "Change", "FY2028E (New)"]
for i, h in enumerate(est_headers):
    est_table.rows[0].cells[i].text = h
est_data = [
    ["Revenue ($B)", "$340.0", "$365.0", "+7.4%", "$440.0"],
    ["Revenue Growth (%)", "57.4%", "69.0%", "+1,160 bps", "20.5%"],
    ["Non-GAAP Gross Margin", "74.5%", "75.0%", "+50 bps", "74.5%"],
    ["Non-GAAP Op Income ($B)", "$225.0", "$245.0", "+8.9%", "$290.0"],
    ["Non-GAAP Op Margin", "66.2%", "67.1%", "+90 bps", "65.9%"],
    ["Non-GAAP EPS ($)", "$8.20", "$8.80", "+7.3%", "$10.50"],
    ["P/E (x) at current price", "26.8x", "25.0x", "-1.8x", "20.9x"],
]
for r, row_data in enumerate(est_data, 1):
    for c, val in enumerate(row_data):
        est_table.rows[r].cells[c].text = val
format_table(est_table)
add_source_line(doc, 'Note: "E" = Estimate. Old estimates from Q4 FY2026 earnings update. Source: Company data, analyst estimates.')

# ════════════════════════════════════════════════════════════════════════════
# PAGES 2-3: DETAILED RESULTS ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "DETAILED RESULTS ANALYSIS", bold=True, font_size=14,
                        color=(0x1B, 0x5E, 0x20), space_after=8)

# Revenue Analysis
add_formatted_paragraph(doc, "Revenue Analysis", bold=True, font_size=12, space_after=6)

add_formatted_paragraph(doc,
    "Revenue of $81.6B grew 85% YoY and 20% QoQ, representing the third consecutive quarter of YoY "
    "acceleration and the fourteenth straight quarter of sequential growth. The $13.5B sequential increase "
    "was a new company record. The beat vs. consensus of $78.8B was primarily driven by stronger-than-expected "
    "Data Center demand, particularly from hyperscale and sovereign AI customers deploying Blackwell systems.",
    font_size=10.5, space_after=6)

# Revenue Breakdown Table
rev_table = doc.add_table(rows=5, cols=7)
rev_table.style = "Light Grid Accent 1"
rev_headers2 = ["Segment", "Q1 FY26", "Q2 FY26", "Q3 FY26", "Q4 FY26", "Q1 FY27", "YoY Chg"]
for i, h in enumerate(rev_headers2):
    rev_table.rows[0].cells[i].text = h
rev_data = [
    ["Total Revenue ($B)", "$44.1", "$46.7", "$57.0", "$68.1", "$81.6", "+85%"],
    ["Data Center ($B)", "$39.1", "$42.0", "$51.2", "$62.3", "$75.2", "+92%"],
    ["Edge Computing ($B)*", "$5.0", "$4.7", "$5.8", "$5.8", "$6.4", "+29%"],
    ["DC % of Total", "89%", "90%", "90%", "91%", "92%", "+3pp"],
]
for r, row_data in enumerate(rev_data, 1):
    for c, val in enumerate(row_data):
        rev_table.rows[r].cells[c].text = val
format_table(rev_table)
add_source_line(doc, "*Edge Computing includes Gaming, Professional Visualization, Automotive & Robotics "
                     "(new segment effective Q1 FY27). Prior periods reclassified for comparability. "
                     "Source: Company filings.")

# Chart 1
add_formatted_paragraph(doc, "", space_after=2)
add_formatted_paragraph(doc, "Figure 1 — Quarterly Revenue Progression", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
doc.add_picture(os.path.join(OUT, "nvda_chart1_quarterly_revenue.png"), width=Inches(5.8))
last_paragraph = doc.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings; consensus estimate from Bloomberg (May 19, 2026).")

# Data Center Deep Dive
add_formatted_paragraph(doc, "Data Center Deep Dive", bold=True, font_size=12, space_after=6, space_before=10)

add_formatted_paragraph(doc,
    "Data Center revenue of $75.2B (+92% YoY, +21% QoQ) was the primary growth driver. Within Data Center, "
    "compute revenue reached $60.4B (+77% YoY, +18% QoQ), driven by accelerating Blackwell GPU deployments. "
    "Networking revenue nearly tripled YoY to $14.8B (+199% YoY, +35% QoQ), reflecting the growing importance "
    "of high-bandwidth interconnects (NVLink, InfiniBand, Spectrum-X Ethernet) as GPU clusters scale to tens of "
    "thousands of GPUs.",
    font_size=10.5, space_after=6)

add_formatted_paragraph(doc,
    "By customer type, hyperscale customers contributed approximately $38B (~50% of DC revenue), while the ACIE "
    "segment (AI Clouds, Industrial, Enterprise) generated $37B, up 31% QoQ. The rapid ACIE growth signals that "
    "AI workload demand is broadening beyond the top hyperscalers into sovereign AI programs, enterprise deployments, "
    "and AI-native cloud providers. CFO Colette Kress noted that \"demand for GB300 NVL72 was particularly strong\" "
    "with frontier builders deploying \"hundreds and thousands\" of Blackwell GPUs.",
    font_size=10.5, space_after=6)

# Profitability Analysis
doc.add_page_break()
add_formatted_paragraph(doc, "Profitability Analysis", bold=True, font_size=12, space_after=6)

add_formatted_paragraph(doc,
    "Non-GAAP gross margin of 75.0% was roughly flat vs. Q4 FY2026 (75.2%) and up dramatically from "
    "Q1 FY2026's H20-impacted 61.0%. The margin stability during Blackwell's fastest-ever product ramp "
    "demonstrates strong pricing discipline and improving yields. GAAP gross margin was 74.9%. "
    "Operating expenses totaled $7.6B (GAAP), with R&D spending at $6.3B — reflecting continued heavy "
    "investment in next-generation GPU architectures (Rubin/Vera) and software platforms.",
    font_size=10.5, space_after=6)

margin_table = doc.add_table(rows=5, cols=6)
margin_table.style = "Light Grid Accent 1"
m_headers = ["Metric", "Q1 FY26", "Q3 FY26", "Q4 FY26", "Q1 FY27", "YoY Chg"]
for i, h in enumerate(m_headers):
    margin_table.rows[0].cells[i].text = h
m_data = [
    ["GAAP Gross Margin", "60.5%", "73.4%", "75.0%", "74.9%", "+1,440 bps"],
    ["Non-GAAP Gross Margin", "61.0%", "73.6%", "75.2%", "75.0%", "+1,400 bps"],
    ["GAAP Operating Margin", "49.0%", "63.2%", "65.0%", "65.6%", "+1,660 bps"],
    ["Non-GAAP Operating Margin", "52.8%", "64.9%", "67.7%", "65.9%", "+1,310 bps"],
]
for r, row_data in enumerate(m_data, 1):
    for c, val in enumerate(row_data):
        margin_table.rows[r].cells[c].text = val
format_table(margin_table)
add_source_line(doc, "Source: Company filings. Q1 FY26 margins impacted by $4.5B H20 inventory charge.")

# Chart 2 & 3
add_formatted_paragraph(doc, "Figure 2 — Quarterly EPS Progression", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart2_quarterly_eps.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings; consensus estimate from Bloomberg (May 19, 2026).")

add_formatted_paragraph(doc, "Figure 3 — Quarterly Margin Trends", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart3_margin_trends.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings.")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 4-5: KEY METRICS & GUIDANCE
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "KEY METRICS & GUIDANCE", bold=True, font_size=14,
                        color=(0x1B, 0x5E, 0x20), space_after=8)

# Key Operating Metrics
add_formatted_paragraph(doc, "Key Operating Metrics", bold=True, font_size=12, space_after=6)

kpi_table = doc.add_table(rows=7, cols=4)
kpi_table.style = "Light Grid Accent 1"
kpi_headers = ["Metric", "Q1 FY27", "Q4 FY26", "YoY / QoQ"]
for i, h in enumerate(kpi_headers):
    kpi_table.rows[0].cells[i].text = h
kpi_data = [
    ["DC Compute Revenue ($B)", "$60.4", "$51.2", "+18% QoQ"],
    ["DC Networking Revenue ($B)", "$14.8", "$11.0", "+35% QoQ"],
    ["Free Cash Flow ($B)", "$49.0", "$34.9", "+40% QoQ"],
    ["Operating Cash Flow ($B)", "~$51.0", "$36.2", "+41% QoQ"],
    ["Share Repurchases ($B)", "~$20.0", "~$11.0", "—"],
    ["Diluted Share Count (B)", "~24.4", "24.4", "Flat"],
]
for r, row_data in enumerate(kpi_data, 1):
    for c, val in enumerate(row_data):
        kpi_table.rows[r].cells[c].text = val
format_table(kpi_table)
add_source_line(doc, "Source: Company earnings release (May 20, 2026); 10-Q filing (May 28, 2026).")

# Charts 4, 5
add_formatted_paragraph(doc, "Figure 4 — Revenue: Data Center vs. Other Segments", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart4_segment_revenue.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings.")

add_formatted_paragraph(doc, "Figure 5 — Data Center Revenue & YoY Growth", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart5_dc_revenue_growth.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings.")

# Guidance Analysis
doc.add_page_break()
add_formatted_paragraph(doc, "Q2 FY2027 Guidance Analysis", bold=True, font_size=12, space_after=6)

add_formatted_paragraph(doc,
    "Management guided Q2 FY2027 revenue to approximately $91.0B (+/-2%), implying $89.2B-$92.8B. "
    "At the midpoint, this represents ~11.5% QoQ growth and significantly exceeds the pre-earnings "
    "Street estimate of ~$85B. The guidance excludes any Data Center compute revenue from China due "
    "to continued U.S. export restrictions — a notable constraint that makes the beat even more "
    "impressive on an organic basis.",
    font_size=10.5, space_after=6)

guide_table = doc.add_table(rows=5, cols=4)
guide_table.style = "Light Grid Accent 1"
g_headers = ["Metric", "Q2 FY27E Guidance", "Q1 FY27 Actual", "QoQ Implied"]
for i, h in enumerate(g_headers):
    guide_table.rows[0].cells[i].text = h
g_data = [
    ["Revenue", "~$91.0B (+/-2%)", "$81.6B", "+11.5%"],
    ["GAAP Gross Margin", "74.9% (+/-50 bps)", "74.9%", "Flat"],
    ["Non-GAAP Gross Margin", "75.0% (+/-50 bps)", "75.0%", "Flat"],
    ["Non-GAAP OpEx", "~$8.3B", "~$7.5B", "+10.7%"],
]
for r, row_data in enumerate(g_data, 1):
    for c, val in enumerate(row_data):
        guide_table.rows[r].cells[c].text = val
format_table(guide_table)
add_source_line(doc, "Source: Company Q1 FY2027 earnings release and CFO commentary (May 20, 2026).")

add_formatted_paragraph(doc,
    "The guidance is notable for several reasons: (1) it implies continued sequential acceleration "
    "despite an increasingly large revenue base; (2) gross margin guidance of ~75% suggests no margin "
    "headwinds from the Blackwell production ramp; (3) OpEx rising to ~$8.3B reflects continued heavy "
    "investment in Rubin/Vera next-gen architectures and CUDA software ecosystem expansion; and (4) the "
    "explicit exclusion of China compute revenue provides a clean organic growth picture.",
    font_size=10.5, space_after=6)

# Chart 6
add_formatted_paragraph(doc, "Figure 6 — Q1 FY2027 Revenue Beat: Consensus to Reported", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart6_beat_waterfall.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings; Bloomberg consensus (May 19, 2026).")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 6-7: UPDATED INVESTMENT THESIS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "UPDATED INVESTMENT THESIS", bold=True, font_size=14,
                        color=(0x1B, 0x5E, 0x20), space_after=8)

add_formatted_paragraph(doc, "Thesis Impact Assessment", bold=True, font_size=12, space_after=6)

add_bullet(doc,
    "Thesis 1: NVIDIA is the primary beneficiary of the AI infrastructure buildout — STRENGTHENED",
    "Q1 FY2027 results emphatically reinforce this thesis. Data Center revenue of $75.2B (+92% YoY) "
    "demonstrates that the AI infrastructure investment cycle continues to accelerate rather than plateau. "
    "The broadening customer base — from hyperscalers ($38B) to ACIE ($37B) — suggests the AI compute "
    "demand is not concentrated in a handful of buyers but expanding across sovereign AI programs, enterprise "
    "deployments, and specialized AI cloud providers. Blackwell's record ramp further cements NVIDIA's "
    "technological leadership, with GB300 NVL72 demand described as \"particularly strong.\"")

add_bullet(doc,
    "Thesis 2: Networking becomes the next major growth vector — STRENGTHENED",
    "Networking revenue of $14.8B (+199% YoY, +35% QoQ) validates our thesis that as GPU clusters scale "
    "to tens of thousands of accelerators, high-bandwidth interconnects become critical. NVLink, InfiniBand, "
    "and Spectrum-X Ethernet solutions are increasingly sold as part of integrated Blackwell systems rather "
    "than standalone products. Networking now represents ~20% of Data Center revenue, up from ~13% a year ago, "
    "demonstrating meaningful TAM expansion beyond GPUs.")

add_bullet(doc,
    "Thesis 3: Gross margins sustained above 73% through product transitions — UNCHANGED",
    "Non-GAAP gross margin of 75.0% remains well above our 73% floor thesis. The stability during the "
    "Hopper-to-Blackwell transition is particularly notable, as new product ramps historically create "
    "margin volatility. Management's Q2 guidance of 75.0% (+/-50 bps) provides visibility into sustained "
    "margins. However, we maintain our conservative 74.5% assumption for the full FY2027 to account for "
    "potential mix shifts as lower-priced inference GPUs scale.")

# Risks Update
add_formatted_paragraph(doc, "Risk Assessment Update", bold=True, font_size=12, space_after=6, space_before=10)

add_formatted_paragraph(doc,
    "Key risks remain largely unchanged but warrant monitoring: (1) Export control risk — the exclusion "
    "of China DC compute revenue from guidance highlights ongoing regulatory headwinds, though NVIDIA has "
    "demonstrated ability to grow robustly without China; (2) Customer concentration — while ACIE growth is "
    "encouraging, hyperscalers still drive ~50% of DC revenue, creating potential cyclical risk if cloud "
    "CapEx slows; (3) Competition — AMD's MI400 and custom ASICs from Google (TPU), Amazon (Trainium), and "
    "Microsoft (Maia) represent long-term competitive threats, though near-term share losses appear minimal; "
    "(4) Valuation — at ~25x FY2027E P/E, the stock prices in significant continued growth, leaving limited "
    "room for execution missteps.",
    font_size=10.5, space_after=6)

# Chart 7, 8
add_formatted_paragraph(doc, "Figure 7 — Quarterly Free Cash Flow", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart7_free_cash_flow.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings.")

add_formatted_paragraph(doc, "Figure 8 — Q1 FY2027 Data Center Revenue Mix", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart8_dc_mix.png"), width=Inches(4.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company earnings release (May 20, 2026).")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 8-10: VALUATION & ESTIMATES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "VALUATION & UPDATED ESTIMATES", bold=True, font_size=14,
                        color=(0x1B, 0x5E, 0x20), space_after=8)

add_formatted_paragraph(doc, "Valuation Framework", bold=True, font_size=12, space_after=6)

add_formatted_paragraph(doc,
    f"Our $280 price target (raised from $260) is based on a blended valuation approach:\n\n"
    f"DCF Analysis (40% weight): Updated DCF with FY2027E-FY2032E cash flows yields a fair value of $295. "
    f"Key assumptions: revenue CAGR of 22% over 5 years, terminal EBIT margin of 62%, WACC of 10.5%, "
    f"terminal growth of 3.5%.\n\n"
    f"P/E Multiple (40% weight): 32x our FY2028E Non-GAAP EPS of $10.50 yields $336. We apply a 20% "
    f"discount for execution risk, yielding $269.\n\n"
    f"EV/EBITDA (20% weight): 25x our FY2028E EBITDA of $310B yields $250 per share after net cash.\n\n"
    f"Blended target: (40% x $295) + (40% x $269) + (20% x $250) = $276, rounded to $280.",
    font_size=10.5, space_after=6)

# Updated Estimates Detail
add_formatted_paragraph(doc, "Detailed Estimate Updates", bold=True, font_size=12, space_after=6, space_before=8)

det_table = doc.add_table(rows=11, cols=5)
det_table.style = "Light Grid Accent 1"
det_headers = ["Metric", "FY27E (Old)", "FY27E (New)", "Change", "FY28E (New)"]
for i, h in enumerate(det_headers):
    det_table.rows[0].cells[i].text = h
det_data = [
    ["Revenue ($B)", "$340.0", "$365.0", "+7.4%", "$440.0"],
    ["  Data Center ($B)", "$310.0", "$335.0", "+8.1%", "$400.0"],
    ["  Edge Computing ($B)", "$30.0", "$30.0", "—", "$40.0"],
    ["Gross Profit ($B)", "$253.3", "$273.8", "+8.1%", "$327.8"],
    ["Non-GAAP Gross Margin", "74.5%", "75.0%", "+50 bps", "74.5%"],
    ["Non-GAAP OpEx ($B)", "$35.0", "$35.0", "—", "$42.0"],
    ["Non-GAAP Op Income ($B)", "$225.0", "$245.0", "+8.9%", "$290.0"],
    ["Non-GAAP Op Margin", "66.2%", "67.1%", "+90 bps", "65.9%"],
    ["Non-GAAP EPS ($)", "$8.20", "$8.80", "+7.3%", "$10.50"],
    ["Free Cash Flow ($B)", "$195.0", "$215.0", "+10.3%", "$250.0"],
]
for r, row_data in enumerate(det_data, 1):
    for c, val in enumerate(row_data):
        det_table.rows[r].cells[c].text = val
format_table(det_table)
add_source_line(doc, "Source: Analyst estimates.")

# Charts 9, 10
add_formatted_paragraph(doc, "Figure 9 — FY2027E Estimate Revisions", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart9_estimate_revisions.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Analyst estimates.")

add_formatted_paragraph(doc, "Figure 10 — Quarterly Revenue Sequential Growth Rate", bold=True, font_size=10,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "nvda_chart10_sequential_growth.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_line(doc, "Source: Company filings.")

# ════════════════════════════════════════════════════════════════════════════
# SOURCES & REFERENCES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_formatted_paragraph(doc, "SOURCES & REFERENCES", bold=True, font_size=14,
                        color=(0x1B, 0x5E, 0x20), space_after=8)

add_formatted_paragraph(doc, "Earnings Materials (Q1 FY2027):", bold=True, font_size=10.5, space_after=4)

src_items = [
    ("Earnings Release (May 20, 2026)",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-first-quarter-fiscal-2027"),
    ("Form 8-K — Press Release (Filed May 20, 2026)",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27pr.htm"),
    ("CFO Commentary (May 20, 2026)",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27cfocommentary.htm"),
    ("Form 10-Q (Filed May 28, 2026)",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000052/nvda-20260426.htm"),
    ("Q1 FY2027 Earnings Call Transcript (May 20, 2026)",
     "https://seekingalpha.com/article/4907259-nvidia-corporation-nvda-q1-2027-earnings-call-transcript"),
]

for title, url in src_items:
    p = doc.add_paragraph()
    run = p.add_run("• ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

add_formatted_paragraph(doc, "", space_after=4)
add_formatted_paragraph(doc, "Prior Period References:", bold=True, font_size=10.5, space_after=4)

prior_items = [
    ("Q4 FY2026 Earnings Release (February 26, 2026)",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-fourth-quarter-and-fiscal-2026"),
    ("Q4 FY2026 Form 8-K — Press Release",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000019/q4fy26pr.htm"),
    ("Q3 FY2026 Earnings Release (November 2025)",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-third-quarter-fiscal-2026"),
]
for title, url in prior_items:
    p = doc.add_paragraph()
    run = p.add_run("• ")
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

add_formatted_paragraph(doc, "", space_after=4)
add_formatted_paragraph(doc, "Consensus & Market Data:", bold=True, font_size=10.5, space_after=4)
add_formatted_paragraph(doc,
    "• Bloomberg consensus estimates as of market close May 19, 2026\n"
    "• Yahoo Finance / yfinance for real-time market data\n"
    "• NVIDIA Investor Relations: https://investor.nvidia.com",
    font_size=10, space_after=6)

# Disclaimer
add_formatted_paragraph(doc, "", space_after=6)
add_formatted_paragraph(doc, "DISCLAIMER", bold=True, font_size=9, color=(0x75, 0x75, 0x75), space_after=2)
add_formatted_paragraph(doc,
    "This report is for informational purposes only and does not constitute investment advice. "
    "The analyst is not a licensed investment advisor. All estimates and price targets are based on "
    "publicly available information and represent the analyst's independent assessment. Past performance "
    "is not indicative of future results. Investors should conduct their own due diligence before making "
    "investment decisions.",
    font_size=8, color=(0x99, 0x99, 0x99), space_after=4)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.join(OUT, "NVDA_Q1_FY2027_Earnings_Update.docx")
doc.save(output_path)
print(f"Report saved: {output_path}")
