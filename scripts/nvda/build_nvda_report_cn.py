#!/usr/bin/env python3
"""Generate NVIDIA Q1 FY2027 Earnings Update Report (Chinese / 中文版)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "NVDA")
os.makedirs(OUT, exist_ok=True)

# ── Market Data ─────────────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Could not fetch market data: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("NVDA")
price_str = f"${mkt['price']}" if mkt['price'] != 'N/A' else 'N/A'
mcap_str = f"${mkt['market_cap']/1e9:.0f}B" if mkt['market_cap'] != 'N/A' else 'N/A'
hi52 = f"${mkt['52w_high']}" if mkt['52w_high'] != 'N/A' else 'N/A'
lo52 = f"${mkt['52w_low']}" if mkt['52w_low'] != 'N/A' else 'N/A'

# ── Helper Functions ────────────────────────────────────────────────────────
SONG = "宋体"
HEI  = "黑体"

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cn_font(run, cn_font=SONG, en_font="Times New Roman", size=11):
    run.font.name = en_font
    run.font.size = Pt(size)
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)

def add_cn_paragraph(doc, text, bold=False, font_size=11, color=None,
                     alignment=None, space_after=6, space_before=0, heading_font=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    cn = HEI if heading_font else SONG
    set_cn_font(run, cn_font=cn, size=font_size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)
    if alignment:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_cn_bullet(doc, header, body, font_size=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run("■ ")
    set_cn_font(r1, cn_font=HEI, size=font_size)
    r1.bold = True
    r2 = p.add_run(header)
    set_cn_font(r2, cn_font=HEI, size=font_size)
    r2.bold = True
    p.add_run("\n")
    r3 = p.add_run(body)
    set_cn_font(r3, cn_font=SONG, size=font_size)
    p.paragraph_format.space_after = Pt(8)
    return p

def format_table_cn(table, header_color="1B5E20"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, header_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.size = Pt(9)
                set_cn_font(r, cn_font=HEI, size=9)
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    set_cn_font(r, cn_font=SONG, size=9)

def add_source_cn(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn_font(run, cn_font=SONG, size=8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
    run.italic = True
    p.paragraph_format.space_after = Pt(4)
    return p

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
                          f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
                          f'<w:sz w:val="16"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'
                          f'<w:t>{text}</w:t></w:r></w:hyperlink>')
    paragraph._p.append(hyperlink)

# ── Build Document ──────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(11)
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════════
# 第1页：业绩摘要
# ════════════════════════════════════════════════════════════════════════════
add_cn_paragraph(doc, "英伟达 (NVDA)", bold=True, font_size=16,
                 color=(0x1B, 0x5E, 0x20), alignment=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=2, heading_font=True)
add_cn_paragraph(doc, "2027财年第一季度业绩更新报告", bold=True, font_size=14,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, heading_font=True)
add_cn_paragraph(doc, "Blackwell加速交付推动收入创纪录816亿美元 — 超预期并上调指引",
                 bold=True, font_size=12, color=(0x1B, 0x5E, 0x20),
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, heading_font=True)

header_data = [
    f"日期：2026年6月6日",
    f"评级：维持 跑赢大盘 (OUTPERFORM)",
    f"股价（截至2026年6月5日）：{price_str}",
    f"目标价：$280（此前$260）— 上调",
    f"市值：{mcap_str}  |  52周区间：{lo52} – {hi52}",
]
for h in header_data:
    add_cn_paragraph(doc, h, font_size=10, space_after=2)

add_cn_paragraph(doc, "", space_after=4)
add_cn_paragraph(doc, "业绩摘要", bold=True, font_size=12,
                 color=(0x1B, 0x5E, 0x20), space_after=4, space_before=6, heading_font=True)

summary_table = doc.add_table(rows=4, cols=4)
summary_table.style = "Light Grid Accent 1"
s_headers = ["指标", "实际值", "市场共识", "差异"]
for i, h in enumerate(s_headers):
    summary_table.rows[0].cells[i].text = h
s_data = [
    ["收入", "$816亿", "$788亿", "+$28亿 (+3.6%)"],
    ["Non-GAAP每股收益", "$1.87", "$1.77", "+$0.10 (+5.7%)"],
    ["Non-GAAP毛利率", "75.0%", "74.8%", "+20个基点"],
]
for r, row_data in enumerate(s_data, 1):
    for c, val in enumerate(row_data):
        summary_table.rows[r].cells[c].text = val
format_table_cn(summary_table)
add_source_cn(doc, "资料来源：公司业绩发布（2026年5月20日）；Bloomberg市场共识（2026年5月19日收盘）。")

add_cn_paragraph(doc, "", space_after=2)

add_cn_bullet(doc,
    "收入创纪录816亿美元，超市场共识3.6%，受益于Blackwell系统在超大规模客户中的强劲需求",
    "2027财年第一季度收入同比增长85%、环比增长20%至816亿美元，超过市场共识788亿美元约28亿美元。"
    "数据中心收入达到752亿美元（同比+92%），占总收入的92%。管理层指出Blackwell部署是\"公司历史上"
    "最快的产品爬坡\"，GB300 NVL72需求尤为强劲。超大规模客户贡献约380亿美元，ACIE（AI云、工业及"
    "企业）贡献370亿美元（环比+31%），表明客户基础正从头部云厂商向更广泛市场扩展。")

add_cn_bullet(doc,
    "Non-GAAP毛利率扩展至75.0%，完全恢复至H20出口管制影响前水平",
    "Non-GAAP毛利率75.0%，较受H20影响的2026财年第一季度（61.0%）同比提升1,400个基点，"
    "与2026财年第四季度（75.2%）基本持平。在公司历史上规模最大的产品代际切换（Hopper→Blackwell）"
    "期间保持毛利率稳定，充分展示了定价能力和制造良率的持续改善。管理层指引第二季度Non-GAAP毛利率"
    "为75.0%（±50个基点），预示盈利能力将在Blackwell爬坡期间持续维持。")

add_cn_bullet(doc,
    "第二季度指引约910亿美元，大幅超出市场预期，信号表明增长仍在加速",
    "管理层指引2027财年第二季度收入约910亿美元（±2%），环比增长约11.5%，显著高于财报前市场预期"
    "的约850亿美元。该指引不包含来自中国的数据中心计算收入（受出口管制影响）。同时宣布800亿美元"
    "新增股票回购授权和将季度股息大幅提高至每股0.25美元（提升25倍），资本回报框架凸显管理层"
    "对持续盈利能力的信心。")

add_cn_bullet(doc,
    "维持跑赢大盘评级，目标价上调至$280，基于上修的FY2027E/FY2028E预测",
    "我们将FY2027E收入预测上调至3,650亿美元（此前3,400亿美元），Non-GAAP每股收益上调至$8.80"
    "（此前$8.20），反映第一季度超预期表现和上调的指引。目标价$280（此前$260）基于FY2028E "
    "Non-GAAP每股收益$10.50的32倍市盈率，依据英伟达在AI基础设施领域的主导地位、推理/主权AI/"
    "边缘机器人等领域不断扩大的可触达市场，以及毛利率持续高于74%的支撑。")

# Updated Estimates Table
add_cn_paragraph(doc, "更新后的财务预测", bold=True, font_size=11,
                 color=(0x1B, 0x5E, 0x20), space_after=4, space_before=10, heading_font=True)

est_table = doc.add_table(rows=8, cols=5)
est_table.style = "Light Grid Accent 1"
est_headers = ["指标", "FY2027E（旧）", "FY2027E（新）", "变动", "FY2028E（新）"]
for i, h in enumerate(est_headers):
    est_table.rows[0].cells[i].text = h
est_data = [
    ["收入（$B）", "$340.0", "$365.0", "+7.4%", "$440.0"],
    ["收入增长率", "57.4%", "69.0%", "+1,160 bps", "20.5%"],
    ["Non-GAAP毛利率", "74.5%", "75.0%", "+50 bps", "74.5%"],
    ["Non-GAAP营业利润（$B）", "$225.0", "$245.0", "+8.9%", "$290.0"],
    ["Non-GAAP营业利润率", "66.2%", "67.1%", "+90 bps", "65.9%"],
    ["Non-GAAP每股收益", "$8.20", "$8.80", "+7.3%", "$10.50"],
    ["市盈率（倍）", "26.8x", "25.0x", "-1.8x", "20.9x"],
]
for r, row_data in enumerate(est_data, 1):
    for c, val in enumerate(row_data):
        est_table.rows[r].cells[c].text = val
format_table_cn(est_table)
add_source_cn(doc, '注："E"=预测。旧预测来自FY2026第四季度业绩更新。资料来源：公司数据、分析师预测。')

# ════════════════════════════════════════════════════════════════════════════
# 第2-3页：详细业绩分析
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_cn_paragraph(doc, "详细业绩分析", bold=True, font_size=14,
                 color=(0x1B, 0x5E, 0x20), space_after=8, heading_font=True)

add_cn_paragraph(doc, "收入分析", bold=True, font_size=12, space_after=6, heading_font=True)
add_cn_paragraph(doc,
    "收入816亿美元，同比增长85%、环比增长20%，为连续第三个季度同比增速加快，也是连续第十四个季度"
    "环比增长。135亿美元的环比增量创下公司新纪录。超出市场共识788亿美元的原因主要是数据中心需求"
    "强于预期，尤其是超大规模客户和主权AI客户部署Blackwell系统的需求。",
    font_size=10.5, space_after=6)

rev_table = doc.add_table(rows=5, cols=7)
rev_table.style = "Light Grid Accent 1"
rev_h = ["业务板块", "Q1 FY26", "Q2 FY26", "Q3 FY26", "Q4 FY26", "Q1 FY27", "同比变化"]
for i, h in enumerate(rev_h):
    rev_table.rows[0].cells[i].text = h
rev_data = [
    ["总收入 ($B)", "$44.1", "$46.7", "$57.0", "$68.1", "$81.6", "+85%"],
    ["数据中心 ($B)", "$39.1", "$42.0", "$51.2", "$62.3", "$75.2", "+92%"],
    ["边缘计算 ($B)*", "$5.0", "$4.7", "$5.8", "$5.8", "$6.4", "+29%"],
    ["数据中心占比", "89%", "90%", "90%", "91%", "92%", "+3pp"],
]
for r, row_data in enumerate(rev_data, 1):
    for c, val in enumerate(row_data):
        rev_table.rows[r].cells[c].text = val
format_table_cn(rev_table)
add_source_cn(doc, "*边缘计算包括游戏、专业可视化、汽车与机器人业务（Q1 FY27起新分部）。此前期间已重新分类以便比较。"
                   "资料来源：公司财报。")

add_cn_paragraph(doc, "图1 — 季度收入走势", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart1_quarterly_revenue.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报；Bloomberg市场共识（2026年5月19日）。")

# DC Deep Dive
add_cn_paragraph(doc, "数据中心业务深度分析", bold=True, font_size=12, space_after=6,
                 space_before=10, heading_font=True)
add_cn_paragraph(doc,
    "数据中心收入752亿美元（同比+92%，环比+21%）是核心增长引擎。其中计算收入达604亿美元"
    "（同比+77%，环比+18%），受Blackwell GPU加速部署驱动。网络收入同比增长近两倍至148亿美元"
    "（同比+199%，环比+35%），反映出随着GPU集群规模扩展至数万片加速器，高带宽互联（NVLink、"
    "InfiniBand、Spectrum-X以太网）变得日益关键。",
    font_size=10.5, space_after=6)
add_cn_paragraph(doc,
    "按客户类型划分，超大规模客户贡献约380亿美元（占数据中心收入的约50%），ACIE板块（AI云、"
    "工业及企业）贡献370亿美元（环比+31%）。ACIE的快速增长表明AI工作负载需求正从头部超大规模"
    "客户向主权AI项目、企业部署和AI原生云提供商等更广泛领域扩展。CFO Colette Kress指出"
    "\"GB300 NVL72需求尤为强劲\"，前沿建设者正部署\"数百和数千片\"Blackwell GPU。",
    font_size=10.5, space_after=6)

# Profitability
doc.add_page_break()
add_cn_paragraph(doc, "盈利能力分析", bold=True, font_size=12, space_after=6, heading_font=True)
add_cn_paragraph(doc,
    "Non-GAAP毛利率75.0%，与FY2026第四季度（75.2%）基本持平，较受H20影响的FY2026第一季度"
    "（61.0%）大幅回升。在Blackwell有史以来最快的产品爬坡期间保持毛利率稳定，展示了强劲的定价"
    "能力和制造良率改善。GAAP毛利率为74.9%。营业费用总额76亿美元（GAAP），其中研发支出63亿美元"
    "——反映出在下一代GPU架构（Rubin/Vera）和软件平台方面的持续大力投入。",
    font_size=10.5, space_after=6)

m_table = doc.add_table(rows=5, cols=6)
m_table.style = "Light Grid Accent 1"
m_h = ["指标", "Q1 FY26", "Q3 FY26", "Q4 FY26", "Q1 FY27", "同比变化"]
for i, h in enumerate(m_h):
    m_table.rows[0].cells[i].text = h
m_data = [
    ["GAAP毛利率", "60.5%", "73.4%", "75.0%", "74.9%", "+1,440 bps"],
    ["Non-GAAP毛利率", "61.0%", "73.6%", "75.2%", "75.0%", "+1,400 bps"],
    ["GAAP营业利润率", "49.0%", "63.2%", "65.0%", "65.6%", "+1,660 bps"],
    ["Non-GAAP营业利润率", "52.8%", "64.9%", "67.7%", "65.9%", "+1,310 bps"],
]
for r, row_data in enumerate(m_data, 1):
    for c, val in enumerate(row_data):
        m_table.rows[r].cells[c].text = val
format_table_cn(m_table)
add_source_cn(doc, "资料来源：公司财报。Q1 FY26毛利率受45亿美元H20库存减值影响。")

add_cn_paragraph(doc, "图2 — 季度每股收益走势", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart2_quarterly_eps.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报；Bloomberg市场共识（2026年5月19日）。")

add_cn_paragraph(doc, "图3 — 季度利润率走势", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart3_margin_trends.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报。")

# ════════════════════════════════════════════════════════════════════════════
# 第4-5页：关键指标与指引
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_cn_paragraph(doc, "关键指标与指引", bold=True, font_size=14,
                 color=(0x1B, 0x5E, 0x20), space_after=8, heading_font=True)

add_cn_paragraph(doc, "关键运营指标", bold=True, font_size=12, space_after=6, heading_font=True)

kpi_table = doc.add_table(rows=7, cols=4)
kpi_table.style = "Light Grid Accent 1"
kpi_h = ["指标", "Q1 FY27", "Q4 FY26", "变化"]
for i, h in enumerate(kpi_h):
    kpi_table.rows[0].cells[i].text = h
kpi_data = [
    ["数据中心计算收入 ($B)", "$60.4", "$51.2", "+18% 环比"],
    ["数据中心网络收入 ($B)", "$14.8", "$11.0", "+35% 环比"],
    ["自由现金流 ($B)", "$49.0", "$34.9", "+40% 环比"],
    ["经营性现金流 ($B)", "~$51.0", "$36.2", "+41% 环比"],
    ["股票回购 ($B)", "~$20.0", "~$11.0", "—"],
    ["摊薄股份数 (B)", "~24.4", "24.4", "持平"],
]
for r, row_data in enumerate(kpi_data, 1):
    for c, val in enumerate(row_data):
        kpi_table.rows[r].cells[c].text = val
format_table_cn(kpi_table)
add_source_cn(doc, "资料来源：公司业绩发布（2026年5月20日）；10-Q（2026年5月28日提交）。")

add_cn_paragraph(doc, "图4 — 收入结构：数据中心 vs. 其他板块", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart4_segment_revenue.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报。")

add_cn_paragraph(doc, "图5 — 数据中心收入及同比增长", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart5_dc_revenue_growth.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报。")

# Guidance
doc.add_page_break()
add_cn_paragraph(doc, "FY2027第二季度指引分析", bold=True, font_size=12, space_after=6, heading_font=True)
add_cn_paragraph(doc,
    "管理层指引FY2027第二季度收入约910亿美元（±2%），隐含环比增长约11.5%，显著高于财报前市场预期"
    "的约850亿美元。该指引不包含来自中国的数据中心计算收入（受美国出口管制影响）——这一限制使得"
    "有机增长表现更加亮眼。",
    font_size=10.5, space_after=6)

g_table = doc.add_table(rows=5, cols=4)
g_table.style = "Light Grid Accent 1"
g_h = ["指标", "Q2 FY27E 指引", "Q1 FY27 实际", "隐含环比"]
for i, h in enumerate(g_h):
    g_table.rows[0].cells[i].text = h
g_data = [
    ["收入", "~$910亿 (±2%)", "$816亿", "+11.5%"],
    ["GAAP毛利率", "74.9% (±50 bps)", "74.9%", "持平"],
    ["Non-GAAP毛利率", "75.0% (±50 bps)", "75.0%", "持平"],
    ["Non-GAAP营业费用", "~$83亿", "~$75亿", "+10.7%"],
]
for r, row_data in enumerate(g_data, 1):
    for c, val in enumerate(row_data):
        g_table.rows[r].cells[c].text = val
format_table_cn(g_table)
add_source_cn(doc, "资料来源：公司Q1 FY2027业绩发布及CFO评论（2026年5月20日）。")

add_cn_paragraph(doc,
    "该指引值得关注的原因包括：（1）在收入基数已经非常庞大的情况下仍暗示持续环比加速；（2）毛利率"
    "指引约75%，表明Blackwell生产爬坡未造成毛利率压力；（3）营业费用升至约83亿美元，反映了在"
    "Rubin/Vera下一代架构和CUDA软件生态扩展方面的持续重投入；（4）明确排除中国计算收入，提供了"
    "清晰的有机增长视角。",
    font_size=10.5, space_after=6)

add_cn_paragraph(doc, "图6 — Q1 FY2027收入超预期分解", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart6_beat_waterfall.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报；Bloomberg市场共识（2026年5月19日）。")

# ════════════════════════════════════════════════════════════════════════════
# 第6-7页：更新投资论点
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_cn_paragraph(doc, "更新投资论点", bold=True, font_size=14,
                 color=(0x1B, 0x5E, 0x20), space_after=8, heading_font=True)
add_cn_paragraph(doc, "论点影响评估", bold=True, font_size=12, space_after=6, heading_font=True)

add_cn_bullet(doc,
    "论点一：英伟达是AI基础设施建设的最大受益者 — 增强",
    "Q1 FY2027业绩有力地验证了这一论点。数据中心收入752亿美元（同比+92%）表明AI基础设施投资"
    "周期仍在加速而非见顶。客户基础持续扩大——从超大规模客户（380亿美元）到ACIE（370亿美元）——"
    "表明AI计算需求并非集中在少数买家，而是正扩展至主权AI项目、企业部署和专业AI云提供商。"
    "Blackwell创纪录的爬坡速度进一步巩固了英伟达的技术领先地位。")

add_cn_bullet(doc,
    "论点二：网络业务成为下一个重要增长引擎 — 增强",
    "网络收入148亿美元（同比+199%，环比+35%）验证了我们的论点：随着GPU集群扩展至数万片加速器，"
    "高带宽互联变得至关重要。NVLink、InfiniBand和Spectrum-X以太网解决方案越来越多地作为"
    "Blackwell集成系统的一部分销售。网络收入现占数据中心收入的约20%，高于一年前的约13%，"
    "展示了GPU之外的显著可触达市场扩展。")

add_cn_bullet(doc,
    "论点三：毛利率在产品换代期间维持73%以上 — 不变",
    "Non-GAAP毛利率75.0%远高于我们73%下限论点。在Hopper到Blackwell换代期间保持稳定尤为突出，"
    "因为新产品爬坡历史上往往造成毛利率波动。管理层第二季度指引75.0%（±50 bps）提供了持续"
    "毛利率的可见性。但我们维持FY2027全年74.5%的保守假设，以应对推理GPU规模化带来的潜在"
    "产品组合变化。")

add_cn_paragraph(doc, "风险评估更新", bold=True, font_size=12, space_after=6,
                 space_before=10, heading_font=True)
add_cn_paragraph(doc,
    "核心风险基本不变但需持续关注：（1）出口管制风险——指引中排除中国数据中心计算收入凸显了持续的"
    "监管逆风，但英伟达已证明有能力在无中国市场的情况下实现强劲增长；（2）客户集中度——尽管ACIE"
    "增长令人鼓舞，超大规模客户仍贡献约50%的数据中心收入，如果云资本开支放缓存在周期性风险；"
    "（3）竞争——AMD的MI400和来自谷歌（TPU）、亚马逊（Trainium）、微软（Maia）的定制ASIC"
    "构成长期竞争威胁，但短期内份额流失似乎有限；（4）估值——约25倍FY2027E市盈率已反映了显著的"
    "持续增长预期，留给执行失误的空间有限。",
    font_size=10.5, space_after=6)

add_cn_paragraph(doc, "图7 — 季度自由现金流", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart7_free_cash_flow.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报。")

add_cn_paragraph(doc, "图8 — Q1 FY2027数据中心收入结构", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart8_dc_mix.png"), width=Inches(4.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司业绩发布（2026年5月20日）。")

# ════════════════════════════════════════════════════════════════════════════
# 第8-10页：估值与预测
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_cn_paragraph(doc, "估值与更新预测", bold=True, font_size=14,
                 color=(0x1B, 0x5E, 0x20), space_after=8, heading_font=True)

add_cn_paragraph(doc, "估值框架", bold=True, font_size=12, space_after=6, heading_font=True)
add_cn_paragraph(doc,
    f"目标价$280（此前$260）基于混合估值方法：\n\n"
    f"DCF分析（40%权重）：基于FY2027E-FY2032E现金流的更新DCF得出公允价值$295。核心假设：5年"
    f"收入复合增长率22%、终端EBIT利润率62%、WACC 10.5%、终端增长率3.5%。\n\n"
    f"市盈率（40%权重）：FY2028E Non-GAAP每股收益$10.50乘以32倍得$336。考虑执行风险给予20%"
    f"折扣，得$269。\n\n"
    f"EV/EBITDA（20%权重）：FY2028E EBITDA 3,100亿美元乘以25倍，扣除净现金后得每股$250。\n\n"
    f"混合目标价：(40% x $295) + (40% x $269) + (20% x $250) = $276，取整为$280。",
    font_size=10.5, space_after=6)

add_cn_paragraph(doc, "详细预测更新", bold=True, font_size=12, space_after=6,
                 space_before=8, heading_font=True)

det_table = doc.add_table(rows=11, cols=5)
det_table.style = "Light Grid Accent 1"
det_h = ["指标", "FY27E（旧）", "FY27E（新）", "变动", "FY28E（新）"]
for i, h in enumerate(det_h):
    det_table.rows[0].cells[i].text = h
det_data = [
    ["收入 ($B)", "$340.0", "$365.0", "+7.4%", "$440.0"],
    ["  数据中心 ($B)", "$310.0", "$335.0", "+8.1%", "$400.0"],
    ["  边缘计算 ($B)", "$30.0", "$30.0", "—", "$40.0"],
    ["毛利润 ($B)", "$253.3", "$273.8", "+8.1%", "$327.8"],
    ["Non-GAAP毛利率", "74.5%", "75.0%", "+50 bps", "74.5%"],
    ["Non-GAAP营业费用 ($B)", "$35.0", "$35.0", "—", "$42.0"],
    ["Non-GAAP营业利润 ($B)", "$225.0", "$245.0", "+8.9%", "$290.0"],
    ["Non-GAAP营业利润率", "66.2%", "67.1%", "+90 bps", "65.9%"],
    ["Non-GAAP每股收益", "$8.20", "$8.80", "+7.3%", "$10.50"],
    ["自由现金流 ($B)", "$195.0", "$215.0", "+10.3%", "$250.0"],
]
for r, row_data in enumerate(det_data, 1):
    for c, val in enumerate(row_data):
        det_table.rows[r].cells[c].text = val
format_table_cn(det_table)
add_source_cn(doc, "资料来源：分析师预测。")

add_cn_paragraph(doc, "图9 — FY2027E预测修订", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart9_estimate_revisions.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：分析师预测。")

add_cn_paragraph(doc, "图10 — 季度收入环比增长率", bold=True, font_size=10,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading_font=True)
doc.add_picture(os.path.join(OUT, "nvda_chart10_sequential_growth.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_source_cn(doc, "资料来源：公司财报。")

# ════════════════════════════════════════════════════════════════════════════
# 资料来源与参考文献
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_cn_paragraph(doc, "资料来源与参考文献", bold=True, font_size=14,
                 color=(0x1B, 0x5E, 0x20), space_after=8, heading_font=True)

add_cn_paragraph(doc, "业绩资料（Q1 FY2027）：", bold=True, font_size=10.5, space_after=4, heading_font=True)

src_items = [
    ("业绩新闻稿（2026年5月20日）",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-first-quarter-fiscal-2027"),
    ("Form 8-K — 新闻稿（2026年5月20日提交）",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27pr.htm"),
    ("CFO评论（2026年5月20日）",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000051/q1fy27cfocommentary.htm"),
    ("Form 10-Q（2026年5月28日提交）",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000052/nvda-20260426.htm"),
    ("Q1 FY2027业绩电话会议记录（2026年5月20日）",
     "https://seekingalpha.com/article/4907259-nvidia-corporation-nvda-q1-2027-earnings-call-transcript"),
]
for title, url in src_items:
    p = doc.add_paragraph()
    run = p.add_run("• ")
    set_cn_font(run, cn_font=SONG, size=10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

add_cn_paragraph(doc, "", space_after=4)
add_cn_paragraph(doc, "往期参考资料：", bold=True, font_size=10.5, space_after=4, heading_font=True)

prior_items = [
    ("FY2026第四季度业绩发布（2026年2月26日）",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-fourth-quarter-and-fiscal-2026"),
    ("FY2026第四季度 Form 8-K",
     "https://www.sec.gov/Archives/edgar/data/0001045810/000104581026000019/q4fy26pr.htm"),
    ("FY2026第三季度业绩发布（2025年11月）",
     "https://nvidianews.nvidia.com/news/nvidia-announces-financial-results-for-third-quarter-fiscal-2026"),
]
for title, url in prior_items:
    p = doc.add_paragraph()
    run = p.add_run("• ")
    set_cn_font(run, cn_font=SONG, size=10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)

add_cn_paragraph(doc, "", space_after=4)
add_cn_paragraph(doc, "共识与市场数据：", bold=True, font_size=10.5, space_after=4, heading_font=True)
add_cn_paragraph(doc,
    "• Bloomberg市场共识预测（截至2026年5月19日收盘）\n"
    "• Yahoo Finance / yfinance 实时市场数据\n"
    "• 英伟达投资者关系：https://investor.nvidia.com",
    font_size=10, space_after=6)

# Disclaimer
add_cn_paragraph(doc, "", space_after=6)
add_cn_paragraph(doc, "免责声明", bold=True, font_size=9, color=(0x75, 0x75, 0x75),
                 space_after=2, heading_font=True)
add_cn_paragraph(doc,
    "本报告仅供信息参考，不构成投资建议。分析师不是持牌投资顾问。所有预测和目标价均基于公开信息，"
    "代表分析师的独立评估。过往表现不代表未来结果。投资者在做出投资决策前应进行独立的尽职调查。",
    font_size=8, color=(0x99, 0x99, 0x99), space_after=4)

# ── Save ────────────────────────────────────────────────────────────────────
output_path = os.path.join(OUT, "NVDA_Q1_FY2027_业绩更新报告_中文版.docx")
doc.save(output_path)
print(f"Report saved: {output_path}")
