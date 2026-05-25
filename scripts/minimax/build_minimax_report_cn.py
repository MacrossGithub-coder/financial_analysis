#!/usr/bin/env python3
"""
MiniMax Group (0100.HK) FY2025 年度业绩 — 中文版 DOCX 业绩更新报告
机构研报格式：约10页，正文宋体，标题黑体
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, yfinance as yf

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/MINIMAX/"
CHARTS = OUT

# ── 市场数据 ─────────────────────────────────────────────────────────────────
def get_market_data(ticker):
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"警告：无法获取 {ticker} 市场数据：{e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("0100.HK")
price = mkt["price"]
mcap = mkt["market_cap"]
mcap_b = round(mcap / 1e9, 1) if isinstance(mcap, (int, float)) else "N/A"
mcap_b_usd = round(mcap_b / 7.8, 1) if isinstance(mcap_b, (int, float)) else "N/A"

# ── 颜色 ─────────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)

# ── 工具函数 ─────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink'); rPr.append(rStyle)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0066CC'); rPr.append(color_el)
    u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single'); rPr.append(u_el)
    sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), '20'); rPr.append(sz_el)
    new_run.append(rPr)
    t_el = OxmlElement('w:t'); t_el.text = text; new_run.append(t_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def heading(doc, text, level=1, color=NAVY, size=None, bold=True, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.bold = bold; run.font.color.rgb = color
    if size is None: size = {1: 16, 2: 13, 3: 11}.get(level, 11)
    run.font.size = Pt(size)
    return p

def body(doc, text, size=10, bold=False, color=DGRAY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.color.rgb = color
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(size); run.font.color.rgb = DGRAY
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = DGRAY
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            p.paragraph_format.space_after = Pt(4)
    else:
        body(doc, f"[图表未找到: {fname}]", italic=True, color=RED)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, '003366')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.size = Pt(9); run.font.bold = True; run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            if i % 2 == 0: set_cell_bg(cell, 'F2F6FA')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                    run.font.size = Pt(9); run.font.color.rgb = DGRAY
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# 构建文档
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
for section in doc.sections:
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── 第1页：业绩摘要 ─────────────────────────────────────────────────────────
heading(doc, 'MiniMax Group Inc.（稀宇科技，0100.HK）', level=1, size=18)
heading(doc, 'FY2025年度业绩更新报告', level=2, size=14, color=BLUE)

price_str = f"HK${price}" if price != "N/A" else "N/A"
body(doc, f'评级：买入  |  目标价：HK$1,100  |  现价：{price_str}  |  市值：约US${mcap_b_usd}B',
     bold=True, size=10, color=NAVY)
body(doc, '行业：科技——AI基础模型  |  交易所：香港交易所  |  报告日期：2026年3月2日',
     size=9, italic=True)

heading(doc, '核心要点', level=2)
bullet(doc, '收入同比大增158.9%至7,900万美元，超出市场一致预期约11%，海外收入占比提升至73%，全球化战略成效显著。')
bullet(doc, '毛利率从12.2%大幅提升至25.4%（提升13.2个百分点），反映模型推理效率优化和云基础设施成本下降。')
bullet(doc, 'Q4 2025收入约2,600万美元（同比+131%），超高盛预期19%；调整后净亏损6,500万美元，较预期的1.3亿美元大幅收窄50%。')
bullet(doc, '销售费用逆势下降40.3%至5,190万美元，在收入近翻三倍的背景下实现，体现产品驱动增长的强劲势头。')
bullet(doc, '截至2026年2月，年化经常性收入（ARR）突破1.5亿美元，暗示2026年收入有望大幅超出市场预期。')
bullet(doc, '现金储备增至10.5亿美元（+19.3%），按当前调整后亏损率计算，可支撑约4.2年运营。')

heading(doc, '业绩概览', level=2)
add_table(doc,
    ['指标', 'FY2025实际', 'FY2024实际', '同比变动', '对比一致预期'],
    [
        ['收入', '7,900万美元', '3,050万美元', '+158.9%', '超出约11%'],
        ['毛利', '2,010万美元', '374万美元', '+437.2%', '超出'],
        ['毛利率', '25.4%', '12.2%', '+13.2pp', '超出'],
        ['调整后净亏损', '(2.509亿美元)', '(2.442亿美元)', '+2.7%', '基本持平'],
        ['研发支出', '2.528亿美元', '1.890亿美元', '+33.8%', '——'],
        ['现金储备', '10.503亿美元', '8.806亿美元', '+19.3%', '——'],
    ],
    col_widths=[3.5, 3.0, 3.0, 2.5, 2.5]
)
body(doc, '数据来源：MiniMax FY2025年度业绩公告（2026年3月2日）；高盛一致预期。',
     size=8, italic=True)

doc.add_page_break()

# ── 第2-3页：详细业绩分析 ───────────────────────────────────────────────────
heading(doc, '详细业绩分析', level=1, size=14)

heading(doc, '收入表现', level=2)
body(doc, 'MiniMax FY2025全年实现收入7,900万美元，同比增长158.9%，较FY2024的3,050万美元大幅增长。这是公司自2026年1月9日在港交所上市以来发布的首份年度报告。业绩超出中金公司预期约11%，验证了AI基础模型平台的强劲商业化势头。')

body(doc, '两大业务板块均实现高速增长：')
bullet(doc, 'AI原生产品（Talkie/星野、Hailuo AI）：收入增长143.4%至5,310万美元（占收入67.2%），受益于累计用户突破2.36亿、覆盖200+国家和地区的快速增长。截至2025年9月，月活用户约2,760万。')
bullet(doc, '开放平台及企业服务：收入飙升197.8%至2,600万美元（占收入32.8%），企业客户和开发者超过21.4万家，遍布100+国家。MiniMax-M2成为OpenRouter平台上首个日处理Token超500亿的中国模型。')

add_chart(doc, 'minimax_chart1_revenue.png', 5.5, '图1：年度收入增长趋势。数据来源：公司公告。')
add_chart(doc, 'minimax_chart2_segment.png', 5.5, '图2：收入结构——按业务板块。数据来源：公司公告。')

heading(doc, '地域多元化', level=2)
body(doc, '海外市场贡献了FY2025收入的73.0%（较FY2024的69.8%进一步提升），这对于一家中国AI公司而言是极为突出的成就。全球化布局大幅降低了中国特有的监管和地缘政治风险，也体现了MiniMax旗下Talkie（AI社交陪伴）和Hailuo AI（视频生成）等消费级AI产品的全球吸引力。公司产品已覆盖超过200个国家和地区。')

add_chart(doc, 'minimax_chart6_geo.png', 5.5, '图3：收入地域分布。数据来源：公司公告。')

doc.add_page_break()

# ── 第4-5页：关键指标与利润率分析 ───────────────────────────────────────────
heading(doc, '关键指标与利润率分析', level=1, size=14)

heading(doc, '毛利率扩张', level=2)
body(doc, '毛利率大幅提升13.2个百分点至25.4%，为MiniMax历史上最重要的盈利能力里程碑。毛利增长437.2%至2,010万美元。改善主要受益于：（1）自研Lightning Attention架构优化模型推理成本；（2）云基础设施利用率提升；（3）收入规模增长摊薄固定服务器成本。毛利率从FY2023的负值，到FY2024的12.2%，再到FY2025的25.4%，展现了清晰的长期盈利路径。')

add_chart(doc, 'minimax_chart3_margin.png', 5.5, '图4：毛利率趋势。数据来源：公司公告。')

heading(doc, '运营费用控制', level=2)
body(doc, '总运营费用相对收入增长呈现良好的成本纪律：')
bullet(doc, '研发支出：增长33.8%至2.528亿美元，研发强度（研发/收入）从619%降至320%，收入增速远超研发增速。428人团队持续推进文本、视频、语音、音乐的多模态AI能力。')
bullet(doc, '销售及分销费用：大幅下降40.3%至5,190万美元（FY2024为8,700万美元），在收入增长159%的背景下实现逆势缩减，表明产品驱动的有机增长模式已经跑通，对付费获客的依赖大幅降低。')
bullet(doc, '行政费用：增长155.9%至3,680万美元，主要由IPO相关专业费用和上市公司合规成本驱动——属一次性阶梯式增加，预计后续将回归正常。')

add_chart(doc, 'minimax_chart4_opex.png', 5.5, '图5：运营费用结构。数据来源：公司公告。')

heading(doc, '调整后净亏损与现金消耗', level=2)
body(doc, '调整后净亏损为2.509亿美元，与FY2024的2.442亿美元基本持平。GAAP净亏损扩大至18.716亿美元，主要因可转换优先股公允价值重估产生15.899亿美元非现金损失（该等优先股已于2026年1月上市时转换为普通股）。值得注意的是，Q4 2025调整后净亏损约6,500万美元，较高盛预期的1.3亿美元大幅收窄50%，显示下半年费用效率显著改善。')

add_chart(doc, 'minimax_chart5_adj_loss.png', 5.5, '图6：调整后净亏损趋势。数据来源：公司公告。')

doc.add_page_break()

# ── 第6-7页：投资论点更新 ───────────────────────────────────────────────────
heading(doc, '投资论点更新', level=1, size=14)

heading(doc, '变化要点——论点影响：正面', level=2, color=GREEN)
body(doc, 'MiniMax首份年度报告从多个维度强化了投资论点：')

bullet(doc, '收入动能加速：Q4收入约2,600万美元（同比+131%，超预期19%），加上截至2026年2月ARR突破1.5亿美元，暗示FY2026收入有望大幅超出高盛2.5亿美元的预期。收入曲线呈加速上行态势。')
bullet(doc, '利润率改善超预期：25.4%的毛利率和销售费用40%的压缩证明MiniMax的单位经济模型改善速度远超同行。高盛预计2029年实现盈亏平衡，但该预测可能偏保守。')
bullet(doc, '全球AI平台战略：73%的海外收入占比使MiniMax堪称全球化程度最高的中国AI基础模型公司，在与OpenAI、Anthropic、Google的全球竞争格局中占据独特定位。')
bullet(doc, '模型竞争力：M2.5（2026年2月发布）编程基准测试成绩达80.2%，逼近Claude Opus 4.6的80.8%，验证了MiniMax在前沿AI技术层面的实力。')
bullet(doc, '充裕的现金储备：10.503亿美元的总流动性按当前调整后亏损率可支撑4年以上运营，有效消除近期稀释风险。')

heading(doc, '主要风险', level=2, color=RED)
bullet(doc, '盈利时间表：调整后净亏损仍高达2.509亿美元。高盛预计2029年实现盈亏平衡——若时间推迟将对股价形成压力。')
bullet(doc, '竞争激烈：AI基础模型领域竞争异常激烈。OpenAI、Google、Anthropic在资源和企业渠道方面拥有显著优势。')
bullet(doc, '监管风险：作为全球运营的中国AI公司，MiniMax面临地缘政治逆风、出口限制和数据主权法规的潜在影响。')
bullet(doc, '客户集中度：消费产品（Talkie、Hailuo）主导收入，企业端收入仅2,600万美元，尚处早期。')
bullet(doc, '资产负债率：FY2025末攀升至343.3%（FY2024末为187.8%），主因优先股负债（已于上市时转换），但仍需持续关注。')

heading(doc, '催化剂', level=2)
bullet(doc, 'MiniMax-M3模型发布（预计2026年下半年）——模型能力的阶梯式跃升')
bullet(doc, 'Hailuo-03新一代视频模型：端到端原生多模态架构')
bullet(doc, '企业平台收入拐点——管理层定位"AI时代的平台型公司"')
bullet(doc, '潜在纳入港股通——释放内地投资者资金流入')
bullet(doc, 'Q1 2026业绩发布（预计2026年6月）——首份上市后季度报告，ARR验证')

doc.add_page_break()

# ── 第8-9页：估值与预测 ─────────────────────────────────────────────────────
heading(doc, '估值与预测', level=1, size=14)

heading(doc, '收入预测', level=2)
add_table(doc,
    ['', 'FY2024实际', 'FY2025实际', 'FY2026预测', 'FY2027预测'],
    [
        ['收入（百万美元）', '30.5', '79.0', '250', '980'],
        ['同比增速', '+323%', '+158.9%', '+216%', '+292%'],
        ['毛利率', '12.2%', '25.4%', '约32%', '约40%'],
        ['调整后净亏损（百万美元）', '(244.2)', '(250.9)', '约(200)', '约(100)'],
        ['ARR（百万美元）', 'N/A', 'N/A', '>150（2026.2月）', '待定'],
    ],
    col_widths=[4.0, 2.5, 2.5, 2.5, 2.5]
)
body(doc, '数据来源：FY2024-2025实际数据来自公司公告；FY2026-2027预测基于高盛研报（2026年3月）。',
     size=8, italic=True)

heading(doc, '估值框架', level=2)
body(doc, '鉴于MiniMax尚未实现盈利，我们主要采用基于收入的估值方法：')
bullet(doc, f'当前市值：按{price_str}计算约US${mcap_b_usd}B')
bullet(doc, 'EV/收入（FY2026E）：基于2.5亿美元收入预测，隐含的估值倍数与高增长AI同行一致。')
bullet(doc, '高盛预计MiniMax全球市场份额将从2026年的0.3%扩大至2030年的2.5%，2029年实现盈亏平衡。')
bullet(doc, '公司股价自HK$165的IPO价格上涨约366%至约HK$768，反映投资者对AI基础设施标的的强烈需求。')

heading(doc, '分析师一致预期', level=2)
add_table(doc,
    ['分析师', '评级', '目标价（港元）', '隐含上行空间'],
    [
        ['高盛', '中性', '1,018', '+32%'],
        ['中金公司', '跑赢行业', '1,109', '+44%'],
        ['国泰海通国际', '增持', '1,012', '+32%'],
        ['东吴证券', '买入', '1,100', '+43%'],
        ['一致预期（12位分析师）', '强烈买入', '1,114', '+45%'],
    ],
    col_widths=[4.0, 2.5, 3.5, 3.0]
)
body(doc, '数据来源：Bloomberg一致预期，截至2026年5月。12位买入，0位卖出。',
     size=8, italic=True)

add_chart(doc, 'minimax_chart10_targets.png', 5.5, '图7：分析师目标价对比当前股价。数据来源：Bloomberg。')

heading(doc, '研发投入背景', level=2)
body(doc, 'MiniMax FY2025研发支出2.528亿美元（占收入320%），在中国AI行业中研发强度位居前列。虽然短期拖累盈利，但研发投入驱动的模型创新管线是商业化增长的核心引擎。研发强度从FY2024的619%大幅下降，证明收入增长远快于研发支出增速（33.8%）。')

add_chart(doc, 'minimax_chart8_rd_vs_rev.png', 5.5, '图8：研发投入与收入对比。数据来源：公司公告。')

doc.add_page_break()

# ── 第10页：用户数据与运营亮点 ───────────────────────────────────────────────
heading(doc, '用户数据与运营亮点', level=1, size=14)

heading(doc, '用户增长', level=2)
body(doc, 'MiniMax累计服务2.36亿+个人用户（覆盖200+国家），21.4万+企业客户和开发者（来自100+国家）。核心产品里程碑：')
bullet(doc, 'Talkie/星野（AI社交陪伴）：截至2025年9月月活约2,000万，为全球最受欢迎的AI陪伴应用之一。')
bullet(doc, 'Hailuo AI（视频生成）：月活约560万，累计生成超6亿个视频。')
bullet(doc, '语音模型：累计生成超2亿小时语音内容。')
bullet(doc, 'M2模型登录OpenRouter：成为首个在该平台日处理Token超500亿的中国模型。')

add_chart(doc, 'minimax_chart7_users.png', 5.5, '图9：累计用户增长。数据来源：公司公告、招股书。')

heading(doc, '现金储备与资产负债表', level=2)
body(doc, '截至2025年12月31日，总现金储备（含受限现金、定期存款及短期金融资产）为10.503亿美元，较2024年末的8.806亿美元增长19.3%。增长主要来自IPO募集资金和持续融资。按调整后年度消耗率约2.51亿美元计算，MiniMax拥有约4.2年运营跑道——远超机构投资者通常要求的2年门槛。')

add_chart(doc, 'minimax_chart9_cash.png', 5.5, '图10：现金储备变化。数据来源：公司公告。')

heading(doc, '员工指标', level=2)
body(doc, 'MiniMax以428名全职员工的精干团队运营，FY2025人均收入约18.5万美元，人均研发支出约59.1万美元。相对于收入和模型能力而言极小的团队规模，充分体现了运营效率。')

doc.add_page_break()

# ── 数据来源 ─────────────────────────────────────────────────────────────────
heading(doc, '数据来源与参考文献', level=1, size=14)

body(doc, '业绩材料（FY2025）：', bold=True, size=10)
p1 = doc.add_paragraph()
p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
run1 = p1.add_run('• MiniMax FY2025年度业绩公告（2026年3月2日）：')
run1.font.name = 'Times New Roman'; run1._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run1.font.size = Pt(9); run1.font.color.rgb = DGRAY
add_hyperlink(p1, 'MiniMax官方公告',
              'https://www.minimax.io/news/minimax-global-announces-full-year-2025-financial-results')

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('• PR Newswire新闻稿（2026年3月2日）：')
run2.font.name = 'Times New Roman'; run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run2.font.size = Pt(9); run2.font.color.rgb = DGRAY
add_hyperlink(p2, 'PR Newswire',
              'https://www.prnewswire.com/apac/news-releases/minimax-announces-full-year-2025-financial-results-302700878.html')

body(doc, '分析师研报：', bold=True, size=10)
p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run('• 高盛——MiniMax（00100.HK）Q4业绩超预期（2026年3月3日）：')
run3.font.name = 'Times New Roman'; run3._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run3.font.size = Pt(9); run3.font.color.rgb = DGRAY
add_hyperlink(p3, '高盛研报',
              'https://finance.sina.com.cn/stock/hkstock/hkgg/2026-03-03/doc-inhpsxzn7541443.shtml')

p4 = doc.add_paragraph()
p4.paragraph_format.space_before = Pt(2); p4.paragraph_format.space_after = Pt(2)
run4 = p4.add_run('• 中金公司——跑赢行业，目标价HK$1,109（2026年3月3日）：')
run4.font.name = 'Times New Roman'; run4._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run4.font.size = Pt(9); run4.font.color.rgb = DGRAY
add_hyperlink(p4, '中金研报',
              'https://finance.sina.com.cn/stock/hkstock/hkgg/2026-03-03/doc-inhpteik7480227.shtml')

body(doc, '其他来源：', bold=True, size=10)
p5 = doc.add_paragraph()
p5.paragraph_format.space_before = Pt(2); p5.paragraph_format.space_after = Pt(2)
run5 = p5.add_run('• Bloomberg一致预期（截至2026年5月）')
run5.font.name = 'Times New Roman'; run5._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run5.font.size = Pt(9); run5.font.color.rgb = DGRAY

p6 = doc.add_paragraph()
p6.paragraph_format.space_before = Pt(2); p6.paragraph_format.space_after = Pt(2)
run6 = p6.add_run('• 市场数据：Yahoo Finance（0100.HK），通过yfinance API获取')
run6.font.name = 'Times New Roman'; run6._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run6.font.size = Pt(9); run6.font.color.rgb = DGRAY

# ── 免责声明 ─────────────────────────────────────────────────────────────────
doc.add_paragraph()
body(doc, '免责声明：本报告仅供参考，不构成投资建议。过往表现不代表未来收益。作者可能持有本报告讨论的证券头寸。'
         '所有数据来源于公开的公司公告和第三方研究报告。',
     size=8, italic=True, color=DGRAY)

# ── 保存 ─────────────────────────────────────────────────────────────────────
fname = OUT + "MINIMAX_FY2025_业绩更新报告_中文版.docx"
doc.save(fname)
print(f"\n✅ 中文报告已保存：{fname}")
