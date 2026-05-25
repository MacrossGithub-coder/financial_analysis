#!/usr/bin/env python3
"""
MiniMax Group (0100.HK) FY2025 Annual Results — English DOCX Earnings Update
Institutional equity research format, 8-12 pages
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, yfinance as yf

OUT = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/MINIMAX/"
CHARTS = OUT

# ── Market Data ──────────────────────────────────────────────────────────────
def get_market_data(ticker):
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Cannot fetch {ticker} market data: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("0100.HK")
price = mkt["price"]
mcap = mkt["market_cap"]
mcap_b = round(mcap / 1e9, 1) if isinstance(mcap, (int, float)) else "N/A"
mcap_b_usd = round(mcap_b / 7.8, 1) if isinstance(mcap_b, (int, float)) else "N/A"

# ── Colors ───────────────────────────────────────────────────────────────────
NAVY  = RGBColor(0x00, 0x33, 0x66)
BLUE  = RGBColor(0x00, 0x66, 0xCC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DGRAY = RGBColor(0x44, 0x44, 0x44)
GREEN = RGBColor(0x1A, 0x6B, 0x3A)
RED   = RGBColor(0xCC, 0x00, 0x00)

# ── Helpers ──────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle'); rStyle.set(qn('w:val'), 'Hyperlink'); rPr.append(rStyle)
    color_el = OxmlElement('w:color'); color_el.set(qn('w:val'), '0066CC'); rPr.append(color_el)
    u_el = OxmlElement('w:u'); u_el.set(qn('w:val'), 'single'); rPr.append(u_el)
    sz_el = OxmlElement('w:sz'); sz_el.set(qn('w:val'), '20'); rPr.append(sz_el)
    new_run.append(rPr)
    t_el = OxmlElement('w:t'); t_el.text = text; new_run.append(t_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def heading(doc, text, level=1, color=NAVY, size=None, bold=True, space_before=6, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = bold; run.font.color.rgb = color
    if size is None: size = {1: 16, 2: 13, 3: 11}.get(level, 11)
    run.font.size = Pt(size)
    return p

def body(doc, text, size=10, bold=False, color=DGRAY, space_before=2, space_after=2, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before); p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.color.rgb = color
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size); run.font.color.rgb = DGRAY
    return p

def add_chart(doc, fname, width=6.0, caption=None):
    path = CHARTS + fname
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.font.size = Pt(8); run.font.italic = True; run.font.color.rgb = DGRAY
            run.font.name = 'Times New Roman'
            p.paragraph_format.space_after = Pt(4)
    else:
        body(doc, f"[Chart not found: {fname}]", italic=True, color=RED)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        set_cell_bg(cell, '003366')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.name = 'Times New Roman'; run.font.size = Pt(9)
                run.font.bold = True; run.font.color.rgb = WHITE
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            cell = table.rows[i + 1].cells[j]
            cell.text = str(val)
            if i % 2 == 0: set_cell_bg(cell, 'F2F6FA')
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'; run.font.size = Pt(9); run.font.color.rgb = DGRAY
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for j, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[j].width = Cm(w)
    return table

# ══════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(10)
for section in doc.sections:
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── PAGE 1: EARNINGS SUMMARY ────────────────────────────────────────────────
heading(doc, 'MiniMax Group Inc. (0100.HK)', level=1, size=18)
heading(doc, 'FY2025 Annual Results — Earnings Update', level=2, size=14, color=BLUE)

price_str = f"HK${price}" if price != "N/A" else "N/A"
body(doc, f'Rating: BUY  |  Price Target: HK$1,100  |  Current Price: {price_str}  |  Mkt Cap: ~US${mcap_b_usd}B',
     bold=True, size=10, color=NAVY)
body(doc, 'Sector: Technology — AI Foundation Models  |  Exchange: HKEX  |  Report Date: March 2, 2026',
     size=9, italic=True)

heading(doc, 'Key Takeaways', level=2)
bullet(doc, 'Revenue surged 158.9% YoY to US$79.0M, beating consensus by ~11%, driven by rapid global expansion with international revenue now at 73% of total.')
bullet(doc, 'Gross margin expanded dramatically to 25.4% (from 12.2%), reflecting improved model efficiency and optimized cloud infrastructure costs.')
bullet(doc, 'Q4 2025 revenue reached ~US$26M (+131% YoY), beating Goldman Sachs estimate by 19%; adjusted net loss of US$65M was 50% better than the US$130M expected.')
bullet(doc, 'Sales & distribution expenses cut 40.3% despite revenue near-tripling, demonstrating strong organic user acquisition and product-market fit.')
bullet(doc, 'ARR exceeded US$150M by February 2026, implying a significant acceleration in revenue run-rate entering 2026.')
bullet(doc, 'Cash position strengthened to US$1,050.3M (+19.3% YoY), providing ample runway for continued R&D investment.')

heading(doc, 'Results Snapshot', level=2)
add_table(doc,
    ['Metric', 'FY2025A', 'FY2024A', 'YoY Change', 'vs Consensus'],
    [
        ['Revenue', 'US$79.0M', 'US$30.5M', '+158.9%', 'Beat ~11%'],
        ['Gross Profit', 'US$20.1M', 'US$3.7M', '+437.2%', 'Beat'],
        ['Gross Margin', '25.4%', '12.2%', '+13.2pp', 'Beat'],
        ['Adj. Net Loss', '(US$250.9M)', '(US$244.2M)', '+2.7%', 'In-line'],
        ['R&D Expense', 'US$252.8M', 'US$189.0M', '+33.8%', '—'],
        ['Cash Balance', 'US$1,050.3M', 'US$880.6M', '+19.3%', '—'],
    ],
    col_widths=[4.0, 3.0, 3.0, 2.5, 2.5]
)
body(doc, 'Source: MiniMax FY2025 Annual Results Press Release, March 2, 2026; Goldman Sachs consensus estimates.',
     size=8, italic=True)

doc.add_page_break()

# ── PAGES 2-3: DETAILED RESULTS ─────────────────────────────────────────────
heading(doc, 'Detailed Results Analysis', level=1, size=14)

heading(doc, 'Revenue Performance', level=2)
body(doc, 'MiniMax delivered FY2025 revenue of US$79.0M, a 158.9% increase from US$30.5M in FY2024, marking the company\'s first annual report since listing on HKEX on January 9, 2026. The result beat CICC\'s estimate by approximately 11%, underscoring the strong monetization momentum of its AI foundation model platform.')

body(doc, 'Revenue growth was broad-based across both business segments:')
bullet(doc, 'AI-Native Products (Talkie/Xingye, Hailuo AI): Revenue grew 143.4% to US$53.1M (67.2% of total), driven by rapid user growth to 236M+ cumulative users across 200+ countries. Monthly active users reached ~27.6M by September 2025.')
bullet(doc, 'Open Platform & Enterprise Services: Revenue surged 197.8% to US$26.0M (32.8% of total), reflecting expanding adoption among 214,000+ enterprise customers and developers. The MiniMax-M2 model became the first Chinese model on OpenRouter to exceed 50 billion daily tokens.')

add_chart(doc, 'minimax_chart1_revenue.png', 5.5, 'Figure 1: Annual revenue progression. Source: Company filings.')
add_chart(doc, 'minimax_chart2_segment.png', 5.5, 'Figure 2: Revenue by business segment. Source: Company filings.')

heading(doc, 'Geographic Diversification', level=2)
body(doc, 'International markets accounted for 73.0% of FY2025 revenue (up from 69.8% in FY2024), a notable achievement for a Chinese AI company. This global diversification significantly reduces China-specific regulatory and geopolitical risks, and reflects the universal appeal of MiniMax\'s consumer-facing AI products such as Talkie (social AI companion) and Hailuo AI (video generation). The company\'s products serve users in over 200 countries and regions.')

add_chart(doc, 'minimax_chart6_geo.png', 5.5, 'Figure 3: Geographic revenue split. Source: Company filings.')

doc.add_page_break()

# ── PAGES 4-5: KEY METRICS & MARGINS ────────────────────────────────────────
heading(doc, 'Key Metrics & Margin Analysis', level=1, size=14)

heading(doc, 'Gross Margin Expansion', level=2)
body(doc, 'Gross margin improved by 13.2 percentage points to 25.4%, the most significant profitability milestone in MiniMax\'s history. Gross profit grew 437.2% to US$20.1M. The improvement was driven by: (1) optimized model inference costs through proprietary Lightning Attention architecture, (2) better cloud infrastructure utilization, and (3) increasing revenue scale diluting fixed server costs. The gross margin trajectory — from negative in FY2023, to 12.2% in FY2024, to 25.4% in FY2025 — demonstrates a clear path toward long-term profitability.')

add_chart(doc, 'minimax_chart3_margin.png', 5.5, 'Figure 4: Gross margin trend. Source: Company filings.')

heading(doc, 'Operating Expense Discipline', level=2)
body(doc, 'Total operating expenses increased modestly relative to revenue growth, reflecting improving cost discipline:')
bullet(doc, 'R&D Expenses: Grew 33.8% to US$252.8M, with R&D intensity (R&D/Revenue) declining from 619% to 320% as revenue scaled faster. The 428-person team continues to advance multi-modal AI capabilities across text, video, speech, and music.')
bullet(doc, 'Sales & Distribution: Declined 40.3% to US$51.9M (from US$87.0M), a remarkable achievement given the 159% revenue growth. This signals strong organic product-led growth and reduced dependence on paid user acquisition.')
bullet(doc, 'Administrative: Increased 155.9% to US$36.8M, primarily driven by IPO-related professional fees and public company compliance costs — a one-time step-up expected to normalize.')

add_chart(doc, 'minimax_chart4_opex.png', 5.5, 'Figure 5: Operating expense breakdown. Source: Company filings.')

heading(doc, 'Adjusted Net Loss & Cash Burn', level=2)
body(doc, 'The adjusted net loss was US$250.9M, roughly flat versus US$244.2M in FY2024. While the GAAP net loss widened to US$1,871.6M, this was primarily driven by US$1,589.9M in non-cash fair value losses on convertible preferred shares (which converted to equity upon the January 2026 IPO). Importantly, Q4 2025 adjusted net loss of ~US$65M came in 50% below Goldman Sachs\' estimate of US$130M, suggesting meaningful expense efficiency gains in the latter half of the year.')

add_chart(doc, 'minimax_chart5_adj_loss.png', 5.5, 'Figure 6: Adjusted net loss trend. Source: Company filings.')

doc.add_page_break()

# ── PAGES 6-7: UPDATED THESIS ───────────────────────────────────────────────
heading(doc, 'Updated Investment Thesis', level=1, size=14)

heading(doc, 'What Changed — Thesis Impact: Positive', level=2, color=GREEN)
body(doc, 'MiniMax\'s inaugural annual report materially strengthens our investment thesis. Several developments stand out:')

bullet(doc, 'Revenue Momentum Accelerating: Q4 revenue of ~US$26M (+131% YoY, beating by 19%) and ARR exceeding US$150M by February 2026 suggest FY2026 revenue could significantly surpass the US$250M Goldman Sachs estimate. The revenue trajectory is inflecting upward.')
bullet(doc, 'Margin Trajectory Exceeding Expectations: The 25.4% gross margin and 40% cut in sales expenses demonstrate that MiniMax\'s unit economics are improving far faster than peers. Breakeven by 2029 (Goldman Sachs estimate) may prove conservative.')
bullet(doc, 'Global AI Platform Strategy: With 73% international revenue, MiniMax is arguably the most globally diversified Chinese AI foundation model company. This positions it uniquely in the global competitive landscape alongside OpenAI, Anthropic, and Google.')
bullet(doc, 'Model Competitiveness: M2.5 (launched February 2026) achieved coding benchmarks (80.2%) comparable to Claude Opus 4.6 (80.8%), validating MiniMax\'s technical capabilities at the frontier.')
bullet(doc, 'Strong Cash Position: US$1,050.3M in total liquidity provides 4+ years of runway at current adjusted loss rates, eliminating near-term dilution risk.')

heading(doc, 'Key Risks', level=2, color=RED)
bullet(doc, 'Profitability Timeline: Adjusted net loss of US$250.9M remains substantial. Goldman Sachs projects breakeven by 2029 — any delay could pressure the stock.')
bullet(doc, 'Competitive Intensity: The AI foundation model space is intensely competitive. OpenAI, Google, and Anthropic have significantly larger resource pools and enterprise distribution.')
bullet(doc, 'Regulatory Risk: As a Chinese AI company operating globally, MiniMax faces potential geopolitical headwinds, export restrictions, and data sovereignty regulations.')
bullet(doc, 'Customer Concentration: Consumer products (Talkie, Hailuo) dominate revenue; enterprise adoption is still nascent at US$26M.')
bullet(doc, 'Debt/Asset Ratio: Climbed to 343.3% at year-end 2025 (from 187.8%), largely from preferred share liabilities that converted at IPO, but warrants monitoring.')

heading(doc, 'Catalysts', level=2)
bullet(doc, 'MiniMax-M3 model launch (expected H2 2026) — potential step-change in model capabilities')
bullet(doc, 'Hailuo-03 next-generation video model with end-to-end native multimodal architecture')
bullet(doc, 'Enterprise platform revenue inflection — management targeting "platform company for the AI era"')
bullet(doc, 'Potential inclusion in HK Stock Connect — unlocking mainland Chinese investor flows')
bullet(doc, 'Q1 2026 results (expected June 2026) — first post-IPO quarterly update, ARR validation')

doc.add_page_break()

# ── PAGES 8-9: VALUATION & ESTIMATES ─────────────────────────────────────────
heading(doc, 'Valuation & Estimates', level=1, size=14)

heading(doc, 'Revenue Estimates', level=2)
add_table(doc,
    ['', 'FY2024A', 'FY2025A', 'FY2026E', 'FY2027E'],
    [
        ['Revenue (US$M)', '30.5', '79.0', '250', '980'],
        ['YoY Growth', '+323%', '+158.9%', '+216%', '+292%'],
        ['Gross Margin', '12.2%', '25.4%', '~32%', '~40%'],
        ['Adj. Net Loss (US$M)', '(244.2)', '(250.9)', '~(200)', '~(100)'],
        ['ARR (US$M)', 'N/A', 'N/A', '>150 (Feb\'26)', 'TBD'],
    ],
    col_widths=[4.0, 2.5, 2.5, 2.5, 2.5]
)
body(doc, 'Source: FY2024-2025 actual from company filings; FY2026-2027E based on Goldman Sachs estimates (March 2026).',
     size=8, italic=True)

heading(doc, 'Valuation Framework', level=2)
body(doc, 'Given MiniMax\'s pre-profit status, we rely on revenue-based valuation methodologies:')
bullet(doc, f'Current Market Cap: ~US${mcap_b_usd}B at HK${price}/share')
bullet(doc, 'EV/Revenue (FY2026E): Based on US$250M revenue estimate, implying a premium multiple consistent with high-growth AI peers.')
bullet(doc, 'Goldman Sachs projects global market share expansion from 0.3% (2026) to 2.5% (2030), with breakeven by 2029.')
bullet(doc, 'The stock has appreciated ~366% from its HK$165 IPO price to ~HK$768, reflecting strong investor appetite for AI infrastructure plays.')

heading(doc, 'Analyst Consensus', level=2)
add_table(doc,
    ['Analyst', 'Rating', 'Target Price (HK$)', 'Implied Upside'],
    [
        ['Goldman Sachs', 'Neutral', '1,018', '+32%'],
        ['CICC', 'Outperform', '1,109', '+44%'],
        ['Guotai Junan Intl', 'Overweight', '1,012', '+32%'],
        ['Soochow Securities', 'Buy', '1,100', '+43%'],
        ['Consensus (12 analysts)', 'Strong Buy', '1,114', '+45%'],
    ],
    col_widths=[4.0, 2.5, 3.5, 3.0]
)
body(doc, 'Source: Bloomberg consensus as of May 2026. 12 Buy, 0 Sell.',
     size=8, italic=True)

add_chart(doc, 'minimax_chart10_targets.png', 5.5, 'Figure 7: Analyst target prices vs current price. Source: Bloomberg.')

heading(doc, 'R&D Investment Context', level=2)
body(doc, 'MiniMax\'s R&D investment of US$252.8M in FY2025 (320% of revenue) is among the highest in the Chinese AI sector on an R&D intensity basis. While this weighs on near-term profitability, it fuels the model innovation pipeline that drives commercial adoption. The decline from 619% R&D intensity in FY2024 demonstrates rapid revenue scaling relative to a more moderate R&D growth rate of 33.8%.')

add_chart(doc, 'minimax_chart8_rd_vs_rev.png', 5.5, 'Figure 8: R&D expense vs revenue. Source: Company filings.')

doc.add_page_break()

# ── PAGE 10: USER METRICS & OPERATING DATA ───────────────────────────────────
heading(doc, 'User Metrics & Operating Highlights', level=1, size=14)

heading(doc, 'User Growth', level=2)
body(doc, 'MiniMax has cumulatively served 236+ million individual users across 200+ countries, and 214,000+ enterprise customers and developers from 100+ countries. Key product milestones:')
bullet(doc, 'Talkie/Xingye (AI social companion): ~20M MAU as of September 2025, one of the most popular AI companion apps globally.')
bullet(doc, 'Hailuo AI (video generation): ~5.6M MAU, having generated 600+ million videos cumulatively.')
bullet(doc, 'Speech models: Generated 200+ million hours of speech cumulatively.')
bullet(doc, 'M2 model on OpenRouter: First Chinese model to exceed 50 billion daily tokens on the platform.')

add_chart(doc, 'minimax_chart7_users.png', 5.5, 'Figure 9: Cumulative user growth. Source: Company filings, prospectus.')

heading(doc, 'Cash Position & Balance Sheet', level=2)
body(doc, 'Total cash and cash equivalents (including restricted cash, time deposits, and short-term financial assets) stood at US$1,050.3M as of December 31, 2025, up from US$880.6M at year-end 2024. The increase was primarily driven by IPO proceeds and continued fundraising. With an adjusted annual burn rate of ~US$251M, MiniMax has approximately 4.2 years of runway — comfortably above the 2-year threshold institutional investors typically require.')

add_chart(doc, 'minimax_chart9_cash.png', 5.5, 'Figure 10: Cash position. Source: Company filings.')

heading(doc, 'Employee Metrics', level=2)
body(doc, 'MiniMax operates with a lean team of 428 full-time employees, translating to approximately US$185K in revenue per employee in FY2025 and US$591K in R&D spending per employee. The small team size relative to revenue and model capabilities underscores operational efficiency.')

doc.add_page_break()

# ── SOURCES ──────────────────────────────────────────────────────────────────
heading(doc, 'Sources & References', level=1, size=14)

body(doc, 'Earnings Materials (FY2025):', bold=True, size=10)
p1 = doc.add_paragraph()
p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
run1 = p1.add_run('• MiniMax FY2025 Annual Results Press Release (March 2, 2026): ')
run1.font.name = 'Times New Roman'; run1.font.size = Pt(9); run1.font.color.rgb = DGRAY
add_hyperlink(p1, 'MiniMax Official Announcement',
              'https://www.minimax.io/news/minimax-global-announces-full-year-2025-financial-results')

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2); p2.paragraph_format.space_after = Pt(2)
run2 = p2.add_run('• PR Newswire Release (March 2, 2026): ')
run2.font.name = 'Times New Roman'; run2.font.size = Pt(9); run2.font.color.rgb = DGRAY
add_hyperlink(p2, 'PR Newswire',
              'https://www.prnewswire.com/apac/news-releases/minimax-announces-full-year-2025-financial-results-302700878.html')

body(doc, 'Analyst Research:', bold=True, size=10)
p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(2); p3.paragraph_format.space_after = Pt(2)
run3 = p3.add_run('• Goldman Sachs — MiniMax (00100.HK) 4Q25 Results Beat (March 3, 2026): ')
run3.font.name = 'Times New Roman'; run3.font.size = Pt(9); run3.font.color.rgb = DGRAY
add_hyperlink(p3, 'Goldman Sachs Report',
              'https://finance.sina.com.cn/stock/hkstock/hkgg/2026-03-03/doc-inhpsxzn7541443.shtml')

p4 = doc.add_paragraph()
p4.paragraph_format.space_before = Pt(2); p4.paragraph_format.space_after = Pt(2)
run4 = p4.add_run('• CICC — Outperform, TP HK$1,109 (March 3, 2026): ')
run4.font.name = 'Times New Roman'; run4.font.size = Pt(9); run4.font.color.rgb = DGRAY
add_hyperlink(p4, 'CICC Report',
              'https://finance.sina.com.cn/stock/hkstock/hkgg/2026-03-03/doc-inhpteik7480227.shtml')

body(doc, 'Other Sources:', bold=True, size=10)
p5 = doc.add_paragraph()
p5.paragraph_format.space_before = Pt(2); p5.paragraph_format.space_after = Pt(2)
run5 = p5.add_run('• Bloomberg consensus estimates as of May 2026')
run5.font.name = 'Times New Roman'; run5.font.size = Pt(9); run5.font.color.rgb = DGRAY

p6 = doc.add_paragraph()
p6.paragraph_format.space_before = Pt(2); p6.paragraph_format.space_after = Pt(2)
run6 = p6.add_run('• Market data: Yahoo Finance (0100.HK) via yfinance API')
run6.font.name = 'Times New Roman'; run6.font.size = Pt(9); run6.font.color.rgb = DGRAY

# ── DISCLAIMER ───────────────────────────────────────────────────────────────
doc.add_paragraph()
body(doc, 'DISCLAIMER: This report is for informational purposes only and does not constitute investment advice. '
         'Past performance is not indicative of future results. The author may hold positions in the securities discussed. '
         'All data sourced from publicly available company filings and third-party research.',
     size=8, italic=True, color=DGRAY)

# ── SAVE ─────────────────────────────────────────────────────────────────────
fname = OUT + "MINIMAX_FY2025_Earnings_Update.docx"
doc.save(fname)
print(f"\n✅ Report saved: {fname}")
