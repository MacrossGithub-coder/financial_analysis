#!/usr/bin/env python3
"""Generate SpaceX (SPCX) Q1 2026 Post-IPO Earnings Update Report (Chinese / 中文版)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "SPCX")
os.makedirs(OUT, exist_ok=True)

SONG = "宋体"
HEI  = "黑体"

# ── Market Data ─────────────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price": round(info.last_price, 2), "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2), "52w_low": round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("SPCX")
price_str = f"${mkt['price']}" if mkt['price'] != 'N/A' else '~$161'
mcap_str = f"${mkt['market_cap']/1e12:.2f}T" if mkt['market_cap'] != 'N/A' else '~$2.1T'
hi52 = f"${mkt['52w_high']}" if mkt['52w_high'] != 'N/A' else '~$177'
lo52 = f"${mkt['52w_low']}" if mkt['52w_low'] != 'N/A' else '$135 (IPO)'

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cn(run, cn=SONG, en="Times New Roman", size=11):
    run.font.name = en; run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rF = rPr.find(qn("w:rFonts"))
    if rF is None:
        rF = parse_xml(f'<w:rFonts {nsdecls("w")}/>'); rPr.insert(0, rF)
    rF.set(qn("w:eastAsia"), cn)

def add_p(doc, text, bold=False, font_size=11, color=None, alignment=None,
          space_after=6, space_before=0, heading=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_cn(run, cn=HEI if heading else SONG, size=font_size)
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment: p.alignment = alignment
    pf = p.paragraph_format; pf.space_after = Pt(space_after); pf.space_before = Pt(space_before)
    return p

def add_bullet(doc, header, body, fs=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run("■ "); set_cn(r1, cn=HEI, size=fs); r1.bold = True
    r2 = p.add_run(header); set_cn(r2, cn=HEI, size=fs); r2.bold = True
    p.add_run("\n")
    r3 = p.add_run(body); set_cn(r3, cn=SONG, size=fs)
    p.paragraph_format.space_after = Pt(8)

def fmt_table(table, hdr="003366"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, hdr)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True; set_cn(r, cn=HEI, size=9)
    for i, row in enumerate(table.rows):
        if i == 0: continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs: set_cn(r, cn=SONG, size=9)

def add_src(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text); set_cn(run, cn=SONG, size=8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75); run.italic = True
    p.paragraph_format.space_after = Pt(4)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
        f'<w:sz w:val="16"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>')
    paragraph._p.append(hl)

# ── Build ───────────────────────────────────────────────────────────────────
doc = Document()
doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ═══════════════ 第1页：业绩摘要 ═══════════════════════════════════════════
add_p(doc, "太空探索技术公司 (SPCX)", bold=True, font_size=16,
      color=(0x00, 0x33, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, heading=True)
add_p(doc, "2026年第一季度 IPO后首份业绩更新报告", bold=True, font_size=14,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, heading=True)
add_p(doc, "星链贡献12亿美元利润，但xAI亏损推动净亏损43亿美元", bold=True, font_size=12,
      color=(0x00, 0x33, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10, heading=True)

for h in [
    "日期：2026年6月14日",
    "评级：首次覆盖，给予 中性 (MARKET PERFORM)",
    f"股价（截至2026年6月13日）：{price_str}",
    "目标价：$150 — 12个月",
    f"市值：{mcap_str}  |  IPO发行价：$135  |  首日收盘价：~$161 (+19%)",
]:
    add_p(doc, h, font_size=10, space_after=2)

add_p(doc, "", space_after=4)
add_p(doc, "2026年第一季度财务摘要", bold=True, font_size=12,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=6, heading=True)

st = doc.add_table(rows=7, cols=4)
st.style = "Light Grid Accent 1"
for i, h in enumerate(["指标", "Q1 2026", "Q1 2025", "同比变化"]):
    st.rows[0].cells[i].text = h
for r, d in enumerate([
    ["收入", "$46.9亿", "$40.7亿", "+15.4%"],
    ["调整后EBITDA", "$11.3亿", "N/A", "—"],
    ["营业利润/（亏损）", "$(19.4)亿", "$0.3亿", "NM"],
    ["净利润/（亏损）", "$(42.8)亿", "$(5.3)亿", "NM"],
    ["基本每股亏损", "$(1.27)", "$(0.18)", "NM"],
    ["自由现金流", "$(91)亿", "N/A", "—"],
], 1):
    for c, v in enumerate(d): st.rows[r].cells[c].text = v
fmt_table(st)
add_src(doc, "资料来源：SpaceX S-1招股说明书（2026年5月20日提交）；修正案第1号（2026年6月1日）。")

add_p(doc, "", space_after=2)

add_bullet(doc,
    "星链是皇冠明珠——Q1贡献33亿美元收入和12亿美元营业利润，用户达1,030万",
    "连接业务板块Q1收入32.6亿美元（占总收入69%），营业利润11.9亿美元，调整后EBITDA 20.9亿美元"
    "（利润率64%）。星链用户达到1,030万，覆盖164个市场，较2025年末的890万增长16%。该板块的"
    "盈利能力实质上在补贴公司其余业务。但值得关注的是，ARPU已从2023年的$99/月降至$66/月，"
    "这反映了星链向低收入市场扩张和竞争定价压力。")

add_bullet(doc,
    "xAI亏损加速：仅Q1即产生25亿美元营业亏损，AI基础设施资本开支达77亿美元",
    "AI板块（xAI、Grok、X平台）贡献约8.15亿美元收入，但产生约25亿美元营业亏损——年化超100亿美元。"
    "Q1 AI资本开支77亿美元（占总资本开支101亿美元的76%），反映了Colossus数据中心（22万片NVIDIA "
    "GPU、300MW功率）的建设。虽然与Anthropic签订的约450亿美元计算合同提供了收入可见性，但现金消耗"
    "速度惊人，该板块盈利路径仍高度不确定。")

add_bullet(doc,
    "太空板块收入同比下降28%至6.19亿美元，受星舰研发周期拖累",
    "猎鹰9号仍是主力运载火箭，但太空板块收入降至6.19亿美元，因公司仅在FY2025就投入30亿美元用于"
    "星舰开发。太空板块Q1营业亏损约6亿美元（年化）。虽然星舰最终商业化可能具有变革性意义，但短期内"
    "太空经济对整体盈利构成稀释。")

add_bullet(doc,
    "首次覆盖给予中性评级，目标价$150——卓越的业务，但估值同样非凡",
    "以约2.1万亿美元市值计算（94倍2025年收入），SpaceX估值远高于任何可比公司。星链的增长和盈利"
    "能力确实出色，但xAI亏损在加速，合计自由现金流严重为负（Q1为-91亿美元）。目标价$150隐含当前"
    "水平约7%下行空间，反映我们认为IPO后的热情已过度反映乐观情景，且近期盈利可见性受xAI投资周期限制。")

# Estimates
add_p(doc, "财务预测", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=10, heading=True)

et = doc.add_table(rows=8, cols=4)
et.style = "Light Grid Accent 1"
for i, h in enumerate(["指标", "FY2025A", "FY2026E", "FY2027E"]):
    et.rows[0].cells[i].text = h
for r, d in enumerate([
    ["收入 ($B)", "$18.7", "$22.0", "$28.0"],
    ["收入增长率", "+33%", "+18%", "+27%"],
    ["调整后EBITDA ($B)", "$6.6", "$5.5", "$9.0"],
    ["调整后EBITDA利润率", "35.2%", "25.0%", "32.1%"],
    ["营业利润/（亏损）($B)", "$(2.6)", "$(6.0)", "$(3.0)"],
    ["净利润/（亏损）($B)", "$(4.9)", "$(10.0)", "$(5.0)"],
    ["资本开支 ($B)", "$20.7", "$35.0", "$30.0"],
], 1):
    for c, v in enumerate(d): et.rows[r].cells[c].text = v
fmt_table(et)
add_src(doc, '注："A"=实际值，"E"=预测值。资料来源：S-1（实际值）、分析师预测（前瞻值）。')

# ═══════════════ 第2-3页：详细业绩分析 ═════════════════════════════════════
doc.add_page_break()
add_p(doc, "详细业绩分析", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "收入分析", bold=True, font_size=12, space_after=6, heading=True)
add_p(doc,
    "SpaceX报告2026年第一季度合并收入46.94亿美元，同比增长15.4%（Q1 2025约40.67亿美元）。"
    "营收增速放缓（对比FY2025全年33%增长）反映太空板块收入下降28%，被星链的持续强劲势头部分抵消。"
    "收入结构继续向连接业务倾斜，目前占总收入的69%，较一年前的约63%进一步上升。",
    font_size=10.5, space_after=6)

rt = doc.add_table(rows=6, cols=5)
rt.style = "Light Grid Accent 1"
for i, h in enumerate(["板块", "Q1 2026", "Q1 2025E", "同比变化", "占比"]):
    rt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["连接业务（星链）", "$32.6亿", "$25.6亿", "+27%", "69%"],
    ["太空（发射/Dragon）", "$6.19亿", "$8.6亿", "-28%", "13%"],
    ["AI（xAI / X）", "$8.15亿", "$6.5亿", "+25%", "18%"],
    ["抵消/其他", "—", "—", "—", "—"],
    ["总收入", "$46.94亿", "$40.67亿", "+15.4%", "100%"],
], 1):
    for c, v in enumerate(d): rt.rows[r].cells[c].text = v
fmt_table(rt)
add_src(doc, "资料来源：S-1招股说明书。Q1 2025板块估计值基于披露的同比增长率推算。")

add_p(doc, "图1 — SpaceX合并收入", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart1_annual_revenue.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

# Segment details
add_p(doc, "连接业务板块（星链）", bold=True, font_size=12, space_after=6, space_before=10, heading=True)
add_p(doc,
    "星链Q1收入32.6亿美元，营业利润11.9亿美元（利润率36%），调整后EBITDA 20.9亿美元（利润率64%）。"
    "截至2026年3月31日，用户基础达到1,030万，较2025年末的890万增长16%。星链在164个市场运营，"
    "控制着轨道上约75%的可机动卫星。\n\n"
    "核心关注点是ARPU下行：月度ARPU从2023年的$99降至Q1 2026的$66（约2.5年下降33%）。这反映了"
    "向新兴市场的地理扩张（当地定价较低）、来自Amazon Kuiper和OneWeb的竞争，以及企业/政府客户"
    "组合变化。虽然用户增长目前足以弥补，但随着高ARPU发达市场宽带市场趋于饱和，这一趋势值得持续关注。",
    font_size=10.5, space_after=6)

add_p(doc, "太空板块", bold=True, font_size=12, space_after=6, space_before=6, heading=True)
add_p(doc,
    "太空板块收入同比下降28%至6.19亿美元。公司在FY2025投入30亿美元用于星舰研发，太空板块全年营业"
    "亏损6.57亿美元。猎鹰9号本身仍然盈利——2025年发射约100次，可靠性很高——但星舰开发计划对板块"
    "经济构成显著拖累。政府合同在FY2025贡献了约59亿美元的公司总收入（约32%），提供了相对稳定的收入基础。",
    font_size=10.5, space_after=6)

add_p(doc, "AI板块（xAI / X）", bold=True, font_size=12, space_after=6, space_before=6, heading=True)
add_p(doc,
    "AI板块（2026年2月xAI合并后纳入合并报表）Q1收入约8.15亿美元，但营业亏损约25亿美元。FY2025 AI"
    "亏损总计64亿美元（收入32亿美元）。该板块的支撑来自与Anthropic签订的每月12.5亿美元计算合同"
    "（至2029年5月，终身价值约450亿美元），提供了显著的收入可见性。Colossus数据中心（22万片NVIDIA "
    "GPU）代表了公司成为主要AI基础设施提供商的押注。AI资本开支在FY2025消耗127亿美元、Q1 2026消耗"
    "77亿美元（占总资本开支的76%），反映了该板块的非凡资本密集度。",
    font_size=10.5, space_after=6)

doc.add_page_break()
add_p(doc, "图2 — 板块收入对比", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=4, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart2_segment_revenue.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

add_p(doc, "图3 — 板块盈利能力（年化）", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart3_segment_profitability.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书；分析师计算。")

# ═══════════════ 第4-5页：关键指标与星链 ═══════════════════════════════════
doc.add_page_break()
add_p(doc, "关键指标与星链分析", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "星链运营指标", bold=True, font_size=12, space_after=6, heading=True)

kt = doc.add_table(rows=6, cols=5)
kt.style = "Light Grid Accent 1"
for i, h in enumerate(["指标", "2023年末", "2024年末", "2025年末", "Q1 2026"]):
    kt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["用户数（百万）", "2.3", "4.4", "8.9", "10.3"],
    ["用户同比增长", "—", "+91%", "+102%", "+49%*"],
    ["月度ARPU ($)", "$99", "$85", "$76", "$66"],
    ["ARPU同比变化", "—", "-14%", "-11%", "-13%*"],
    ["覆盖市场数", "~60", "~100", "~150", "164"],
], 1):
    for c, v in enumerate(d): kt.rows[r].cells[c].text = v
fmt_table(kt)
add_src(doc, "*Q1 2026增长率为对比Q1 2025估计值。资料来源：S-1招股说明书。")

add_p(doc, "图4 — 星链用户增长", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart4_starlink_subs.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

add_p(doc, "图5 — 星链ARPU趋势", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart5_starlink_arpu.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

add_p(doc, "图6 — FY2025收入结构", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart6_revenue_mix.png"), width=Inches(4.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

# ═══════════════ 第6-7页：资本配置与现金流 ═════════════════════════════════
doc.add_page_break()
add_p(doc, "资本配置与现金流", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "资本开支分析", bold=True, font_size=12, space_after=6, heading=True)
add_p(doc,
    "SpaceX在FY2025部署了207亿美元资本开支，超过其187亿美元的总收入。Q1 2026资本开支达101亿美元"
    "——年化约400亿美元——主要由AI基础设施支出驱动。AI板块消耗了FY2025资本开支的61%（127亿美元）"
    "和Q1 2026的76%（77亿美元），用于Colossus数据中心建设。连接业务资本开支42亿美元用于卫星发射"
    "和地面基础设施，太空资本开支38亿美元主要用于星舰开发。\n\n"
    "由此产生的自由现金流严重为负：仅Q1 2026就达-91亿美元。IPO后，SpaceX的现金头寸将因约750亿美元"
    "IPO募资而大幅增强，按当前消耗速率可提供约两年的资金跑道。但在不进行额外融资或显著减少AI支出的"
    "情况下维持这一投资水平将面临挑战。",
    font_size=10.5, space_after=6)

add_p(doc, "图7 — FY2025各板块资本开支", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart7_capex_breakdown.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

add_p(doc, "图8 — Q1 2026现金流概览", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart8_cash_flow.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书；分析师计算。")

add_p(doc, "资产负债表要点（IPO前）", bold=True, font_size=12, space_after=6, space_before=10, heading=True)
bt = doc.add_table(rows=6, cols=2)
bt.style = "Light Grid Accent 1"
for i, h in enumerate(["项目", "截至2026年3月31日"]):
    bt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["现金及等价物", "$237亿"], ["长期债务", "$287亿"],
    ["营运资本", "~$53亿"], ["累计亏损", "$(413)亿"],
    ["IPO后备考现金*", "~$980亿"],
], 1):
    for c, v in enumerate(d): bt.rows[r].cells[c].text = v
fmt_table(bt)
add_src(doc, "*备考数据包含约750亿美元IPO募资（扣除承销费用前）。资料来源：S-1招股说明书。")

# ═══════════════ 第8-9页：投资论点与风险 ═══════════════════════════════════
doc.add_page_break()
add_p(doc, "投资论点与风险", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "看多论点", bold=True, font_size=12, space_after=6, heading=True)

add_bullet(doc,
    "星链是一代人的连接特许经营权，具备清晰的500亿美元以上收入路径",
    "凭借1,030万用户（同比增长49%）和约50亿未充分覆盖宽带用户的可触达市场，星链的用户增长跑道"
    "很长。企业、海事、航空和军事/政府应用在消费者之外增加了新的市场层次。如果星链达到5,000万用户、"
    "ARPU $70，则可产生约420亿美元年收入，EBITDA利润率超过60%。")

add_bullet(doc,
    "星舰开启太空经济的全新时代",
    "完全可复用可将发射成本降低10-100倍，催生空间站、太空制造、旅游和火星殖民等新市场。政府合同"
    "（NASA、国防部）在星舰达到商业可行性期间提供近期收入稳定性。")

add_bullet(doc,
    "通过xAI/Colossus的AI基础设施定位拥有450亿美元以上合同收入",
    "仅Anthropic合同就提供约450亿美元三年期收入。如果SpaceX能将xAI定位为顶级AI基础设施提供商，"
    "该板块可能在FY2028-2029达到盈利。")

add_p(doc, "看空论点", bold=True, font_size=12, space_after=6, space_before=8, heading=True)

add_bullet(doc,
    "94倍市销率估值不容许任何失望",
    "以约2.1万亿美元市值对应187亿美元收入，SpaceX的市销率比大型科技公司高10-20倍。即使Palantir"
    "约55倍市销率相比之下都显得便宜。任何执行失误或增长减速都可能触发显著的估值倍数压缩。")

add_bullet(doc,
    "xAI现金消耗惊人，盈利路径不明确",
    "单季度25亿美元营业亏损、77亿美元资本开支——AI板块的现金消耗速度可能在2-3年内消耗大量IPO募资。"
    "计算基础设施市场竞争激烈（AWS、Azure、GCP），Grok面临来自OpenAI、Anthropic和谷歌的强劲竞争。")

add_bullet(doc,
    "马斯克85%投票控制权带来治理风险",
    "马斯克持有42%经济利益但通过双重股权结构掌握85%投票权。投资者实际上无法影响战略决策、董事会"
    "组成或资本配置。马斯克的其他事务（特斯拉、政治活动）造成分心和关键人物风险。")

add_p(doc, "图9 — 各板块调整后EBITDA（FY2025）", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart9_ebitda_segment.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：S-1招股说明书。")

# ═══════════════ 第10页：估值 ══════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "估值与目标价", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "估值框架", bold=True, font_size=12, space_after=6, heading=True)
add_p(doc,
    "SpaceX的估值面临独特挑战。作为三板块综合体（增长和盈利特征差异巨大），分部加总法（SOTP）最为"
    "合适。\n\n"
    "连接业务（星链）——8,000亿-1万亿美元价值：对FY2027E收入160亿美元施以20倍估值，反映星链的"
    "市场主导地位、60%以上EBITDA利润率和强劲用户增长。这高于卫星同业但低于高增长SaaS公司的估值。\n\n"
    "太空——1,500亿-2,000亿美元价值：对FY2027E太空收入50亿美元施以10倍估值，加上1,000亿美元"
    "星舰长期潜力的期权价值（对执行风险和时间表不确定性进行大幅折扣）。\n\n"
    "AI（xAI）——800亿-1,200亿美元价值：2026年2月xAI合并隐含约800亿美元独立估值。鉴于亏损加速"
    "和竞争逆风，我们维持该水平作为基准。\n\n"
    "SOTP公允价值：1.03万亿-1.32万亿美元 → 每股约$80-$100\n\n"
    "目标价$150（高于SOTP）包含50%\"平台溢价\"，反映SpaceX集成太空-连接-AI生态系统的独特战略价值，"
    "但仍低于当前约$161交易水平，体现我们认为IPO后热情已过度反映近期基本面的观点。",
    font_size=10.5, space_after=6)

add_p(doc, "图10 — IPO估值：市销率对比同业", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8, heading=True)
doc.add_picture(os.path.join(OUT, "spcx_chart10_valuation_comps.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "资料来源：Bloomberg、分析师预测。市销率基于过去12个月收入计算。")

# ═══════════════ 资料来源 ══════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "资料来源与参考文献", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8, heading=True)

add_p(doc, "SEC文件：", bold=True, font_size=10.5, space_after=4, heading=True)
for title, url in [
    ("S-1招股说明书（2026年5月20日提交）",
     "https://www.sec.gov/Archives/edgar/data/1181412/000162828026036936/spaceexplorationtechnologi.htm"),
    ("Form 8-K — Q1 2026业绩新闻稿（2026年5月7日提交）",
     "https://www.marketbeat.com/earnings/reports/2026-5-7-spacex-stock/"),
]:
    p = doc.add_paragraph()
    run = p.add_run("• "); set_cn(run, cn=SONG, size=10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

add_p(doc, "", space_after=4)
add_p(doc, "新闻与分析：", bold=True, font_size=10.5, space_after=4, heading=True)
for title, url in [
    ("SpaceX IPO Takes Off — CNBC（2026年6月12日）",
     "https://www.cnbc.com/2026/06/12/spacex-ipo-spcx-live-updates.html"),
    ("SpaceX提交IPO招股书 — Yahoo Finance（2026年5月20日）",
     "https://finance.yahoo.com/markets/article/spacex-files-ipo-prospectus-offering-a-peek-into-its-finances-205406189.html"),
    ("深度解析SpaceX S-1：三家公司，一个利润 — SatNews（2026年5月21日）",
     "https://satnews.com/2026/05/21/inside-spacexs-s-1-three-companies-one-profit-1-75-trillion/"),
    ("SpaceX Pre-IPO财务六图 — Morningstar",
     "https://www.morningstar.com/stocks/6-charts-spacexs-s-1-financials"),
]:
    p = doc.add_paragraph()
    run = p.add_run("• "); set_cn(run, cn=SONG, size=10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

add_p(doc, "", space_after=4)
add_p(doc, "市场数据：", bold=True, font_size=10.5, space_after=4, heading=True)
add_p(doc,
    "• Bloomberg / Yahoo Finance / yfinance 实时市场数据\n"
    "• SpaceX投资者关系：https://investor.spacex.com\n"
    "• Nasdaq上市页面：https://www.nasdaq.com/market-activity/stocks/spcx",
    font_size=10, space_after=6)

add_p(doc, "", space_after=6)
add_p(doc, "免责声明", bold=True, font_size=9, color=(0x75, 0x75, 0x75), space_after=2, heading=True)
add_p(doc,
    "本报告仅供信息参考，不构成投资建议。分析师不是持牌投资顾问。所有预测和目标价均基于公开信息，"
    "代表分析师的独立评估。过往表现不代表未来结果。投资者在做出投资决策前应进行独立的尽职调查。"
    "分析师在SPCX无持仓。",
    font_size=8, color=(0x99, 0x99, 0x99), space_after=4)

path = os.path.join(OUT, "SPCX_Q1_2026_业绩更新报告_中文版.docx")
doc.save(path)
print(f"Report saved: {path}")
