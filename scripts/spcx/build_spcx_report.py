#!/usr/bin/env python3
"""Generate SpaceX (SPCX) Q1 2026 Post-IPO Earnings Update Report (English)."""

import os
import yfinance as yf
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

OUT = os.path.join(os.path.dirname(__file__), "..", "..", "output", "SPCX")
os.makedirs(OUT, exist_ok=True)

# ── Market Data ─────────────────────────────────────────────────────────────
def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        return {
            "price":      round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high":   round(info.year_high, 2),
            "52w_low":    round(info.year_low, 2),
        }
    except Exception as e:
        print(f"Warning: Could not fetch market data for {ticker}: {e}")
        return {"price": "N/A", "market_cap": "N/A", "52w_high": "N/A", "52w_low": "N/A"}

mkt = get_market_data("SPCX")
price_str = f"${mkt['price']}" if mkt['price'] != 'N/A' else '~$161'
mcap_str = f"${mkt['market_cap']/1e12:.2f}T" if mkt['market_cap'] != 'N/A' else '~$2.1T'
hi52 = f"${mkt['52w_high']}" if mkt['52w_high'] != 'N/A' else '~$177'
lo52 = f"${mkt['52w_low']}" if mkt['52w_low'] != 'N/A' else '$135 (IPO)'

# ── Helpers ─────────────────────────────────────────────────────────────────
def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def add_p(doc, text, bold=False, font_size=11, color=None,
          alignment=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if alignment: p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    return p

def add_bullet(doc, header, body, fs=10.5):
    p = doc.add_paragraph()
    r1 = p.add_run("■ "); r1.font.name = "Times New Roman"; r1.font.size = Pt(fs); r1.bold = True
    r2 = p.add_run(header); r2.font.name = "Times New Roman"; r2.font.size = Pt(fs); r2.bold = True
    p.add_run("\n")
    r3 = p.add_run(body); r3.font.name = "Times New Roman"; r3.font.size = Pt(fs)
    p.paragraph_format.space_after = Pt(8)
    return p

def fmt_table(table, hdr_color="003366"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        set_cell_shading(cell, hdr_color)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True; r.font.size = Pt(9); r.font.name = "Times New Roman"
    for i, row in enumerate(table.rows):
        if i == 0: continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9); r.font.name = "Times New Roman"

def add_src(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Times New Roman"; run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75); run.italic = True
    p.paragraph_format.space_after = Pt(4)

def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(
        f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">'
        f'<w:r><w:rPr><w:rStyle w:val="Hyperlink"/><w:color w:val="0563C1"/><w:u w:val="single"/>'
        f'<w:sz w:val="16"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/></w:rPr>'
        f'<w:t>{text}</w:t></w:r></w:hyperlink>')
    paragraph._p.append(hyperlink)

# ── Build Document ──────────────────────────────────────────────────────────
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Times New Roman"; style.font.size = Pt(11)
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

# ════════════════════════════════════════════════════════════════════════════
# PAGE 1: EARNINGS SUMMARY
# ════════════════════════════════════════════════════════════════════════════
add_p(doc, "SPACE EXPLORATION TECHNOLOGIES CORP. (SPCX)", bold=True, font_size=16,
      color=(0x00, 0x33, 0x66), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_p(doc, "Q1 2026 POST-IPO EARNINGS UPDATE", bold=True, font_size=14,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
add_p(doc, "Starlink Prints $1.2B Profit, but xAI Losses Drive $(4.3)B Net Loss",
      bold=True, font_size=12, color=(0x00, 0x33, 0x66),
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

for h in [
    "Date: June 14, 2026",
    "Rating: INITIATE AT MARKET PERFORM",
    f"Price (as of Jun 13, 2026): {price_str}",
    f"Price Target: $150 — 12-Month",
    f"Market Cap: {mcap_str}  |  IPO Price: $135  |  Day-1 Close: ~$161 (+19%)",
]:
    add_p(doc, h, font_size=10, space_after=2)

add_p(doc, "", space_after=4)
add_p(doc, "Q1 2026 FINANCIAL SUMMARY", bold=True, font_size=12,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=6)

# Summary table
st = doc.add_table(rows=7, cols=4)
st.style = "Light Grid Accent 1"
for i, h in enumerate(["Metric", "Q1 2026", "Q1 2025", "YoY Change"]):
    st.rows[0].cells[i].text = h
for r, d in enumerate([
    ["Revenue", "$4.69B", "$4.07B", "+15.4%"],
    ["Adj. EBITDA", "$1.13B", "N/A", "—"],
    ["Operating Income / (Loss)", "$(1.94)B", "$0.03B", "NM"],
    ["Net Income / (Loss)", "$(4.28)B", "$(0.53)B", "NM"],
    ["Basic Loss Per Share", "$(1.27)", "$(0.18)", "NM"],
    ["Free Cash Flow", "$(9.1)B", "N/A", "—"],
], 1):
    for c, v in enumerate(d):
        st.rows[r].cells[c].text = v
fmt_table(st)
add_src(doc, "Source: SpaceX S-1 Registration Statement (filed May 20, 2026); Amendment No.1 (June 1, 2026).")

add_p(doc, "", space_after=2)

add_bullet(doc,
    "Starlink is the crown jewel — $3.3B revenue and $1.2B operating profit in Q1, with 10.3M subscribers",
    "The Connectivity segment generated $3.26B in Q1 revenue (69% of total) with $1.19B operating income "
    "and $2.09B Adjusted EBITDA (64% margin). Starlink subscribers reached 10.3M across 164 markets, up "
    "from 8.9M at year-end 2025. The segment's profitability effectively subsidizes the rest of the business. "
    "However, ARPU has declined from $99/month in 2023 to $66/month as Starlink expands into lower-income "
    "markets and faces competitive pricing pressure, a trend that bears watching.")

add_bullet(doc,
    "xAI losses accelerate: $(2.5)B operating loss in Q1 alone, with $7.7B CapEx on AI infrastructure",
    "The AI segment (xAI, Grok, X platform) posted ~$815M revenue but incurred ~$(2.5)B in operating "
    "losses — annualising to $(10)B. Q1 AI CapEx of $7.7B (76% of total $10.1B CapEx) reflects the "
    "buildout of the Colossus data center (220K NVIDIA GPUs, 300MW). While the $45B Anthropic compute "
    "contract provides revenue visibility, the cash burn rate is extraordinary and the path to segment "
    "profitability remains highly uncertain.")

add_bullet(doc,
    "Space segment revenue declined 28% YoY to $619M, weighed by Starship R&D cycle",
    "Falcon 9 remains the workhorse launch vehicle but Space segment revenue dropped to $619M as the "
    "company invested $3.0B in FY2025 alone on Starship development. The Space segment posted a $(0.6)B "
    "operating loss in Q1 (annualised). While Starship's eventual commercialisation could be transformative, "
    "near-term Space economics are dilutive to overall profitability.")

add_bullet(doc,
    "Initiating at Market Perform with $150 PT — extraordinary business, extraordinary valuation",
    "At ~$2.1T market cap (94x 2025 revenue), SpaceX trades at a significant premium to any comparable "
    "company. Starlink's growth and profitability are genuinely exceptional, but xAI losses are accelerating "
    "and aggregate FCF is deeply negative ($(9.1)B in Q1). Our $150 PT implies ~7% downside from current "
    "levels, reflecting our view that the IPO pop has priced in optimistic scenarios and near-term "
    "earnings visibility is limited by the xAI investment cycle.")

# Estimates Table
add_p(doc, "FINANCIAL ESTIMATES", bold=True, font_size=11,
      color=(0x00, 0x33, 0x66), space_after=4, space_before=10)

et = doc.add_table(rows=8, cols=4)
et.style = "Light Grid Accent 1"
for i, h in enumerate(["Metric", "FY2025A", "FY2026E", "FY2027E"]):
    et.rows[0].cells[i].text = h
for r, d in enumerate([
    ["Revenue ($B)", "$18.7", "$22.0", "$28.0"],
    ["Revenue Growth", "+33%", "+18%", "+27%"],
    ["Adj. EBITDA ($B)", "$6.6", "$5.5", "$9.0"],
    ["Adj. EBITDA Margin", "35.2%", "25.0%", "32.1%"],
    ["Operating Income / (Loss) ($B)", "$(2.6)", "$(6.0)", "$(3.0)"],
    ["Net Income / (Loss) ($B)", "$(4.9)", "$(10.0)", "$(5.0)"],
    ["CapEx ($B)", "$20.7", "$35.0", "$30.0"],
], 1):
    for c, v in enumerate(d):
        et.rows[r].cells[c].text = v
fmt_table(et)
add_src(doc, 'Note: "A" = Actual, "E" = Estimate. Source: S-1 filing (actuals), analyst estimates (forward).')

# ════════════════════════════════════════════════════════════════════════════
# PAGES 2-3: DETAILED RESULTS ANALYSIS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "DETAILED RESULTS ANALYSIS", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "Revenue Analysis", bold=True, font_size=12, space_after=6)
add_p(doc,
    "SpaceX reported consolidated Q1 2026 revenue of $4.694B, up 15.4% YoY from $4.07B in Q1 2025. "
    "The slower topline growth (vs. 33% FY2025 annual growth) reflects a 28% decline in the Space segment "
    "partially offset by continued strong Starlink momentum. Revenue composition continues to shift toward "
    "Connectivity, which now represents 69% of total revenue, up from approximately 63% a year ago.",
    font_size=10.5, space_after=6)

rt = doc.add_table(rows=6, cols=5)
rt.style = "Light Grid Accent 1"
for i, h in enumerate(["Segment", "Q1 2026", "Q1 2025E", "YoY Change", "% of Total"]):
    rt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["Connectivity (Starlink)", "$3,260M", "$2,560M", "+27%", "69%"],
    ["Space (Launch/Dragon)", "$619M", "$860M", "-28%", "13%"],
    ["AI (xAI / X)", "$815M", "$650M", "+25%", "18%"],
    ["Eliminations / Other", "—", "—", "—", "—"],
    ["Total Revenue", "$4,694M", "$4,067M", "+15.4%", "100%"],
], 1):
    for c, v in enumerate(d):
        rt.rows[r].cells[c].text = v
fmt_table(rt)
add_src(doc, "Source: S-1 Registration Statement. Q1 2025 segment estimates derived from YoY growth rates disclosed.")

add_p(doc, "Figure 1 — SpaceX Consolidated Revenue", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart1_annual_revenue.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement (FY2024-2025 actuals); FY2023 estimated from industry sources.")

# Segment Deep Dives
add_p(doc, "Connectivity Segment (Starlink)", bold=True, font_size=12, space_after=6, space_before=10)
add_p(doc,
    "Starlink delivered $3.26B in Q1 revenue with $1.19B operating income (36% margin) and $2.09B "
    "Adjusted EBITDA (64% margin). The subscriber base reached 10.3M as of March 31, 2026, up from "
    "8.9M at year-end 2025 (+16% in one quarter). Starlink operates in 164 markets and controls "
    "approximately 75% of all maneuverable satellites in orbit.\n\n"
    "The key concern is ARPU compression: monthly ARPU declined from $99 in 2023 to $66 in Q1 2026 "
    "(-33% over ~2.5 years). This reflects geographic expansion into emerging markets with lower "
    "pricing, competitive offerings from Amazon Kuiper and OneWeb, and enterprise/government mix shifts. "
    "While subscriber growth more than compensates, the trajectory warrants monitoring as the addressable "
    "market for high-ARPU developed-market broadband approaches saturation.",
    font_size=10.5, space_after=6)

add_p(doc, "Space Segment", bold=True, font_size=12, space_after=6, space_before=6)
add_p(doc,
    "Space segment revenue declined 28% YoY to $619M. The company invested $3.0B in Starship R&D "
    "during FY2025, and the Space segment posted a $(657)M operating loss for the full year. Falcon 9 "
    "remains profitable on a standalone basis — it launched ~100 times in 2025 with high reliability — "
    "but the Starship development program is a significant drag on segment economics. Government "
    "contracts contributed approximately $5.9B of total company revenue in FY2025 (~32%), providing "
    "a relatively stable revenue base.",
    font_size=10.5, space_after=6)

add_p(doc, "AI Segment (xAI / X)", bold=True, font_size=12, space_after=6, space_before=6)
add_p(doc,
    "The AI segment (consolidated following the xAI merger in February 2026) generated ~$815M in Q1 "
    "revenue but posted ~$(2.5)B in operating losses. FY2025 AI losses totaled $(6.4)B on $3.2B "
    "revenue. The segment is anchored by a $1.25B/month compute contract with Anthropic through "
    "May 2029 (~$45B lifetime value), providing significant revenue visibility. The Colossus data "
    "center (220K NVIDIA GPUs) represents the company's bet on becoming a major AI infrastructure "
    "provider. AI CapEx consumed $12.7B in FY2025 and $7.7B in Q1 2026 alone (76% of total CapEx), "
    "reflecting the extraordinary capital intensity of this segment.",
    font_size=10.5, space_after=6)

doc.add_page_break()
add_p(doc, "Figure 2 — Revenue by Segment", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=4)
doc.add_picture(os.path.join(OUT, "spcx_chart2_segment_revenue.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

add_p(doc, "Figure 3 — Segment Profitability (Annualised)", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart3_segment_profitability.png"), width=Inches(5.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement; analyst calculations.")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 4-5: KEY METRICS & STARLINK
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "KEY METRICS & STARLINK ANALYSIS", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "Starlink Operating Metrics", bold=True, font_size=12, space_after=6)

kt = doc.add_table(rows=6, cols=5)
kt.style = "Light Grid Accent 1"
for i, h in enumerate(["Metric", "YE 2023", "YE 2024", "YE 2025", "Q1 2026"]):
    kt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["Subscribers (M)", "2.3", "4.4", "8.9", "10.3"],
    ["YoY Sub Growth", "—", "+91%", "+102%", "+49%*"],
    ["Monthly ARPU ($)", "$99", "$85", "$76", "$66"],
    ["ARPU YoY Change", "—", "-14%", "-11%", "-13%*"],
    ["Markets Served", "~60", "~100", "~150", "164"],
], 1):
    for c, v in enumerate(d):
        kt.rows[r].cells[c].text = v
fmt_table(kt)
add_src(doc, "*Q1 2026 growth rates are vs. Q1 2025 estimated. Source: S-1 Registration Statement.")

add_p(doc, "Figure 4 — Starlink Subscriber Growth", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart4_starlink_subs.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

add_p(doc, "Figure 5 — Starlink ARPU Trend", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart5_starlink_arpu.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

add_p(doc, "Figure 6 — FY2025 Revenue Mix", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart6_revenue_mix.png"), width=Inches(4.8))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 6-7: CAPITAL ALLOCATION & CASH FLOW
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "CAPITAL ALLOCATION & CASH FLOW", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "Capital Expenditure Analysis", bold=True, font_size=12, space_after=6)
add_p(doc,
    "SpaceX deployed $20.7B in CapEx during FY2025, exceeding its total revenue of $18.7B. In Q1 2026, "
    "CapEx reached $10.1B — annualising to ~$40B — driven overwhelmingly by AI infrastructure spending. "
    "The AI segment consumed $12.7B (61%) of FY2025 CapEx and $7.7B (76%) of Q1 2026 CapEx for the "
    "Colossus data center buildout. Connectivity CapEx of $4.2B funded satellite launches and ground "
    "infrastructure, while Space CapEx of $3.8B primarily funded Starship development.\n\n"
    "The resulting free cash flow was deeply negative: $(9.1)B in Q1 2026 alone. Post-IPO, SpaceX's "
    "cash position will be significantly bolstered by the ~$75B in IPO proceeds, providing approximately "
    "two years of runway at the current burn rate. However, sustaining this level of investment without "
    "additional capital raises or a meaningful reduction in AI spending will be challenging.",
    font_size=10.5, space_after=6)

add_p(doc, "Figure 7 — FY2025 CapEx by Segment", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart7_capex_breakdown.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

add_p(doc, "Figure 8 — Q1 2026 Cash Flow Overview", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart8_cash_flow.png"), width=Inches(5.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement; analyst calculations.")

# Balance Sheet
add_p(doc, "Balance Sheet Highlights (Pre-IPO)", bold=True, font_size=12, space_after=6, space_before=10)
bt = doc.add_table(rows=6, cols=2)
bt.style = "Light Grid Accent 1"
for i, h in enumerate(["Item", "As of Mar 31, 2026"]):
    bt.rows[0].cells[i].text = h
for r, d in enumerate([
    ["Cash & Cash Equivalents", "$23.7B"],
    ["Long-Term Debt", "$28.7B"],
    ["Working Capital", "~$5.3B"],
    ["Accumulated Deficit", "$(41.3)B"],
    ["Post-IPO Pro-Forma Cash*", "~$98B"],
], 1):
    for c, v in enumerate(d):
        bt.rows[r].cells[c].text = v
fmt_table(bt)
add_src(doc, "*Pro-forma includes ~$75B IPO proceeds (before underwriting fees). Source: S-1 Registration Statement.")

# ════════════════════════════════════════════════════════════════════════════
# PAGES 8-9: INVESTMENT THESIS & RISKS
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "INVESTMENT THESIS & RISKS", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "Bull Case", bold=True, font_size=12, space_after=6)

add_bullet(doc,
    "Starlink is a generational connectivity franchise with a clear path to $50B+ revenue",
    "With 10.3M subscribers growing 49% YoY and an addressable market of ~5B underserved broadband "
    "users, Starlink has a long runway for subscriber growth. Enterprise, maritime, aviation, and "
    "government/military applications add TAM layers beyond consumer. If Starlink reaches 50M "
    "subscribers at $70 ARPU, it generates ~$42B annual revenue with 60%+ EBITDA margins.")

add_bullet(doc,
    "Starship unlocks an entirely new era of space economics",
    "Full reusability could reduce launch costs by 10-100x, enabling new markets in space stations, "
    "manufacturing, tourism, and Mars colonisation. Government contracts (NASA, DoD) provide near-term "
    "revenue stability while Starship reaches commercial viability.")

add_bullet(doc,
    "AI infrastructure positioning via xAI/Colossus has $45B+ contracted revenue",
    "The Anthropic contract alone provides ~$45B over three years. If SpaceX can position xAI as "
    "a top-tier AI infrastructure provider, the segment could reach profitability by FY2028-2029.")

add_p(doc, "Bear Case", bold=True, font_size=12, space_after=6, space_before=8)

add_bullet(doc,
    "94x P/S valuation leaves zero room for disappointment",
    "At ~$2.1T market cap on $18.7B revenue, SpaceX trades at a P/S multiple 10-20x higher than "
    "large-cap tech. Even Palantir at ~55x P/S looks cheap by comparison. Any execution misstep "
    "or growth deceleration could trigger significant multiple compression.")

add_bullet(doc,
    "xAI cash burn is extraordinary and path to profitability is unclear",
    "$(2.5)B operating loss in a single quarter, $7.7B CapEx — the AI segment is burning cash at a "
    "rate that could consume a large portion of the IPO proceeds within 2-3 years. The compute "
    "infrastructure market is intensely competitive (AWS, Azure, GCP), and Grok faces strong "
    "competition from OpenAI, Anthropic, and Google.")

add_bullet(doc,
    "Elon Musk's 85% voting control creates governance risk",
    "Musk holds 42% economic interest but 85% voting power via dual-class shares. Investors have "
    "no practical ability to influence strategic decisions, board composition, or capital allocation. "
    "Musk's other commitments (Tesla, political activities) create distraction and key-man risk.")

add_p(doc, "Figure 9 — Adjusted EBITDA by Segment (FY2025)", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart9_ebitda_segment.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: S-1 Registration Statement.")

# ════════════════════════════════════════════════════════════════════════════
# PAGE 10: VALUATION
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "VALUATION & PRICE TARGET", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "Valuation Framework", bold=True, font_size=12, space_after=6)
add_p(doc,
    "SpaceX presents unique valuation challenges. As a three-segment conglomerate with vastly different "
    "growth and profitability profiles, a sum-of-the-parts (SOTP) approach is most appropriate.\n\n"
    "Connectivity (Starlink) — $800B-$1.0T value: We apply 20x FY2027E revenue of $16B, reflecting "
    "Starlink's dominant market position, 60%+ EBITDA margins, and strong subscriber growth. This "
    "represents a premium to satellite peers but a discount to high-growth SaaS multiples.\n\n"
    "Space — $150B-$200B value: We apply 10x FY2027E Space revenue of $5B, plus a $100B option value "
    "for Starship's long-term potential (heavily discounted for execution risk and timeline uncertainty).\n\n"
    "AI (xAI) — $80B-$120B value: The February 2026 xAI merger implied ~$80B standalone valuation. "
    "Given accelerating losses and competitive headwinds, we maintain this level as our base case.\n\n"
    "SOTP Fair Value: $1.03T-$1.32T → ~$80-$100/share\n\n"
    "Our $150 price target (above SOTP) incorporates a 50% \"platform premium\" for the unique strategic "
    "value of SpaceX's integrated space-connectivity-AI ecosystem, but remains below the current "
    "~$161 trading level, reflecting our view that post-IPO enthusiasm has overshot near-term fundamentals.",
    font_size=10.5, space_after=6)

add_p(doc, "Figure 10 — IPO Valuation: P/S Multiple vs. Peers", bold=True, font_size=10,
      alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2, space_before=8)
doc.add_picture(os.path.join(OUT, "spcx_chart10_valuation_comps.png"), width=Inches(5.5))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
add_src(doc, "Source: Bloomberg, analyst estimates. P/S based on trailing twelve-month revenue.")

# ════════════════════════════════════════════════════════════════════════════
# SOURCES & REFERENCES
# ════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_p(doc, "SOURCES & REFERENCES", bold=True, font_size=14,
      color=(0x00, 0x33, 0x66), space_after=8)

add_p(doc, "SEC Filings:", bold=True, font_size=10.5, space_after=4)
sec_items = [
    ("S-1 Registration Statement (Filed May 20, 2026)",
     "https://www.sec.gov/Archives/edgar/data/1181412/000162828026036936/spaceexplorationtechnologi.htm"),
    ("Form 8-K — Q1 2026 Press Release (Filed May 7, 2026)",
     "https://www.marketbeat.com/earnings/reports/2026-5-7-spacex-stock/"),
]
for title, url in sec_items:
    p = doc.add_paragraph()
    run = p.add_run("• "); run.font.name = "Times New Roman"; run.font.size = Pt(10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

add_p(doc, "", space_after=4)
add_p(doc, "News & Analysis:", bold=True, font_size=10.5, space_after=4)
news_items = [
    ("SpaceX IPO Takes Off — CNBC (June 12, 2026)",
     "https://www.cnbc.com/2026/06/12/spacex-ipo-spcx-live-updates.html"),
    ("SpaceX files IPO prospectus — Yahoo Finance (May 20, 2026)",
     "https://finance.yahoo.com/markets/article/spacex-files-ipo-prospectus-offering-a-peek-into-its-finances-205406189.html"),
    ("Inside SpaceX's S-1: Three Companies, One Profit — SatNews (May 21, 2026)",
     "https://satnews.com/2026/05/21/inside-spacexs-s-1-three-companies-one-profit-1-75-trillion/"),
    ("6 Charts on SpaceX's Pre-IPO Financials — Morningstar",
     "https://www.morningstar.com/stocks/6-charts-spacexs-s-1-financials"),
    ("SpaceX Just Revealed Its Finances — Yahoo Finance",
     "https://finance.yahoo.com/markets/stocks/articles/spacex-just-revealed-finances-warning-203600433.html"),
]
for title, url in news_items:
    p = doc.add_paragraph()
    run = p.add_run("• "); run.font.name = "Times New Roman"; run.font.size = Pt(10)
    add_hyperlink(p, url, title)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.left_indent = Cm(0.5)

add_p(doc, "", space_after=4)
add_p(doc, "Market Data:", bold=True, font_size=10.5, space_after=4)
add_p(doc,
    "• Bloomberg / Yahoo Finance / yfinance for real-time market data\n"
    "• SpaceX Investor Relations: https://investor.spacex.com\n"
    "• Nasdaq listing page: https://www.nasdaq.com/market-activity/stocks/spcx",
    font_size=10, space_after=6)

# Disclaimer
add_p(doc, "", space_after=6)
add_p(doc, "DISCLAIMER", bold=True, font_size=9, color=(0x75, 0x75, 0x75), space_after=2)
add_p(doc,
    "This report is for informational purposes only and does not constitute investment advice. "
    "The analyst is not a licensed investment advisor. All estimates and price targets are based on "
    "publicly available information and represent the analyst's independent assessment. Past performance "
    "is not indicative of future results. Investors should conduct their own due diligence before making "
    "investment decisions. The analyst has no position in SPCX.",
    font_size=8, color=(0x99, 0x99, 0x99), space_after=4)

# ── Save ────────────────────────────────────────────────────────────────────
path = os.path.join(OUT, "SPCX_Q1_2026_Earnings_Update.docx")
doc.save(path)
print(f"Report saved: {path}")
