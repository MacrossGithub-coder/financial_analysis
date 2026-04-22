#!/usr/bin/env python3
"""
阳光电源 (300274.SZ) Q4 2025 / FY2025 业绩更新报告 — 中文版 DOCX
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/300274"
CHARTS = BASE
OUT  = os.path.join(BASE, "300274_Q4_FY2025_业绩更新报告_中文版.docx")

# ── helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part   = paragraph.part
    r_id   = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl     = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r      = OxmlElement("w:r")
    rPr    = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)
    return hl

def para_fmt(p, left_indent=0, space_before=0, space_after=6, line_spacing=None):
    pf = p.paragraph_format
    pf.left_indent  = Pt(left_indent)
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)

def add_run(para, text, bold=False, italic=False, size=10, color=None, zh=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    # CJK font
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    if zh:
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
    else:
        rFonts.set(qn("w:eastAsia"), "宋体")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_bold_heading(para, text, size=11, color=(31, 78, 121)):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "黑体")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    run.font.color.rgb = RGBColor(*color)
    return run

def add_heading_para(doc, text, size=12, color=(31, 78, 121)):
    p = doc.add_paragraph()
    para_fmt(p, space_before=8, space_after=4)
    add_bold_heading(p, text, size=size, color=color)
    return p

def add_divider(doc, color="1F4E79"):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pb  = p._p
    pPr = pb.get_or_add_pPr()
    pBd = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"),  "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBd.append(bot)
    pPr.append(pBd)
    return p

def add_chart(doc, filename, width=6.0, caption=""):
    path = os.path.join(CHARTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        para_fmt(p, space_before=4, space_after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    else:
        p = doc.add_paragraph(f"[图表未找到: {filename}]")
    if caption:
        cp = doc.add_paragraph()
        para_fmt(cp, space_before=0, space_after=6)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cp, caption, size=8, italic=True, color=(128, 128, 128))
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg="1F4E79", alt_bg="E9EEF4"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8)
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:eastAsia"), "黑体")
            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rPr.append(rFonts)
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri+1].cells
        bg = alt_bg if ri % 2 == 1 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            cells[ci].text = str(cell_text)
            if bg != "FFFFFF":
                set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(8)
                rPr = run._r.get_or_add_rPr()
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:eastAsia"), "宋体")
                rFonts.set(qn("w:ascii"), "Times New Roman")
                rFonts.set(qn("w:hAnsi"), "Times New Roman")
                rPr.append(rFonts)
    if col_widths:
        for ri2, row2 in enumerate(tbl.rows):
            for ci2, (cell2, w) in enumerate(zip(row2.cells, col_widths)):
                cell2.width = Inches(w)
    return tbl

def src_note(doc, text):
    p = doc.add_paragraph(text)
    para_fmt(p, space_before=1, space_after=6)
    for run in p.runs:
        run.font.size = Pt(7.5)
        run.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
    return p

# ═════════════════════════════════════════════════════════════════════════════
# 构建文档
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(27.94)

# ─────────────────────────────────────────────────────────────────────────────
# 第1页：业绩摘要
# ─────────────────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
para_fmt(title_p, space_before=4, space_after=2)
add_bold_heading(title_p, "阳光电源股份有限公司（300274.SZ）", size=14, color=(31, 78, 121))

sub_p = doc.add_paragraph()
para_fmt(sub_p, space_before=0, space_after=2)
add_bold_heading(sub_p, "2025年第四季度 / 2025年全年业绩更新报告", size=11, color=(46, 117, 182))

meta_p = doc.add_paragraph()
para_fmt(meta_p, space_before=0, space_after=8)
add_run(meta_p, "2026年4月22日  ·  股票研究  ·  中国新能源 / 新能源设备行业", size=9, italic=True, color=(128, 128, 128))

add_divider(doc)

# 评级/目标价表
rtbl = doc.add_table(rows=1, cols=4)
rtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
rtbl.style = "Table Grid"
rt_cells = rtbl.rows[0].cells
rt_data = [
    ("评级", "跑赢大盘 ★", "1F4E79"),
    ("股价（2026/4/21）", "136.91元", "2E75B6"),
    ("目标价", "158.00元", "375623"),
    ("隐含上涨空间", "+15.4%", "375623"),
]
for cell, (label, val, color_hex) in zip(rt_cells, rt_data):
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "\n")
    r1.bold = False; r1.font.size = Pt(8)
    r1.font.color.rgb = RGBColor(200, 220, 240)
    rPr1 = r1._r.get_or_add_rPr()
    rF1 = OxmlElement("w:rFonts"); rF1.set(qn("w:eastAsia"), "宋体")
    rPr1.append(rF1)
    r2 = p.add_run(val)
    r2.bold = True; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(255, 255, 255)
    rPr2 = r2._r.get_or_add_rPr()
    rF2 = OxmlElement("w:rFonts"); rF2.set(qn("w:eastAsia"), "黑体")
    rPr2.append(rF2)

doc.add_paragraph()

# 业绩快照表
snap_h_p = doc.add_paragraph()
para_fmt(snap_h_p, space_before=4, space_after=4)
add_bold_heading(snap_h_p, "2025年Q4及全年业绩快照", size=10.5, color=(31, 78, 121))

snap_headers = ["指标", "2025年实际", "市场预期（一致预期）", "超预期/（低于预期）", "同比变动"]
snap_rows = [
    ["营业收入（亿元）",         "891.84亿",  "900亿",   "（8.16亿/-0.9%）",  "+14.6%"],
    ["归母净利润（亿元）",        "134.61亿",  "154.52亿","（19.91亿/-12.9%）","+22.0%"],
    ["基本每股收益（元）",        "6.55元",    "7.50元E", "（0.95元/-12.7%）", "+23.1%"],
    ["毛利率（%）",              "31.83%",    "约30%",   "+183bps",           "+631bps"],
    ["净利率（%）",              "15.17%",    "约17%",   "-183bps",           "+488bps"],
    ["Q4营收（亿元）",           "227.82亿",  "约240亿", "（约12亿/-5.1%）",  "-18.4%"],
    ["Q4归母净利润（亿元）",      "15.80亿",   "约35亿",  "（约19亿/-54.9%）", "-54.0%"],
]
make_table(doc, snap_headers, snap_rows, col_widths=[2.0, 1.2, 1.6, 1.6, 1.0])
src_note(doc, "资料来源：阳光电源2025年年度报告（2026-04-01）；市场一致预期来自Bloomberg/富途（发布前）")

doc.add_paragraph()

# 核心观点
kt_p = doc.add_paragraph()
para_fmt(kt_p, space_before=4, space_after=4)
add_bold_heading(kt_p, "核心观点", size=10.5, color=(31, 78, 121))

bullets_cn = [
    ("■ 业绩低于预期 — 全年净利润低于一致预期12.9%；Q4单季利润同比骤降54%",
     "2025年全年营收891.84亿元，基本符合预期（仅低0.9%），但归母净利润134.61亿元较市场预期154.52亿元低近20亿元，"
     "差距几乎完全源自Q4的利润率急剧恶化。Q4归母净利润仅15.8亿元，同比下降54%、环比下降62%，为2023年Q2以来最低季度利润。"
     "具体归因：①收入结构偏移至低毛利的新能源投资开发项目（Q4集中交付）；②储能毛利率从Q3约41%骤降约17个百分点至约24%，"
     "主因Q3确认了大批高毛利海外项目；③碳酸锂价格Q4有所上涨，存量固定价格合同无法及时传导成本。"),
    ("■ 全年整体业绩稳健 — 营收+14.6%，净利润+22.0%，毛利率创历史新高",
     "全年营收及利润均实现双位数增长，毛利率31.83%（同比+6.31pcts）和净利率15.17%（同比+4.88pcts）均创近年新高，"
     "体现了以高毛利海外储能项目为主的产品结构升级成效。经营现金流每股达8.16元（同比+40.2%），投资回报率ROIC升至25.01%。"),
    ("■ 储能仍是核心增长引擎 — FY2025出货43GWh（+54% YoY），海外+89%；2026目标超60GWh",
     "2025年储能系统出货量43GWh，同比增长53.6%，其中海外出货36GWh（同比+89%），占总出货量84%。"
     "公司在2026年3月31日电话会议上指引2026年储能出货目标超60GWh（同比+40%），并预计全球储能市场增速为30%-50%。"
     "国内储能出货7GWh（同比-22%），受国内公用事业储能项目节奏较慢影响，但不改变全球化战略方向。"),
    ("■ 维持「跑赢大盘」评级，目标价下调至158元（原175元）；Q4低迷不改长期逻辑",
     "Q4利润率冲击属于一次性结构因素叠加原材料时序影响，而非竞争格局或需求端的实质恶化。"
     "我们下调2026年净利润预测至167.7亿元（原180亿元），并相应调整目标价至158元（对应20x 2026E市盈率）。"
     "当前股价136.91元隐含15.4%上涨空间，维持「跑赢大盘」评级。"),
]

for header_cn, body_cn in bullets_cn:
    h_p = doc.add_paragraph()
    para_fmt(h_p, space_before=4, space_after=2)
    add_bold_heading(h_p, header_cn, size=9.5, color=(31, 78, 121))
    b_p = doc.add_paragraph()
    para_fmt(b_p, left_indent=12, space_before=0, space_after=8)
    add_run(b_p, body_cn, size=9.5)

# 预测更新表
add_divider(doc)
ue_p = doc.add_paragraph()
para_fmt(ue_p, space_before=6, space_after=4)
add_bold_heading(ue_p, "财务预测更新", size=10.5, color=(31, 78, 121))

ue_headers_cn = ["指标", "2024年实际", "2025年实际", "2025年预期（前）", "2026年预测（新）", "2026年预测（前）"]
ue_rows_cn = [
    ["营业收入（亿元）",   "778.6",  "891.8",   "900.0",   "1020.0",  "1080.0"],
    ["收入增速（%）",      "+7.8%",  "+14.6%",  "+15.6%",  "+14.4%",  "+21.1%"],
    ["毛利率（%）",        "25.5%",  "31.83%",  "约30%",   "32.0%",   "32.5%"],
    ["归母净利润（亿元）", "110.4",  "134.61",  "154.52",  "167.7",   "180.0"],
    ["净利润增速（%）",    "+16.9%", "+22.0%",  "+40.0%",  "+24.6%",  "+33.6%"],
    ["基本EPS（元）",      "5.32",   "6.55",    "7.50",    "8.09",    "8.70"],
    ["P/E @ 136.91元",     "25.7x",  "20.9x",   "—",       "16.9x",   "—"],
]
make_table(doc, ue_headers_cn, ue_rows_cn, col_widths=[1.9, 1.0, 1.0, 1.3, 1.3, 1.3])
src_note(doc, "注：前期预测为发布年报前内部测算。A=实际值，E=预测值。资料来源：阳光电源2025年年报；分析师预测")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 第2-3页：详细业绩分析
# ─────────────────────────────────────────────────────────────────────────────
add_heading_para(doc, "第一部分：详细业绩分析", size=12)
add_divider(doc)

add_chart(doc, "300274_chart1_revenue.png", width=6.2,
          caption="图1 — 季度营收走势 | 资料来源：阳光电源2025年年度报告；StockAnalysis.com")

add_heading_para(doc, "1.1  营收表现", size=11, color=(46, 117, 182))

rev_paras_cn = [
    ("FY2025全年营收891.84亿元（同比+14.6%）——基本符合预期。",
     "2025年全年营收较2024年777.86亿元增长14.6%，略低于市场一致预期的900亿元（差距约8亿元，-0.9%）。"
     "营收增速较2023年79.5%的高增速出现较大幅度回落，主要反映基数效应及国内市场阶段性疲软。"
     "从结构来看，储能系统占比由2024年约32%提升至41.8%，逆变器占比由37%降至34.9%，"
     "新能源投资开发占比由27%大幅下降至18.6%（主要源于2025年全年该业务收入165.6亿元，同比-21.1%）。"),
    ("Q4 2025单季营收227.82亿元（同比-18.4%，环比-0.4%）——同比大幅下滑。",
     "Q4营收227.82亿元，较2024年Q4的279.1亿元大幅下降18.4%，是2025年内单季营收最低点（Q1为190.4亿元）。"
     "同比下滑主要源于以下因素：①2024年Q4基数较高（同比+8.0%）；②国内逆变器出货量全年同比下降约19%，"
     "Q4受此影响明显；③储能海外项目节奏集中于前三季度，Q4确认量有所下降。"
     "环比基本持平（-0.4% vs Q3的228.7亿元），表明Q4的同比下滑主要是高基数正常化，而非需求实质下滑。"),
]
for (bold_text, body_text) in rev_paras_cn:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_bold_heading(p, bold_text, size=9.5, color=(31, 78, 121))
    add_run(p, body_text, size=9.5)

# 分业务营收拆分表
rb_h_cn = ["业务板块", "2024年（亿元）", "2024年占比", "2025年（亿元）", "2025年占比", "同比增速"]
rb_r_cn = [
    ["储能系统",              "约249",  "32.0%", "372.87", "41.8%", "+49.7%"],
    ["光伏逆变器及电力电子设备", "约288", "37.0%", "311.36", "34.9%", "+8.1%"],
    ["新能源投资开发",         "约210",  "27.0%", "165.59", "18.6%", "-21.1%"],
    ["其他业务",              "约32",   "4.0%",  "42.02",  "4.7%",  "+31.3%"],
    ["合计",                  "778.6",  "100%",  "891.84", "100%",  "+14.6%"],
]
make_table(doc, rb_h_cn, rb_r_cn, col_widths=[2.3, 1.2, 0.9, 1.2, 0.9, 1.0])
src_note(doc, "资料来源：阳光电源2025年年度报告（2026-04-01）；2024年分业务数据为基于公司披露的估算值")

add_chart(doc, "300274_chart4_segments.png", width=6.2,
          caption="图4 — 各业务板块营收对比（2024A vs 2025A）| 资料来源：阳光电源2025年年度报告")
add_chart(doc, "300274_chart2_net_profit.png", width=6.2,
          caption="图2 — 季度归母净利润走势 | 资料来源：阳光电源2025年年度报告；分析师季度拆分估算")

add_heading_para(doc, "1.2  盈利能力分析", size=11, color=(46, 117, 182))

prof_paras_cn = [
    ("FY2025全年毛利率31.83%（同比+6.31pcts）——盈利能力大幅提升。",
     "全年毛利率31.83%较2024年25.52%提升超6个百分点，主要驱动因素包括：①海外高毛利储能项目占比持续提升，"
     "海外项目议价能力优于国内；②产品结构优化，高附加值储能产品取代低毛利国内逆变器成为营收主体；"
     "③规模效应与运营杠杆。经营现金流每股8.16元（+40.2%），ROIC达25.01%，验证了资本配置效率的持续改善。"),
    ("Q4 2025归母净利润15.8亿元（同比-54%）——三重负面因素叠加。",
     "Q4归母净利润15.8亿元是2025年最为突出的业绩亮点（负面），较市场预期约35亿元存在约19亿元差距，"
     "远低于Q4 2024的34.4亿元（同比-54%），亦大幅低于Q3 2025的41.47亿元（环比-62%）。"
     "管理层在2026年3月31日电话会上明确了三大原因：\n"
     "① 收入结构变化：新能源投资开发业务大项目集中在Q4交付，拉低了整体毛利率；\n"
     "② 储能毛利率环比大幅下降约17个百分点至约24%，主因Q3确认了海外高毛利项目导致基数偏高；\n"
     "③ Q4碳酸锂价格有一定涨幅，存量项目价格无法及时传导，挤压了固定价格合同的利润空间。"),
]
for (bold_text, body_text) in prof_paras_cn:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_bold_heading(p, bold_text, size=9.5, color=(31, 78, 121))
    add_run(p, body_text, size=9.5)

mg_h_cn = ["利润率指标", "Q1'25A", "Q2'25A", "Q3'25A", "Q4'25A", "FY2025A", "FY2024A", "同比变动"]
mg_r_cn = [
    ["毛利率（%）",     "29.4%", "32.1%", "35.9%", "24.2%", "31.83%", "25.52%", "+631bps"],
    ["净利率（%）",     "17.9%", "17.7%", "18.1%",  "6.9%", "15.17%", "10.29%", "+488bps"],
    ["ROIC（%）",       "—",     "—",     "—",      "—",    "25.01%", "约20%",  "+约500bps"],
    ["每股经营性现金流","—",     "—",     "—",      "—",    "8.16元", "5.82元", "+40.2%"],
]
make_table(doc, mg_h_cn, mg_r_cn, col_widths=[1.8, 0.8, 0.8, 0.8, 0.8, 0.9, 0.9, 0.9])
src_note(doc, "资料来源：阳光电源2025年年度报告；公司投资者电话会议（2026-03-31）")

add_chart(doc, "300274_chart3_margins.png", width=6.2,
          caption="图3 — 季度利润率走势 | 资料来源：阳光电源2025年年度报告；公司投资者电话会议")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 第4-5页：核心运营指标与业绩指引
# ─────────────────────────────────────────────────────────────────────────────
add_heading_para(doc, "第二部分：核心运营指标与业绩指引", size=12)
add_divider(doc)

add_heading_para(doc, "2.1  出货量表现", size=11, color=(46, 117, 182))

ship_paras_cn = [
    ("储能系统：43GWh（同比+53.6%）——全球龙头地位进一步巩固。",
     "2025年储能系统出货43GWh，超出我们约40GWh的预测，同比增速达53.6%。"
     "其中海外出货36GWh（同比+89%），占总出货量84%，充分体现了阳光电源海外市场拓展的显著成效。"
     "国内储能出货7GWh（同比-22%），主要受制于国内公用事业储能项目开工节奏放缓。"
     "国内外出货量的反向走势清晰印证了公司向更高毛利国际市场转型的战略成功。"),
    ("光伏逆变器：143GW（同比-2.7%）——国内承压拖累总量。",
     "2025年逆变器出货量143GW，较2024年147GW小幅下降2.7%，其中国内出货量57GW（同比-19%），"
     "反映国内太阳能装机节奏放缓及竞争加剧。但逆变器收入仍增长8.1%至311.36亿元，"
     "体现了产品结构优化（高价值组串逆变器及大型地面电站解决方案）及海外定价稳定的正面影响。"),
]
for (bold_text, body_text) in ship_paras_cn:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_bold_heading(p, bold_text, size=9.5, color=(31, 78, 121))
    add_run(p, body_text, size=9.5)

ship_h_cn = ["指标", "2023年A", "2024年A", "2025年A", "同比增速", "2026年目标"]
ship_r_cn = [
    ["储能系统（GWh）",  "16.0", "28.0", "43.0", "+53.6%", ">60GWh"],
    ["— 国内（GWh）",   "7.0",  "9.0",  "7.0",  "-22.2%", "约10GWh"],
    ["— 海外（GWh）",   "9.0",  "19.0", "36.0", "+89.5%", ">50GWh"],
    ["光伏逆变器（GW）", "130",  "147",  "143",  "-2.7%",  "约155GW"],
    ["— 国内（GW）",    "n/a",  "71",   "57",   "-19.7%", "约60GW"],
    ["— 海外（GW）",    "n/a",  "76",   "86",   "+13.2%", ">90GW"],
]
make_table(doc, ship_h_cn, ship_r_cn, col_widths=[2.0, 0.9, 0.9, 0.9, 1.0, 1.3])
src_note(doc, "资料来源：阳光电源2025年年度报告（2026-04-01）；2026年目标来自管理层指引（2026-03-31电话会议）")

add_chart(doc, "300274_chart5_storage_shipments.png", width=6.0,
          caption="图5 — 储能系统出货量（GWh）| 资料来源：阳光电源2025年年度报告；管理层指引")
add_chart(doc, "300274_chart6_inverter_shipments.png", width=5.5,
          caption="图6 — 光伏逆变器出货量（GW）| 资料来源：阳光电源2025年年度报告")

add_heading_para(doc, "2.2  管理层指引与展望", size=11, color=(46, 117, 182))

guid_h_cn = ["指标", "FY2025A", "2026年指引/目标", "我们的预测（FY2026E）", "评估"]
guid_r_cn = [
    ["储能出货量（GWh）",    "43GWh",    ">60GWh（+40% YoY）", "62GWh",     "可信；市场需求驱动强劲"],
    ["全球储能市场增速",      "—",        "30%-50%",           "约40%",     "可实现；历史执行力支撑"],
    ["营业收入（亿元）",     "891.84亿", "无正式指引",         "1020亿",    "同比+14.4%估算"],
    ["归母净利润（亿元）",   "134.61亿", "无正式指引",         "167.7亿",   "同比+24.6%估算"],
    ["储能毛利率",            "约31.8%",  "H1 2026预期恢复",    "全年>30%", "若碳酸锂价格企稳，可实现"],
]
make_table(doc, guid_h_cn, guid_r_cn, col_widths=[1.8, 1.2, 1.6, 1.5, 1.5])
src_note(doc, "资料来源：阳光电源投资者电话会议（2026-03-31）；分析师预测")

guid_p_cn = doc.add_paragraph()
para_fmt(guid_p_cn, space_before=6, space_after=6)
add_bold_heading(guid_p_cn, "管理层关于Q4低迷的解释：", size=9.5, color=(31, 78, 121))
add_run(guid_p_cn,
        "管理层在电话会上对Q4业绩低迷进行了详细解释，明确指出三大因素均具有一次性、阶段性特征，"
        "公司内在竞争力与商业模式未受影响。具体而言：①新能源投资开发业务大项目集中在Q4确认收入，"
        "该业务毛利率低于储能产品，拉低整体利润率；②储能业务毛利率Q4仅约24%，"
        "环比下降约17个百分点，主因Q3集中确认了多个高毛利海外大型项目，形成较高基数；"
        "③碳酸锂Q4价格上涨，存量固定价格合同在成本端承压。管理层表示，随着碳酸锂价格趋于平稳、"
        "新海外项目签约单价改善，2026年上半年储能毛利率有望回升，并对实现全年指引目标表达了信心。",
        size=9.5)

add_chart(doc, "300274_chart8_storage_margins.png", width=6.0,
          caption="图8 — 储能业务季度毛利率分析（Q4压缩原因）| 资料来源：阳光电源2025年年度报告；投资者电话会议")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 第6-7页：投资逻辑更新
# ─────────────────────────────────────────────────────────────────────────────
add_heading_para(doc, "第三部分：投资逻辑更新", size=12)
add_divider(doc)

intro_cn = doc.add_paragraph()
para_fmt(intro_cn, space_before=4, space_after=6)
add_run(intro_cn,
        "我们对阳光电源的投资逻辑建立在三大支柱之上：①全球储能市场领导地位与海外扩张；"
        "②逆变器业务的持续竞争优势；③全球能源转型带来的长期结构性需求。"
        "我们据此评估本次年报对各支柱的影响。",
        size=9.5)

thesis_cn = [
    ("■ 支柱一：全球储能市场领导地位 — 进一步强化",
     "进一步强化",
     "2025年储能出货43GWh（+53.6%），海外出货量同比增长89%，阳光电源全球储能龙头地位得到充分印证。"
     "海外收入占比提升至储能板块主体，有效规避了国内市场的价格竞争压力。2026年超60GWh的指引（+40% YoY）"
     "有强劲的全球项目储备支撑，包括美国、欧洲、澳大利亚及中东市场。Q4毛利率压缩是时序因素，"
     "并非竞争格局或需求端的实质恶化——若非如此，公司不会同时给出2026年超60GWh的高增速指引。"),
    ("■ 支柱二：逆变器业务现金奶牛 — 基本维持",
     "基本维持",
     "逆变器出货量143GW（-2.7%），国内同比下降19%，主要反映国内公用事业太阳能装机节奏放缓及竞争激烈，"
     "但逆变器收入仍增长8.1%至311.36亿元，表明海外市场ASP稳定及产品结构优化有效对冲了国内量的压力。"
     "阳光电源在全球地面电站逆变器市场份额约30%，且持续积累数字化运维服务等经常性收入来源，"
     "逆变器业务的稳定性并未因本次年报结果而受损。"),
    ("■ 支柱三：能源转型长期结构性需求 — 进一步强化",
     "进一步强化",
     "全球能源存储部署加速，2026年全球储能市场预计增速30%-50%，美国IRA政策驱动的需求、"
     "欧洲电网稳定性投资及亚太与中东市场的大型公用事业储能项目均提供多年期能见度。"
     "值得关注的是，AI数据中心（AIDC）储能备电需求正在成为新的增量市场，"
     "阳光电源已于2025年开始布局AIDC配套产品线（据华润证券2025年11月研报），"
     "这一趋势为公司储能业务提供了超越传统公用事业市场的期权价值。"),
]

for header_cn, status_cn, body_cn in thesis_cn:
    h_p = doc.add_paragraph()
    para_fmt(h_p, space_before=6, space_after=2)
    add_bold_heading(h_p, header_cn, size=10, color=(31, 78, 121))
    if "强化" in status_cn:
        sc = (55, 86, 35)
    elif "维持" in status_cn:
        sc = (31, 78, 121)
    else:
        sc = (192, 0, 0)
    s_p = doc.add_paragraph()
    para_fmt(s_p, left_indent=12, space_before=0, space_after=2)
    add_bold_heading(s_p, f"评估：{status_cn}", size=9.5, color=sc)
    b_p = doc.add_paragraph()
    para_fmt(b_p, left_indent=12, space_before=0, space_after=8)
    add_run(b_p, body_cn, size=9.5)

add_heading_para(doc, "3.1  风险更新", size=11, color=(46, 117, 182))
risks_cn = [
    ("▲ 新增/加剧风险：", "碳酸锂价格波动仍是近期毛利率的潜在风险，Q4已验证这一点。若2026年碳酸锂价格再度上涨，存量固定价格合同将再度承压。"),
    ("▲ 新增/加剧风险：", "美国关税与贸易政策风险有所升温。若对中国制造储能设备的关税进一步扩大，阳光电源的美国业务可能面临增量阻力。"),
    ("▼ 有所缓解：",      "国内逆变器价格竞争压力相对可控；2025年量虽下降19%，但ASP仍稳定，表明最严峻的价格压缩阶段可能已过。"),
    ("→ 基本维持：",      "来自宁德时代、比亚迪储能、华为数字能源的储能竞争持续，但阳光电源在海外市场的品牌认可度、安全记录和资质认证壁垒仍构成有效护城河。"),
]
for (label, risk_text) in risks_cn:
    r_p = doc.add_paragraph()
    para_fmt(r_p, left_indent=12, space_before=3, space_after=3)
    add_bold_heading(r_p, label, size=9.5, color=(31, 78, 121))
    add_run(r_p, risk_text, size=9.5)

add_chart(doc, "300274_chart7_beat_miss.png", width=6.2,
          caption="图7 — FY2025业绩与市场预期对比 | 资料来源：阳光电源2025年年度报告；Bloomberg/富途发布前一致预期")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 第8-10页：估值与盈利预测
# ─────────────────────────────────────────────────────────────────────────────
add_heading_para(doc, "第四部分：估值与盈利预测", size=12)
add_divider(doc)

add_heading_para(doc, "4.1  目标价更新", size=11, color=(46, 117, 182))
pt_cn = doc.add_paragraph()
para_fmt(pt_cn, space_before=4, space_after=6)
add_run(pt_cn,
        "我们将12个月目标价下调至158元（原175元），反映年报低于预期后的盈利预测下修。估值方法如下：",
        size=9.5)

pt_method_cn = [
    ("主要方法：市盈率估值（权重60%）。",
     "以20x目标P/E乘以我们修订后的2026年每股收益预测8.09元，得到内在价值161.8元。"
     "目标倍数低于历史均值约25x，折扣反映：①Q4 2025后近期利润率不确定性；"
     "②若碳酸锂价格持续高位或美国关税风险落地，存在进一步盈利下修可能；"
     "③宏观经济整体增速放缓背景下给予适当安全边际。"),
    ("辅助方法：EV/EBITDA估值（权重40%）。",
     "基于我们预测的2026年EBITDA约185亿元，以行业均值约13x EV/EBITDA估算公允价值约155元，"
     "与P/E估值结果基本吻合。加权平均目标价158元，较当前价格136.91元隐含上涨空间15.4%。"),
]
for (bold_text, body_text) in pt_method_cn:
    m_p = doc.add_paragraph()
    para_fmt(m_p, left_indent=12, space_before=3, space_after=6)
    add_bold_heading(m_p, bold_text, size=9.5, color=(31, 78, 121))
    add_run(m_p, body_text, size=9.5)

val_h_cn = ["估值方法", "权重", "盈利输入", "目标倍数", "隐含价值（元）"]
val_r_cn = [
    ["市盈率（NTM P/E）", "60%", "2026E EPS 8.09元",      "20.0x", "161.8元"],
    ["EV/EBITDA",         "40%", "2026E EBITDA 185亿元",  "13.0x", "153.5元"],
    ["加权目标价",        "100%", "—",                    "—",     "158.0元"],
    ["当前股价",          "—",   "2026年4月21日",          "—",     "136.91元"],
    ["隐含涨幅",          "—",   "—",                     "—",     "+15.4%"],
]
make_table(doc, val_h_cn, val_r_cn, col_widths=[1.8, 0.7, 2.0, 0.9, 1.6])
src_note(doc, "资料来源：分析师预测；Bloomberg；市场数据截至2026年4月21日")

add_chart(doc, "300274_chart10_valuation.png", width=6.2,
          caption="图10 — NTM市盈率估值区间 | 资料来源：Bloomberg；StockAnalysis.com；分析师预测")

add_heading_para(doc, "4.2  详细盈利预测更新", size=11, color=(46, 117, 182))

det_h_cn = ["利润表指标（亿元）", "FY2024A", "FY2025A", "FY2026E（新）", "FY2026E（前）", "变动（%）"]
det_r_cn = [
    ["营业收入",            "778.6",  "891.8",  "1020.0",  "1080.0",  "-5.6%"],
    ["收入增速（%）",       "+7.8%",  "+14.6%", "+14.4%",  "+21.1%",  "—"],
    ["毛利润",              "198.8",  "284.0E", "326.4",   "351.0",   "-7.1%"],
    ["毛利率（%）",         "25.5%",  "31.83%", "32.0%",   "32.5%",   "-50bps"],
    ["EBITDA",              "145",    "185E",   "205",     "220",     "-6.8%"],
    ["EBITDA利润率（%）",   "18.6%",  "20.7%E", "20.1%",   "20.4%",   "-30bps"],
    ["经营利润",            "135",    "170E",   "195",     "210",     "-7.1%"],
    ["经营利润率（%）",     "17.3%",  "19.1%E", "19.1%",   "19.4%",   "-30bps"],
    ["归母净利润",          "110.4",  "134.61", "167.7",   "180.0",   "-6.8%"],
    ["净利率（%）",         "14.2%",  "15.2%",  "16.4%",   "16.7%",   "-30bps"],
    ["基本EPS（元）",       "5.32",   "6.55",   "8.09",    "8.70",    "-7.0%"],
    ["P/E @ 136.91元",      "25.7x",  "20.9x",  "16.9x",   "15.7x",   "—"],
    ["EV/EBITDA（估算）",   "约19x",  "约15x",  "约14x",   "约13x",   "—"],
]
make_table(doc, det_h_cn, det_r_cn, col_widths=[2.0, 1.0, 1.0, 1.2, 1.2, 1.0])
src_note(doc, "注：A=实际值，E=预测值；2025年EBITDA及经营利润为分析师测算（年报中未单独列示）。资料来源：阳光电源2025年年度报告；分析师预测")

add_chart(doc, "300274_chart9_estimates.png", width=6.2,
          caption="图9 — 盈利预测修正对比 | 资料来源：阳光电源2025年年度报告；分析师预测")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# 资料来源
# ─────────────────────────────────────────────────────────────────────────────
add_heading_para(doc, "资料来源与参考文献", size=12)
add_divider(doc)

src_intro_cn = doc.add_paragraph()
para_fmt(src_intro_cn, space_before=4, space_after=6)
add_run(src_intro_cn, "本报告所有财务数据均来自公司官方公告及知名金融数据提供商，以下为各核心来源的可点击链接：", size=9.5)

sources_cn = [
    ("公司官方文件：", [
        ("阳光电源2025年年度报告摘要（CNINFO）", "https://static.cninfo.com.cn/finalpage/2026-04-01/1225066677.PDF",
         "2026年4月1日发布，所有FY2025财务数据的主要来源"),
        ("投资者电话会议纪要（新浪财经）", "https://finance.sina.com.cn/roll/2026-04-01/doc-inhsziht8333473.shtml",
         "2026年3月31日，Q4毛利率解释、2026年业绩指引"),
        ("阳光电源2025年年报分析（东方财富财富号）", "https://caifuhao.eastmoney.com/news/20260402060936774616180",
         "2026年4月2日，业务板块详细分析与盈利拆分"),
    ]),
    ("财务数据来源：", [
        ("StockAnalysis.com — 季度营收历史", "https://stockanalysis.com/quote/she/300274/revenue/",
         "Q1 2024至Q4 2025季度营收数据"),
        ("StockAnalysis.com — 概况与估值", "https://stockanalysis.com/quote/she/300274/",
         "市值、市盈率、52周区间、流通股数"),
        ("富途牛牛 — 阳光电源2025年报储能评论", "https://news.futunn.com/en/post/71013862/sungrow-power-300274-2025-annual-report-commentary-energy-storage-systems",
         "储能板块评论及市场一致预期"),
        ("华尔街见闻 — 阳光电源2025年年报", "https://wallstreetcn.com/articles/3768898",
         "营收近900亿，净利135亿，储能构网技术分析"),
    ]),
    ("卖方研究及市场评论：", [
        ("华润证券 — Q3 2025季报点评", "https://finance.sina.com.cn/roll/2025-11-10/doc-infwxhkx3003604.shtml",
         "2025年11月，AIDC产品布局，盈利能力趋势"),
        ("新浪财经 — Q4归母净利润骤降分析", "https://finance.sina.com.cn/roll/2026-04-01/doc-inhsziht8333473.shtml",
         "2026年4月，Q4 2025季度归母净利润深度解读"),
    ]),
]

for section_title_cn, items_cn in sources_cn:
    s_p = doc.add_paragraph()
    para_fmt(s_p, space_before=8, space_after=3)
    add_bold_heading(s_p, section_title_cn, size=10, color=(31, 78, 121))
    for title_cn, url, note_cn in items_cn:
        item_p = doc.add_paragraph()
        para_fmt(item_p, left_indent=18, space_before=3, space_after=3)
        add_run(item_p, "• ", size=9.5)
        add_hyperlink(item_p, title_cn, url)
        add_run(item_p, f"  —  {note_cn}", size=9, italic=True, color=(100, 100, 100))

doc.add_paragraph()
add_divider(doc)

disc_cn = doc.add_paragraph()
para_fmt(disc_cn, space_before=6, space_after=4)
add_run(disc_cn,
        "重要声明：本报告仅供参考，不构成投资建议。所有财务预测均为分析师测算，存在不确定性。"
        "超预期/低于预期分析基于发布前市场一致预期（来源：Bloomberg、富途）。"
        "市场数据截至2026年4月21日。阳光电源股份有限公司（300274.SZ）在深圳证券交易所上市。"
        "历史业绩不代表未来表现。",
        size=8, italic=True, color=(128, 128, 128))

doc.save(OUT)
print(f"中文报告已保存：{OUT}")
