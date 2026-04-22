#!/usr/bin/env python3
"""
阳光电源 (Sungrow Power Supply, 300274.SZ)
Q4 2025 / FY2025 Earnings Update — English DOCX Report
"""

import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE   = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/300274"
CHARTS = BASE
OUT    = os.path.join(BASE, "300274_Q4_FY2025_Earnings_Update.docx")

# ── helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    part   = paragraph.part
    r_id   = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hl     = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    r      = OxmlElement("w:r")
    rPr    = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hl.append(r)
    paragraph._p.append(hl)
    return hl

def para_fmt(p, left_indent=0, space_before=0, space_after=6, line_spacing=None):
    pf = p.paragraph_format
    pf.left_indent  = Pt(left_indent)
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)

def add_run(para, text, bold=False, italic=False, size=10, color=None, font="Times New Roman"):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def add_heading(doc, text, level=1, size=13, color=(31, 78, 121), bold=True):
    p = doc.add_paragraph()
    para_fmt(p, space_before=8, space_after=4)
    run = add_run(p, text, bold=bold, size=size, color=color)
    return p

def add_divider(doc, color="1F4E79"):
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pb  = p._p
    pPr = pb.get_or_add_pPr()
    pBd = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"),  "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBd.append(bot)
    pPr.append(pBd)
    return p

def add_chart(doc, filename, width=6.0, caption=""):
    path = os.path.join(CHARTS, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        para_fmt(p, space_before=4, space_after=2)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=Inches(width))
    else:
        p = doc.add_paragraph(f"[Chart not found: {filename}]")
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para_fmt(cp, space_before=0, space_after=6)
        add_run(cp, "", size=8, italic=True, color=(128, 128, 128))
    return p

def make_table(doc, headers, rows, col_widths=None, header_bg="1F4E79", alt_bg="E9EEF4"):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = tbl.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], header_bg)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(8.5)
            run.font.name = "Times New Roman"
    # Data rows
    for ri, row in enumerate(rows):
        cells = tbl.rows[ri+1].cells
        bg = alt_bg if ri % 2 == 1 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            cells[ci].text = str(cell_text)
            if bg != "FFFFFF":
                set_cell_bg(cells[ci], bg)
            p = cells[ci].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(8.5)
                run.font.name = "Times New Roman"
    # Column widths
    if col_widths:
        for ri2, row2 in enumerate(tbl.rows):
            for ci2, (cell2, w) in enumerate(zip(row2.cells, col_widths)):
                cell2.width = Inches(w)
    return tbl

# ═════════════════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ═════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.page_width    = Cm(21.59)
    section.page_height   = Cm(27.94)

# ─────────────────────────────────────────────────────────────────────────────
# PAGE 1: EARNINGS SUMMARY
# ─────────────────────────────────────────────────────────────────────────────

# Title block
title_p = doc.add_paragraph()
para_fmt(title_p, space_before=4, space_after=2)
add_run(title_p, "SUNGROW POWER SUPPLY CO., LTD. (阳光电源)", bold=True, size=14, color=(31, 78, 121))

sub_p = doc.add_paragraph()
para_fmt(sub_p, space_before=0, space_after=2)
add_run(sub_p, "300274.SZ  |  Q4 2025 / FY2025 EARNINGS UPDATE", bold=True, size=11, color=(46, 117, 182))

meta_p = doc.add_paragraph()
para_fmt(meta_p, space_before=0, space_after=8)
add_run(meta_p, "April 22, 2026  ·  Equity Research  ·  China Clean Energy / New Energy Equipment",
        size=9, italic=True, color=(128, 128, 128))

add_divider(doc)

# Rating / PT box (simulated as a table)
rtbl = doc.add_table(rows=1, cols=4)
rtbl.alignment = WD_TABLE_ALIGNMENT.CENTER
rtbl.style = "Table Grid"
rt_cells = rtbl.rows[0].cells
rt_data = [
    ("Rating", "OUTPERFORM ★", "1F4E79"),
    ("Price (Apr 21, 2026)", "CNY 136.91", "2E75B6"),
    ("Price Target", "CNY 158.00", "375623"),
    ("Implied Upside", "+15.4%", "375623"),
]
for cell, (label, val, color_hex) in zip(rt_cells, rt_data):
    set_cell_bg(cell, color_hex)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "\n")
    r1.bold = False; r1.font.size = Pt(8); r1.font.color.rgb = RGBColor(200, 220, 240); r1.font.name = "Times New Roman"
    r2 = p.add_run(val)
    r2.bold = True; r2.font.size = Pt(10); r2.font.color.rgb = RGBColor(255, 255, 255); r2.font.name = "Times New Roman"

doc.add_paragraph()

# ── EARNINGS SNAPSHOT TABLE ───────────────────────────────────────────────────
snap_p = doc.add_paragraph()
para_fmt(snap_p, space_before=4, space_after=4)
add_run(snap_p, "Q4 2025 / FY2025 EARNINGS SNAPSHOT", bold=True, size=10.5, color=(31, 78, 121))

snap_headers = ["Metric", "FY2025 Actual", "FY2025 Consensus", "Beat / (Miss)", "YoY Change"]
snap_rows = [
    ["Revenue (CNY B)",      "¥89.18B",  "¥90.0B",  "(¥0.82B / -0.9%)",  "+14.6%"],
    ["Net Profit Attr. (CNY B)", "¥13.46B", "¥15.45B", "(¥1.99B / -12.9%)", "+22.0%"],
    ["EPS (CNY/share)",      "¥6.55",    "¥7.50E",  "(¥0.95 / -12.7%)",  "+23.1%"],
    ["Gross Margin (%)",     "31.83%",   "~30.0%",  "+183 bps",           "+631 bps YoY"],
    ["Net Margin (%)",       "15.17%",   "~17.0%",  "-183 bps",           "+488 bps YoY"],
    ["Q4 Revenue (CNY B)",   "¥22.78B",  "~¥24.0B", "(¥1.22B / -5.1%)",  "-18.4% YoY"],
    ["Q4 Net Profit (CNY B)","¥1.58B",   "~¥3.5B",  "(¥1.92B / -54.9%)", "-54.0% YoY"],
]
make_table(doc, snap_headers, snap_rows, col_widths=[2.0, 1.3, 1.3, 1.6, 1.3])
src_p = doc.add_paragraph("Source: Sungrow 2025 Annual Report (2026-04-01); Consensus estimates from Bloomberg/Futunn (pre-announcement)")
para_fmt(src_p, space_before=1, space_after=8)
add_run(src_p, "", size=7.5, italic=True, color=(128, 128, 128))
for run in src_p.runs:
    run.font.size = Pt(7.5)
    run.italic = True
    run.font.color.rgb = RGBColor(128, 128, 128)

doc.add_paragraph()

# ── KEY TAKEAWAYS ─────────────────────────────────────────────────────────────
kt_p = doc.add_paragraph()
para_fmt(kt_p, space_before=4, space_after=4)
add_run(kt_p, "KEY TAKEAWAYS", bold=True, size=10.5, color=(31, 78, 121))

bullets = [
    ("■ MISS — FY2025 net profit missed consensus by 12.9%; Q4 earnings collapsed -54% YoY on margin compression",
     "FY2025 revenue of ¥89.18B came in broadly in line with consensus (miss of only ¥0.82B, -0.9%), but net profit attributable to shareholders of ¥13.46B was materially below the ¥15.45B consensus — a ¥1.99B shortfall driven entirely by Q4's dramatic margin implosion. Q4 net profit fell -54% YoY and -62% QoQ to just ¥1.58B as storage gross margins compressed ~17 percentage points QoQ (to ~24%) due to: (1) revenue mix shift toward lower-margin new energy investment & development projects concentrated in Q4; (2) elevated lithium carbonate prices that could not be passed through on legacy fixed-price contracts; and (3) the unwinding of Q3's unusually high-margin overseas storage project base."),
    ("■ FY2025 full-year results solid — revenue +15%, net profit +22%, gross margins at multi-year highs",
     "Full-year 2025 results were respectable: revenue grew 14.6% YoY to ¥89.18B, net profit grew 22.0% to ¥13.46B, and gross margin expanded significantly to 31.83% (+631bps YoY), reflecting a favorable product mix driven by high-margin overseas energy storage shipments. Operating cash flow per share grew 40.2% YoY to ¥8.16, demonstrating strong cash generation. ROIC improved to 25.01%, highlighting capital efficiency."),
    ("■ Energy storage remains the engine — 43 GWh shipped (+54% YoY), overseas +89% YoY; 2026 target: >60 GWh",
     "Energy storage system shipments surged 53.6% YoY to 43 GWh in FY2025, driven by explosive overseas growth (+89% YoY to 36 GWh). Domestic storage declined 22% YoY (7 GWh), but this was offset by international demand diversification. The 43 GWh positions Sungrow as a global leader. For 2026, management guided energy storage shipments of >60 GWh (+40% YoY), supported by continued global market growth expected at 30-50%."),
    ("■ Maintaining OUTPERFORM with revised CNY 158 price target (from CNY 175); Q4 weakness does not impair long-term thesis",
     "While Q4's margin disappointment is sharp, we view it as a one-time mix and raw material timing issue rather than a structural deterioration. The Q4 sequential margin decline is unlikely to persist given: (i) overseas storage demand trajectory is intact; (ii) lithium prices have since stabilized; (iii) Sungrow's dominant market position in global storage. We reduce our FY2026E net profit to ¥16.8B (from ¥18.0B) but retain our OUTPERFORM rating. Our revised ¥158 price target implies 20x FY2026E P/E, a slight discount to the historical average of ~25x to account for near-term margin uncertainty."),
]

for i, (header, body) in enumerate(bullets):
    b_p = doc.add_paragraph()
    para_fmt(b_p, space_before=4, space_after=2)
    add_run(b_p, header, bold=True, size=9.5, color=(31, 78, 121))
    body_p = doc.add_paragraph()
    para_fmt(body_p, left_indent=12, space_before=0, space_after=8)
    add_run(body_p, body, size=9.5)

# ── UPDATED ESTIMATES TABLE ───────────────────────────────────────────────────
add_divider(doc)
ue_p = doc.add_paragraph()
para_fmt(ue_p, space_before=6, space_after=4)
add_run(ue_p, "UPDATED FINANCIAL ESTIMATES", bold=True, size=10.5, color=(31, 78, 121))

ue_headers = ["Metric", "FY2024A", "FY2025A", "FY2025E (Prior)", "FY2026E (New)", "FY2026E (Prior)"]
ue_rows = [
    ["Revenue (CNY B)",          "¥77.86",  "¥89.18",  "¥90.0",   "¥102.0",  "¥108.0"],
    ["Revenue Growth (%)",       "+7.8%",   "+14.6%",  "+15.6%",  "+14.4%",  "+20.0%"],
    ["Gross Margin (%)",         "25.5%",   "31.83%",  "~30.0%",  "32.0%",   "32.5%"],
    ["Net Profit Attr. (CNY B)", "¥11.04",  "¥13.46",  "¥15.45",  "¥16.77",  "¥18.0"],
    ["Net Profit Growth (%)",    "+16.9%",  "+22.0%",  "+40.0%",  "+24.6%",  "+33.6%"],
    ["EPS (CNY)",                "¥5.32",   "¥6.55",   "¥7.50",   "¥8.09",   "¥8.70"],
    ["P/E (x) @ CNY 136.91",    "25.7x",   "20.9x",   "—",       "16.9x",   "—"],
]
make_table(doc, ue_headers, ue_rows, col_widths=[2.0, 1.1, 1.1, 1.3, 1.3, 1.3])
src2_p = doc.add_paragraph("Note: Prior estimates = internal estimates before FY2025 annual report. A = Actual, E = Estimate.")
para_fmt(src2_p, space_before=1, space_after=2)
for run in src2_p.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)
src3_p = doc.add_paragraph("Source: Sungrow 2025 Annual Report; Analyst estimates")
para_fmt(src3_p, space_before=0, space_after=8)
for run in src3_p.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGES 2-3: DETAILED RESULTS ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "SECTION 1: DETAILED RESULTS ANALYSIS", size=12)
add_divider(doc)

# Chart 1 — Revenue
add_chart(doc, "300274_chart1_revenue.png", width=6.2,
          caption="Figure 1 — Quarterly Revenue Progression | Source: Sungrow 2025 Annual Report; StockAnalysis.com")

# Revenue narrative
add_heading(doc, "1.1  Revenue Performance", size=11, color=(46, 117, 182))
rev_text = [
    ("FY2025 Revenue: ¥89.18B (+14.6% YoY) — In Line With Consensus. ",
     "Full-year 2025 revenue reached ¥89.18 billion, growing 14.6% year-over-year from ¥77.86B in FY2024. This growth, while solid in absolute terms, marks a meaningful deceleration from the exceptional 79.5% growth seen in FY2023. Consensus had anticipated ¥90.0B, resulting in a marginal miss of ¥0.82B (-0.9%). The revenue mix shifted materially toward energy storage (now 41.8% of total vs. ~32% in FY2024), while PV inverter share declined."),
    ("Q4 2025 Revenue: ¥22.78B (-18.4% YoY, -0.4% QoQ) — Significant YoY Shortfall. ",
     "Q4 2025 revenue of ¥22.78B was substantially below the prior year's Q4 2024 of ¥27.91B — a decline of 18.4% YoY. This was the weakest quarterly revenue since Q1 2025 (¥19.04B). The YoY decline reflects a very tough Q4 2024 base (which itself had grown 8.0% YoY), as well as weaker domestic inverter demand (-19% YoY shipments for FY2025 domestic). Sequentially, revenue was essentially flat (-0.4% QoQ vs. Q3's ¥22.87B), suggesting Q4 weakness was primarily YoY normalization rather than a sequential deterioration in demand."),
]
for (bold_text, body_text) in rev_text:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_run(p, bold_text, bold=True, size=9.5)
    add_run(p, body_text, size=9.5)

# Revenue breakdown table
rb_h = ["Segment", "FY2024A (CNY B)", "FY2024A (%)", "FY2025A (CNY B)", "FY2025A (%)", "YoY Growth"]
rb_r = [
    ["Energy Storage Systems", "¥24.9",  "32.0%", "¥37.29", "41.8%", "+49.7%"],
    ["PV Inverters & Power Electronics", "¥28.8", "37.0%", "¥31.14", "34.9%", "+8.1%"],
    ["New Energy Investment & Dev.", "¥21.0", "27.0%", "¥16.56", "18.6%", "-21.1%"],
    ["Other", "¥3.2", "4.0%", "¥4.20", "4.7%", "+31.3%"],
    ["Total Revenue", "¥77.86", "100%", "¥89.18", "100%", "+14.6%"],
]
make_table(doc, rb_h, rb_r, col_widths=[2.3, 1.3, 1.0, 1.3, 1.0, 1.0])
src4 = doc.add_paragraph("Source: Sungrow 2025 Annual Report (2026-04-01); FY2024 segment estimates from company disclosures")
para_fmt(src4, space_before=1, space_after=8)
for run in src4.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

add_chart(doc, "300274_chart4_segments.png", width=6.2,
          caption="Figure 4 — Revenue by Segment FY2024A vs. FY2025A | Source: Sungrow 2025 Annual Report")

# Chart 2 — Net profit
add_chart(doc, "300274_chart2_net_profit.png", width=6.2,
          caption="Figure 2 — Quarterly Net Profit | Source: Sungrow 2025 Annual Report; Analyst quarterly estimates")

# Profitability section
add_heading(doc, "1.2  Profitability Analysis", size=11, color=(46, 117, 182))
prof_text = [
    ("FY2025 Gross Margin: 31.83% (+631bps YoY) — Significant Structural Improvement. ",
     "The full-year gross margin of 31.83% represents a 6.31 percentage point improvement versus the 25.52% reported in FY2024. This expansion was driven by: (1) the growing share of high-margin overseas energy storage projects, where Sungrow commands premium pricing; (2) a favorable product mix shift away from commoditized domestic inverters toward value-added storage systems; and (3) operational leverage as fixed costs are spread over a larger revenue base. FY2025 full-year results validate Sungrow's strategy of international expansion."),
    ("Q4 2025 Net Profit Miss: ¥1.58B (-54% YoY) — Three Converging Headwinds. ",
     "Q4 2025 net profit of ¥1.58B was dramatically below Q4 2024's ¥3.44B, resulting in a -54% YoY decline and -62% sequential drop. Management attributed this to three concurrent factors: (1) Revenue mix degradation: New energy investment & development (NEI) projects — which have lower margins than storage products — were disproportionately weighted in Q4, pulling blended margins lower; (2) Storage margin compression: Energy storage gross margin fell approximately 17 percentage points QoQ to ~24% in Q4 2025, compared to an elevated ~41% in Q3 2025 when high-margin overseas contracts were being recognized; (3) Lithium carbonate price spike: Spot lithium carbonate prices rose in Q4 2025, squeezing margins on legacy fixed-price contracts where cost increases could not be immediately passed through."),
]
for (bold_text, body_text) in prof_text:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_run(p, bold_text, bold=True, size=9.5)
    add_run(p, body_text, size=9.5)

# Margin table
mg_h = ["Metric", "Q1'25A", "Q2'25A", "Q3'25A", "Q4'25A", "FY2025A", "FY2024A", "YoY Chg"]
mg_r = [
    ["Gross Margin (%)",     "29.4%", "32.1%", "35.9%", "24.2%", "31.83%", "25.52%", "+631 bps"],
    ["Net Margin (%)",       "17.9%", "17.7%", "18.1%",  "6.9%", "15.17%", "10.29%", "+488 bps"],
    ["ROIC (%)",             "—",     "—",     "—",      "—",    "25.01%", "~20.0%", "+~500 bps"],
    ["Op. CF/Share (CNY)",   "—",     "—",     "—",      "—",    "¥8.16",  "¥5.82",  "+40.2%"],
]
make_table(doc, mg_h, mg_r, col_widths=[1.8, 0.8, 0.8, 0.8, 0.8, 0.9, 0.9, 0.9])
src5 = doc.add_paragraph("Source: Sungrow 2025 Annual Report; Company investor call (2026-03-31)")
para_fmt(src5, space_before=1, space_after=4)
for run in src5.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

add_chart(doc, "300274_chart3_margins.png", width=6.2,
          caption="Figure 3 — Quarterly Margin Trends | Source: Sungrow 2025 Annual Report; Company investor call")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGES 4-5: KEY METRICS & GUIDANCE
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "SECTION 2: KEY OPERATING METRICS & GUIDANCE", size=12)
add_divider(doc)

add_heading(doc, "2.1  Shipment Performance", size=11, color=(46, 117, 182))

ship_text = [
    ("Energy Storage: 43 GWh (+53.6% YoY) — Global Leadership Solidified. ",
     "FY2025 energy storage system shipments of 43 GWh exceeded our estimate of ~40 GWh and grew 53.6% from 28 GWh in FY2024. Overseas shipments surged 89% YoY to 36 GWh (84% of total), reflecting successful global market penetration in the US, Europe, and APAC. Domestic storage contracted 22% YoY to 7 GWh, consistent with softer Chinese utility storage project timing in FY2025. The divergence between domestic and overseas trends underscores Sungrow's strategic pivot toward higher-margin international markets."),
    ("PV Inverters: 143 GW (-2.7% YoY) — Domestic Headwind Dominates. ",
     "PV inverter shipments of 143 GW in FY2025 declined slightly (-2.7%) from 147 GW in FY2024, with domestic shipments falling 19% YoY to 57 GW — a reflection of the more competitive and price-pressured domestic Chinese solar market. However, this was partially offset by international inverter demand. Inverter revenue still grew 8.1% YoY to ¥31.14B as mix improved toward higher-value products (string inverters, utility-scale solutions) and international pricing held up better."),
]
for (bold_text, body_text) in ship_text:
    p = doc.add_paragraph()
    para_fmt(p, space_before=4, space_after=6)
    add_run(p, bold_text, bold=True, size=9.5)
    add_run(p, body_text, size=9.5)

# Shipments table
ship_h = ["Metric", "FY2023A", "FY2024A", "FY2025A", "YoY Growth", "2026 Target"]
ship_r = [
    ["Energy Storage (GWh)",         "16.0",  "28.0",  "43.0",  "+53.6%", ">60.0 GWh"],
    ["— Domestic (GWh)",             "7.0",   "9.0",   "7.0",   "-22.2%",  "~10.0"],
    ["— Overseas (GWh)",             "9.0",   "19.0",  "36.0",  "+89.5%",  ">50.0"],
    ["PV Inverters (GW)",            "130.0", "147.0", "143.0", "-2.7%",   "~155.0 GW"],
    ["— Domestic (GW)",              "n/a",   "71.0",  "57.0",  "-19.7%",  "~60.0"],
    ["— Overseas (GW)",              "n/a",   "76.0",  "86.0",  "+13.2%",  ">90.0"],
]
make_table(doc, ship_h, ship_r, col_widths=[2.2, 1.0, 1.0, 1.0, 1.0, 1.2])
src6 = doc.add_paragraph("Source: Sungrow 2025 Annual Report (2026-04-01); 2026 targets from management guidance (2026-03-31 investor call); FY2023 estimates from prior reports")
para_fmt(src6, space_before=1, space_after=6)
for run in src6.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

add_chart(doc, "300274_chart5_storage_shipments.png", width=6.0,
          caption="Figure 5 — Energy Storage Shipments (GWh) | Source: Sungrow 2025 Annual Report; Management guidance")
add_chart(doc, "300274_chart6_inverter_shipments.png", width=5.5,
          caption="Figure 6 — PV Inverter Shipments (GW) | Source: Sungrow 2025 Annual Report")

# Guidance section
add_heading(doc, "2.2  Management Guidance & Outlook", size=11, color=(46, 117, 182))

guid_p = doc.add_paragraph()
para_fmt(guid_p, space_before=4, space_after=4)
add_run(guid_p, "Management Guidance Summary (Investor Call, March 31, 2026)", bold=True, size=9.5, color=(31, 78, 121))

guid_h = ["Metric", "FY2025A", "FY2026 Guidance / Target", "Our Estimate (FY2026E)", "Assessment"]
guid_r = [
    ["Storage Shipments (GWh)", "43 GWh", ">60 GWh (+40% YoY)", "62 GWh", "Credible; market tailwinds strong"],
    ["Global Storage Mkt Growth", "—", "30%–50% YoY", "~40% YoY", "Achievable; track record supportive"],
    ["Revenue (CNY B)", "¥89.18B", "No formal guidance", "¥102.0B", "+14.4% YoY est."],
    ["Net Profit (CNY B)", "¥13.46B", "No formal guidance", "¥16.77B", "+24.6% YoY est."],
    ["Storage Gross Margin", "31.83% blended", "Recovery expected H1'26", ">30% full-year", "Supportive if Li prices stable"],
]
make_table(doc, guid_h, guid_r, col_widths=[1.8, 1.2, 1.8, 1.5, 1.5])
src7 = doc.add_paragraph("Source: Sungrow investor call (March 31, 2026); Analyst estimates")
para_fmt(src7, space_before=1, space_after=4)
for run in src7.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

guid_commentary = doc.add_paragraph()
para_fmt(guid_commentary, space_before=6, space_after=6)
add_run(guid_commentary,
        "Management Commentary on Q4 Weakness: ", bold=True, size=9.5)
add_run(guid_commentary,
        "On the March 31, 2026 investor call, management attributed Q4 2025's margin compression explicitly to "
        "three transient factors: (1) an elevated concentration of new energy investment & development projects "
        "in Q4 that carried structurally lower margins than energy storage hardware; (2) the reversal of Q3's "
        "unusually high gross margins which benefited from the recognition of several high-value overseas storage "
        "contracts; and (3) a temporary increase in lithium carbonate procurement costs in Q4 that could not be "
        "offset through contract price adjustments on near-term deliveries. Management expressed confidence that "
        "these factors are non-recurring and indicated storage gross margins should recover in H1 2026 as "
        "lithium prices have subsequently stabilized and new overseas project bookings carry improved pricing terms.",
        size=9.5)

add_chart(doc, "300274_chart8_storage_margins.png", width=6.0,
          caption="Figure 8 — Energy Storage Gross Margin by Quarter | Source: Sungrow 2025 Annual Report; Investor call")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGES 6-7: UPDATED INVESTMENT THESIS
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "SECTION 3: UPDATED INVESTMENT THESIS", size=12)
add_divider(doc)

thesis_p = doc.add_paragraph()
para_fmt(thesis_p, space_before=4, space_after=6)
add_run(thesis_p,
        "Our investment thesis on Sungrow rests on three pillars: (1) global energy storage dominance with superior "
        "international expansion; (2) inverter market leadership providing durable cash flow; and (3) long-term "
        "structural demand driven by global energy transition. We assess each pillar's status in light of Q4/FY2025 results.",
        size=9.5)

thesis_bullets = [
    ("■ PILLAR 1: Global Energy Storage Leadership — STRENGTHENED",
     "STRENGTHENED",
     "FY2025 storage shipments of 43 GWh, up 53.6% YoY with overseas volumes surging 89%, reaffirm Sungrow's position "
     "as the world's #1 energy storage system supplier. The international market pivot is delivering — Sungrow now derives "
     "84% of storage shipments from overseas, commanding premium pricing compared to domestic peers constrained by "
     "commoditized pricing in China. The 2026 guidance of >60 GWh (+40% YoY) is backed by a strong global project pipeline "
     "including the US, Europe, Australia, and the Middle East. Q4's margin compression was a timing issue, NOT a structural "
     "erosion of Sungrow's competitive position in storage. If anything, the rapid overseas mix shift over FY2025 "
     "(domestic storage -22%, overseas +89%) demonstrates that Sungrow is executing its internationalization strategy ahead "
     "of plan."),
    ("■ PILLAR 2: PV Inverter Cash Cow — UNCHANGED",
     "UNCHANGED",
     "PV inverter shipments of 143 GW in FY2025 (-2.7% YoY) reflect domestic market headwinds — specifically, China's "
     "curtailed utility solar installations and intense price competition from domestic rivals in the residential segment. "
     "Inverter revenue still grew 8.1% YoY to ¥31.14B due to favorable international pricing and product mix improvement "
     "toward higher-value string inverters and utility-scale solutions. The inverter segment remains the company's "
     "highest-volume, highest-market-share business (Sungrow holds ~30% global share in utility-scale inverters) and "
     "generates stable, recurring maintenance and digital service revenue. We do not see inverter competitive dynamics "
     "deteriorating structurally in the FY2025 results."),
    ("■ PILLAR 3: Structural Demand from Energy Transition — STRENGTHENED",
     "STRENGTHENED",
     "Global energy storage deployment continues to accelerate, with the market expected to grow 30-50% in 2026 according "
     "to management's commentary. US IRA-driven demand, European grid stability investments, and utility-scale storage "
     "buildouts in APAC and the Middle East all provide multi-year visibility. Sungrow's 2025 annual report highlights "
     "that global energy storage installations reached a record level, and the company expects this trajectory to continue. "
     "Separately, the AIDC (AI Data Center) market is an emerging demand driver for energy storage backup systems — "
     "Sungrow disclosed it is actively building AIDC-compatible storage product lines, as highlighted in the Q3 analyst "
     "report from Huarun Securities (November 2025). This secular trend adds optionality beyond utility storage."),
]

for header, status, body in thesis_bullets:
    h_p = doc.add_paragraph()
    para_fmt(h_p, space_before=6, space_after=2)
    add_run(h_p, header, bold=True, size=10, color=(31, 78, 121))
    status_color = (55, 86, 35) if "STRENGTHEN" in status else (31, 78, 121) if "UNCHANGED" in status else (192, 0, 0)
    s_p = doc.add_paragraph()
    para_fmt(s_p, left_indent=12, space_before=0, space_after=2)
    add_run(s_p, f"Status: {status}", bold=True, size=9.5, color=status_color)
    b_p = doc.add_paragraph()
    para_fmt(b_p, left_indent=12, space_before=0, space_after=8)
    add_run(b_p, body, size=9.5)

# Risks
add_heading(doc, "3.1  Risks Update", size=11, color=(46, 117, 182))
risks = [
    ("▲ New / Elevated: ",  "Lithium carbonate price volatility remains a near-term margin risk, as demonstrated by Q4 2025. If Li prices re-accelerate in 2026, fixed-price legacy contracts could compress margins again. Monitoring closely."),
    ("▲ New / Elevated: ", "US tariff and trade policy risk has increased given ongoing US-China trade tensions. Sungrow's US operations may face incremental headwinds if tariffs on Chinese-manufactured energy storage equipment are broadened in 2026."),
    ("▼ Mitigated: ",       "Domestic PV inverter pricing risk has stabilized; while volume was weak (-19% domestic), ASP held firm in 2025 due to product mix improvement, suggesting the worst of domestic pricing compression may be behind us."),
    ("→ Unchanged: ",       "Competition from CATL, BYD Energy Storage, and Huawei Digital in energy storage remains intense. However, Sungrow's bankability, safety track record, and overseas certifications maintain meaningful competitive moats internationally."),
]
for (label, risk_text) in risks:
    r_p = doc.add_paragraph()
    para_fmt(r_p, left_indent=12, space_before=3, space_after=3)
    add_run(r_p, label, bold=True, size=9.5)
    add_run(r_p, risk_text, size=9.5)

add_chart(doc, "300274_chart7_beat_miss.png", width=6.2,
          caption="Figure 7 — FY2025 Beat/Miss vs. Consensus | Source: Sungrow 2025 Annual Report; Bloomberg pre-earnings consensus")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# PAGES 8-10: VALUATION & UPDATED ESTIMATES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "SECTION 4: VALUATION & UPDATED ESTIMATES", size=12)
add_divider(doc)

add_heading(doc, "4.1  Price Target Update", size=11, color=(46, 117, 182))
pt_p = doc.add_paragraph()
para_fmt(pt_p, space_before=4, space_after=6)
add_run(pt_p,
        "We revise our 12-month price target to CNY 158.00 (from CNY 175.00), reflecting downward estimate revisions "
        "after the FY2025 results miss. Our methodology: ",
        size=9.5)

pt_method = [
    ("Primary: P/E Valuation (60% weight). ",
     "We apply a target P/E of 20x on our revised FY2026E EPS of ¥8.09, yielding an intrinsic value of ¥161.8. "
     "This target multiple represents a discount to the stock's 3-year historical average P/E of ~25x to reflect: "
     "(i) near-term margin uncertainty following Q4 2025; (ii) the risk of further earnings estimate cuts if "
     "lithium prices remain elevated or US tariff risk materializes; (iii) the global growth slowdown macro backdrop."),
    ("Secondary: EV/EBITDA (40% weight). ",
     "On an EV/EBITDA basis, Sungrow trades at approximately 12.5x FY2026E EBITDA (estimated ¥18.5B). "
     "Applying a peer-inline multiple of 13x yields a fair value of ~¥155, consistent with our P/E-based estimate. "
     "Our blended price target of ¥158 implies 15.4% upside from the current price of ¥136.91."),
]
for (bold_text, body_text) in pt_method:
    m_p = doc.add_paragraph()
    para_fmt(m_p, left_indent=12, space_before=3, space_after=6)
    add_run(m_p, bold_text, bold=True, size=9.5)
    add_run(m_p, body_text, size=9.5)

# Valuation summary table
val_h = ["Methodology", "Weight", "EPS / EBITDA Input", "Multiple", "Implied Value (CNY)"]
val_r = [
    ["P/E (NTM)",         "60%", "FY2026E EPS ¥8.09",     "20.0x",  "¥161.8"],
    ["EV/EBITDA",         "40%", "FY2026E EBITDA ¥18.5B", "13.0x",  "¥153.5"],
    ["Blended PT",        "100%", "—",                    "—",       "¥158.0"],
    ["Current Price",     "—",   "April 21, 2026",        "—",       "¥136.91"],
    ["Implied Upside",    "—",   "—",                     "—",       "+15.4%"],
]
make_table(doc, val_h, val_r, col_widths=[1.8, 0.7, 2.0, 0.9, 1.6])
src8 = doc.add_paragraph("Source: Analyst estimates; Bloomberg; Market data as of April 21, 2026")
para_fmt(src8, space_before=1, space_after=6)
for run in src8.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

add_chart(doc, "300274_chart10_valuation.png", width=6.2,
          caption="Figure 10 — NTM P/E Valuation Band | Source: Bloomberg; StockAnalysis.com; Analyst estimates")

# Updated detailed estimates
add_heading(doc, "4.2  Detailed Updated Estimates", size=11, color=(46, 117, 182))

det_h = ["Income Statement (CNY B)", "FY2024A", "FY2025A", "FY2026E (New)", "FY2026E (Old)", "Change (%)"]
det_r = [
    ["Revenue",                  "¥77.86",  "¥89.18",  "¥102.0",  "¥108.0",  "-5.6%"],
    ["Revenue Growth (%)",       "+7.8%",   "+14.6%",  "+14.4%",  "+21.1%",  "—"],
    ["Gross Profit",             "¥19.9",   "¥28.4",   "¥32.6",   "¥35.1",   "-7.1%"],
    ["Gross Margin (%)",         "25.5%",   "31.83%",  "32.0%",   "32.5%",   "-50 bps"],
    ["EBITDA",                   "¥14.5",   "¥18.5E",  "¥20.5",   "¥22.0",   "-6.8%"],
    ["EBITDA Margin (%)",        "18.6%",   "20.7%E",  "20.1%",   "20.4%",   "-30 bps"],
    ["Operating Profit",         "¥13.5",   "¥17.0E",  "¥19.5",   "¥21.0",   "-7.1%"],
    ["Operating Margin (%)",     "17.3%",   "19.1%E",  "19.1%",   "19.4%",   "-30 bps"],
    ["Net Profit Attr.",         "¥11.04",  "¥13.46",  "¥16.77",  "¥18.00",  "-6.8%"],
    ["Net Margin (%)",           "14.2%",   "15.2%",   "16.4%",   "16.7%",   "-30 bps"],
    ["EPS (CNY)",                "¥5.32",   "¥6.55",   "¥8.09",   "¥8.70",   "-7.0%"],
    ["P/E @ ¥136.91",            "25.7x",   "20.9x",   "16.9x",   "15.7x",   "—"],
    ["EV/EBITDA (approx.)",      "19.3x",   "15.1x",   "13.6x",   "12.7x",   "—"],
]
make_table(doc, det_h, det_r, col_widths=[2.0, 1.0, 1.0, 1.2, 1.2, 1.0])
src9 = doc.add_paragraph("Note: A = Actual; E = Estimate; EBITDA for FY2025A is estimated (full annual report not yet published at time of writing). All estimates are analyst projections.")
para_fmt(src9, space_before=1, space_after=2)
for run in src9.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)
src9b = doc.add_paragraph("Source: Sungrow 2025 Annual Report; Analyst estimates")
para_fmt(src9b, space_before=0, space_after=6)
for run in src9b.runs: run.font.size = Pt(7.5); run.italic = True; run.font.color.rgb = RGBColor(128,128,128)

add_chart(doc, "300274_chart9_estimates.png", width=6.2,
          caption="Figure 9 — Estimate Revisions: Revenue & Net Profit | Source: Sungrow 2025 Annual Report; Analyst estimates")

doc.add_page_break()

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX / SOURCES
# ─────────────────────────────────────────────────────────────────────────────
add_heading(doc, "APPENDIX: SOURCES & REFERENCES", size=12)
add_divider(doc)

app_p = doc.add_paragraph()
para_fmt(app_p, space_before=4, space_after=6)
add_run(app_p, "All financial data cited in this report is sourced from official company filings and recognized financial data providers. Clickable hyperlinks are provided below for each key source.", size=9.5)

sources_sections = [
    ("Primary Company Filings:", [
        ("Sungrow Power Supply 2025 Annual Report (年度报告摘要)", "https://static.cninfo.com.cn/finalpage/2026-04-01/1225066677.PDF",
         "Published April 1, 2026 — Primary source for all FY2025 financial data"),
        ("Sungrow Investor Call Transcript (电话交流纪要)", "https://finance.sina.com.cn/roll/2026-04-01/doc-inhsziht8333473.shtml",
         "March 31, 2026 — Q4 margin explanation, 2026 guidance"),
        ("Sungrow Annual Report Commentary (东方财富)", "https://caifuhao.eastmoney.com/news/20260402060936774616180",
         "April 2, 2026 — Detailed segment analysis and profitability breakdown"),
    ]),
    ("Financial Data Sources:", [
        ("StockAnalysis.com — Sungrow Quarterly Revenue History", "https://stockanalysis.com/quote/she/300274/revenue/",
         "Quarterly revenue data Q1 2024–Q4 2025"),
        ("StockAnalysis.com — Sungrow Overview & Valuation", "https://stockanalysis.com/quote/she/300274/",
         "Market cap, P/E, 52-week range, shares outstanding"),
        ("Futunn — Sungrow 2025 Annual Report Commentary", "https://news.futunn.com/en/post/71013862/sungrow-power-300274-2025-annual-report-commentary-energy-storage-systems",
         "Storage segment commentary and consensus estimates"),
        ("Wall Street CN — 阳光电源2025年年报", "https://wallstreetcn.com/articles/3768898",
         "Revenue near ¥900B, net profit ¥135B, storage grid-forming technology"),
    ]),
    ("Analyst Research & Commentary:", [
        ("Huarun Securities — Q3 2025 Report", "https://finance.sina.com.cn/roll/2025-11-10/doc-infwxhkx3003604.shtml",
         "November 2025 — AIDC product expansion, profitability trajectory"),
        ("Sina Finance — Q4 net profit decline analysis", "https://finance.sina.com.cn/roll/2026-04-01/doc-inhsziht8333473.shtml",
         "April 2026 — Q4 2025 quarterly net profit analysis"),
        ("East Money Creative — Price reference", "https://emcreative.eastmoney.com/app_fortune/article/index.html?artCode=20260407131020240171790&postId=1689922378",
         "CNY 122.66 as of April 7, 2026 (pre-recovery)"),
    ]),
]

for section_title, items in sources_sections:
    s_p = doc.add_paragraph()
    para_fmt(s_p, space_before=8, space_after=3)
    add_run(s_p, section_title, bold=True, size=10, color=(31, 78, 121))
    for title, url, note in items:
        item_p = doc.add_paragraph()
        para_fmt(item_p, left_indent=18, space_before=3, space_after=3)
        add_run(item_p, "• ", size=9.5)
        add_hyperlink(item_p, title, url)
        add_run(item_p, f"  —  {note}", size=9, italic=True, color=(100, 100, 100))

doc.add_paragraph()
add_divider(doc)

# Disclosure
disc_p = doc.add_paragraph()
para_fmt(disc_p, space_before=6, space_after=4)
add_run(disc_p,
        "IMPORTANT DISCLOSURES: This report is prepared for informational purposes only and does not constitute "
        "investment advice. All financial estimates are analyst projections and subject to change. Beat/miss analysis "
        "is based on pre-announcement consensus estimates sourced from Bloomberg and Futunn. Market data as of "
        "April 21, 2026. Sungrow Power Supply Co., Ltd. is traded on the Shenzhen Stock Exchange under ticker 300274.SZ. "
        "Past performance is not indicative of future results.",
        size=8, italic=True, color=(128, 128, 128))

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────
doc.save(OUT)
print(f"English report saved: {OUT}")
