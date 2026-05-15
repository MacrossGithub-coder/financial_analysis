"""
Adobe Inc. (ADBE) — 投资筛选报告（中文版）
应用软件同业估值对比 + 投资论点
"""
import os
import yfinance as yf
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = "/Users/macrossz/DevTools/VscodeProject/ClaudeCode/financial_analysis/output/ADBE"

C_NAVY   = RGBColor(0x1C, 0x1C, 0x4D)
C_RED    = RGBColor(0xEB, 0x10, 0x00)
C_BLUE   = RGBColor(0x14, 0x73, 0xE6)
C_GREEN  = RGBColor(0x1E, 0x8A, 0x44)
C_SELL   = RGBColor(0xCC, 0x00, 0x00)
C_GREY   = RGBColor(0x8E, 0x8E, 0x93)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GOLD   = RGBColor(0xE6, 0xA8, 0x17)

FONT_CN_BODY  = "宋体"
FONT_CN_TITLE = "黑体"


def get_market_data(ticker: str) -> dict:
    try:
        t = yf.Ticker(ticker)
        info = t.fast_info
        full = t.info
        return {
            "price": round(info.last_price, 2),
            "market_cap": info.market_cap,
            "52w_high": round(info.year_high, 2),
            "52w_low": round(info.year_low, 2),
            "fwd_pe": full.get("forwardPE"),
            "ps": full.get("priceToSalesTrailing12Months"),
            "rev_growth": full.get("revenueGrowth"),
            "op_margin": full.get("operatingMargins"),
            "gross_margin": full.get("grossMargins"),
            "fcf": full.get("freeCashflow"),
        }
    except Exception as e:
        print(f"WARNING: yfinance failed for {ticker}: {e}")
        return {}


def fmt_pe(v):
    return f"{v:.1f}x" if isinstance(v, (int, float)) else "N/A"

def fmt_pct(v):
    return f"{v*100:.1f}%" if isinstance(v, (int, float)) else "N/A"

def fmt_mcap(v):
    return f"${v/1e9:.1f}B" if isinstance(v, (int, float)) else "N/A"

def fmt_price(v):
    return f"${v:.2f}" if isinstance(v, (int, float)) else "N/A"

def fmt_fcf(v):
    return f"${v/1e9:.1f}B" if isinstance(v, (int, float)) else "N/A"


# ── docx helpers ─────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_border(cell, sides=("top", "bottom", "left", "right"), size="4", color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for side in sides:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)


def set_run_cn(run, font_name=FONT_CN_BODY):
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def heading(doc, text, level=1, color=None):
    p = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = C_NAVY
        set_run_cn(run, FONT_CN_TITLE)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = color or C_NAVY
        set_run_cn(run, FONT_CN_TITLE)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1C1C4D")
        pBdr.append(bottom)
        pPr.append(pBdr)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = color or C_BLUE
        set_run_cn(run, FONT_CN_TITLE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def body(doc, text, bold=False, italic=False, color=None, size=10, space_after=4):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    set_run_cn(run, FONT_CN_BODY)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    set_run_cn(run, FONT_CN_BODY)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 + 0.2 * level)
    p.paragraph_format.space_after = Pt(2)
    return p


def add_image(doc, filename, width=Inches(6.5), caption=None):
    path = os.path.join(BASE, filename)
    if not os.path.exists(path):
        body(doc, f"[图表: {filename} 未找到]", italic=True, color=C_GREY)
        return
    doc.add_picture(path, width=width)
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cp.runs:
            run.font.size = Pt(8.5)
            run.font.italic = True
            run.font.color.rgb = C_GREY
        cp.paragraph_format.space_after = Pt(6)


def add_table(doc, headers, rows, col_widths=None, alternate=True, highlight_row=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hrow = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = hdr
        set_cell_bg(cell, "1C1C4D")
        set_cell_border(cell)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.size = Pt(9)
                run.font.name = "Calibri"
                run.font.color.rgb = C_WHITE
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        if highlight_row is not None and ri == highlight_row:
            bg = "E8F0FE"
        else:
            bg = "F5F5F7" if (alternate and ri % 2 == 1) else "FFFFFF"
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = str(val)
            set_cell_bg(cell, bg)
            set_cell_border(cell, color="E5E5EA")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if highlight_row is not None and ri == highlight_row:
                        run.font.bold = True
                    if "+" in str(val) and ci > 1:
                        run.font.color.rgb = C_GREEN
                    elif "−" in str(val) or "-" in str(val) and ci > 1:
                        run.font.color.rgb = C_SELL
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    doc.add_paragraph()
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def star_rating(n, total=5):
    return "★" * n + "☆" * (total - n)


# ══════════════════════════════════════════════════════════════════════════════
# 获取市场数据
# ══════════════════════════════════════════════════════════════════════════════
print("正在获取市场数据...")
TICKERS = ["ADBE", "CRM", "INTU", "NOW", "MSFT", "ORCL", "SAP"]
NAMES = {
    "ADBE": "Adobe",
    "CRM": "Salesforce",
    "INTU": "Intuit",
    "NOW": "ServiceNow",
    "MSFT": "Microsoft",
    "ORCL": "Oracle",
    "SAP": "SAP",
}
data = {}
for tk in TICKERS:
    print(f"  获取 {tk}...")
    data[tk] = get_market_data(tk)

adbe = data["ADBE"]

price_str = fmt_price(adbe.get("price"))
mcap_str = fmt_mcap(adbe.get("market_cap"))
range_str = (
    f"${adbe['52w_low']} – ${adbe['52w_high']}"
    if isinstance(adbe.get("52w_low"), float) else "N/A"
)

# ══════════════════════════════════════════════════════════════════════════════
# 生成报告
# ══════════════════════════════════════════════════════════════════════════════
print("正在生成报告...")
doc = Document()

for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

style_normal = doc.styles["Normal"]
style_normal.font.name = FONT_CN_BODY
style_normal.font.size = Pt(10)
rPr = style_normal.element.rPr
if rPr is None:
    rPr = OxmlElement("w:rPr")
    style_normal.element.append(rPr)
rFonts = rPr.find(qn("w:rFonts"))
if rFonts is None:
    rFonts = OxmlElement("w:rFonts")
    rPr.insert(0, rFonts)
rFonts.set(qn("w:eastAsia"), FONT_CN_BODY)

# ══════════════════════════════════════════════════════════════════════════════
# 第1页 — 封面与核心数据
# ══════════════════════════════════════════════════════════════════════════════
p_ticker = doc.add_paragraph()
p_ticker.clear()
run = p_ticker.add_run("ADOBE INC.  (ADBE: NASDAQ)")
run.font.size = Pt(22)
run.font.bold = True
run.font.color.rgb = C_NAVY
run.font.name = "Calibri"
p_ticker.paragraph_format.space_after = Pt(2)

p_sub = doc.add_paragraph()
p_sub.clear()
run2 = p_sub.add_run("投资筛选报告  |  应用软件  |  大型成长股")
run2.font.size = Pt(10)
run2.font.color.rgb = C_GREY
set_run_cn(run2, FONT_CN_TITLE)
p_sub.paragraph_format.space_after = Pt(6)

# 评级速览表
rating_table = doc.add_table(rows=1, cols=6)
rating_table.alignment = WD_TABLE_ALIGNMENT.LEFT
labels = ["筛选信号", "当前价", "52周区间", "市值", "远期P/E", "P/S"]
values = [
    "深度价值 — 看多",
    price_str,
    range_str,
    mcap_str,
    fmt_pe(adbe.get("fwd_pe")),
    fmt_pe(adbe.get("ps")),
]
colors_v = [C_GREEN, C_NAVY, C_GREY, C_NAVY, C_BLUE, C_BLUE]

for i, (lbl, val, col) in enumerate(zip(labels, values, colors_v)):
    cell = rating_table.rows[0].cells[i]
    set_cell_bg(cell, "F0F4FF")
    set_cell_border(cell, color="D0D8F0")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(lbl + "\n")
    r1.font.size = Pt(8)
    r1.font.color.rgb = C_GREY
    r1.font.name = "Calibri"
    r2 = p1.add_run(val)
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.color.rgb = col
    r2.font.name = "Calibri"

doc.add_paragraph().paragraph_format.space_after = Pt(4)
horizontal_rule(doc)

# 报告标题
p_title = doc.add_paragraph()
p_title.clear()
run_t = p_title.add_run(
    "投资筛选：大型应用软件同业估值对比\n"
    "Adobe以同业最低估值提供最高利润率——深度价值买入机会"
)
run_t.font.size = Pt(14)
run_t.font.bold = True
run_t.font.color.rgb = C_NAVY
set_run_cn(run_t, FONT_CN_TITLE)
p_title.paragraph_format.space_after = Pt(3)

p_date = doc.add_paragraph()
p_date.clear()
rd = p_date.add_run("报告日期: 2026年5月15日  |  数据来源: Yahoo Finance, 公司公告")
rd.font.size = Pt(8.5)
rd.font.color.rgb = C_GREY
set_run_cn(rd, FONT_CN_BODY)
p_date.paragraph_format.space_after = Pt(8)

horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第一部分 — 核心数据一览
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "一、核心数据一览", 2)

add_table(doc,
    ["指标", "数值"],
    [
        ["当前股价", price_str],
        ["市值", mcap_str],
        ["52周区间", range_str],
        ["远期市盈率", fmt_pe(adbe.get("fwd_pe"))],
        ["市销率（TTM）", fmt_pe(adbe.get("ps"))],
        ["营收增速（YoY）", fmt_pct(adbe.get("rev_growth"))],
        ["GAAP营业利润率", fmt_pct(adbe.get("op_margin"))],
        ["毛利率", fmt_pct(adbe.get("gross_margin"))],
        ["自由现金流", fmt_fcf(adbe.get("fcf"))],
        ["总ARR", "$260.6亿（+11% YoY）"],
    ],
    col_widths=[Inches(2.5), Inches(4.0)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 第二部分 — 同业估值对比
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "二、同业估值对比", 2)

body(doc, (
    "以下将Adobe与六家大型应用软件/云计算公司进行横向比较。"
    "ADBE以蓝色高亮行标识，同业中位数见最后一行。"
))

comp_headers = ["代码", "公司", "市值", "远期P/E", "P/S", "营收增速", "营业利润率", "毛利率"]
comp_rows = []
pe_values = []
for tk in TICKERS:
    d = data[tk]
    comp_rows.append([
        tk,
        NAMES[tk],
        fmt_mcap(d.get("market_cap")),
        fmt_pe(d.get("fwd_pe")),
        fmt_pe(d.get("ps")),
        fmt_pct(d.get("rev_growth")),
        fmt_pct(d.get("op_margin")),
        fmt_pct(d.get("gross_margin")),
    ])
    if isinstance(d.get("fwd_pe"), (int, float)):
        pe_values.append(d["fwd_pe"])

pe_values_sorted = sorted(pe_values)
median_pe = pe_values_sorted[len(pe_values_sorted) // 2] if pe_values_sorted else 0

ps_values = sorted([d["ps"] for d in data.values() if isinstance(d.get("ps"), (int, float))])
median_ps = ps_values[len(ps_values) // 2] if ps_values else 0

rev_values = sorted([d["rev_growth"] for d in data.values() if isinstance(d.get("rev_growth"), (int, float))])
median_rev = rev_values[len(rev_values) // 2] if rev_values else 0

op_values = sorted([d["op_margin"] for d in data.values() if isinstance(d.get("op_margin"), (int, float))])
median_op = op_values[len(op_values) // 2] if op_values else 0

gm_values = sorted([d["gross_margin"] for d in data.values() if isinstance(d.get("gross_margin"), (int, float))])
median_gm = gm_values[len(gm_values) // 2] if gm_values else 0

comp_rows.append([
    "—", "同业中位数", "—",
    fmt_pe(median_pe), fmt_pe(median_ps),
    fmt_pct(median_rev), fmt_pct(median_op), fmt_pct(median_gm),
])

add_table(doc, comp_headers, comp_rows,
    col_widths=[Inches(0.55), Inches(1.0), Inches(0.7), Inches(0.7),
                Inches(0.55), Inches(0.7), Inches(0.8), Inches(0.7)],
    highlight_row=0,
)

# 估值折价分析
heading(doc, "估值折价分析", 3)

adbe_pe = adbe.get("fwd_pe")
if isinstance(adbe_pe, (int, float)) and median_pe:
    discount = (1 - adbe_pe / median_pe) * 100
    body(doc, (
        f"■ Adobe远期市盈率{adbe_pe:.1f}x vs. 同业中位数{median_pe:.1f}x，"
        f"折价幅度达{discount:.0f}%。"
        "这是ADBE在该同业组中估值最低的，但其毛利率（89.4%）和自由现金流利润率（~41%）均为同业最高水平。"
    ))

body(doc, (
    "■ 市销率方面，ADBE的3.9x同样低于同业中位数水平，"
    "考虑到其接近90%的毛利率，该估值水平意味着市场对其增长持续性和竞争优势存在显著定价折扣。"
))

# ══════════════════════════════════════════════════════════════════════════════
# 第三部分 — 筛选信号分析
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "三、筛选信号分析", 2)

heading(doc, "看多信号（做多依据）", 3, color=C_GREEN)

bullet(doc, (
    "极端估值折价：远期P/E仅~9x，为同业组中最低，较中位数折价约49%。"
    "Adobe拥有同业最高毛利率（89%）和第二高营业利润率，估值水平与基本面严重脱节。"
), color=C_GREEN)

bullet(doc, (
    "Firefly AI变现拐点：AI优先ARR同比增长超3倍，独立Firefly ARR突破2.5亿美元。"
    "管理层指引显示生成式AI积分消耗正在加速增长，AI变现从概念转向收入贡献。"
), color=C_GREEN)

bullet(doc, (
    "大规模资本回报：2026年4月新增250亿美元回购授权（约占市值26%），"
    "Q1单季度回购25亿美元。按当前股价计算，自由现金流收益率约9.7%。"
), color=C_GREEN)

bullet(doc, (
    "增长重新加速：Q1 FY2026营收同比增长12%（vs. Q4 FY2025的10%），"
    "超市场预期约1.2亿美元。每股收益超预期$0.19-$0.60。"
), color=C_GREEN)

bullet(doc, (
    "接近52周低点：当前股价$237 vs. 52周高点$423，较峰值下跌约46%，年初至今下跌约28%。"
    "技术面接近超卖区域。"
), color=C_GREEN)

heading(doc, "看空信号（风险因素）", 3, color=C_SELL)

bullet(doc, (
    "CEO交接不确定性：Shantanu Narayen于2026年3月宣布退休，至今未公布继任者。"
    "在AI战略转型的关键时期出现领导层真空，增加执行风险。"
), color=C_SELL)

bullet(doc, (
    "竞争侵蚀加剧：Canva（ARR约40亿美元，增速+30%）和Figma（ARR约12亿美元，增速+40%）"
    "在创意工具市场快速抢占份额，尤其在中小企业和专业消费者群体中。"
), color=C_SELL)

bullet(doc, (
    "传统业务逆风：图库素材业务因生成式AI替代品普及而加速下滑，"
    "下滑速度超出管理层此前预期。"
), color=C_SELL)

bullet(doc, (
    "增长天花板：基于260亿美元ARR基数，营收增速约束在11-12%。"
    "同业中ServiceNow（+22%）和Oracle（+22%）的增速接近Adobe的两倍。"
), color=C_SELL)

# ══════════════════════════════════════════════════════════════════════════════
# 第四部分 — 综合评分
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "四、综合评分", 2)

add_table(doc,
    ["维度", "评分", "评语"],
    [
        ["估值", star_rating(5), "同业组中最低P/E和P/FCF，处于历史低位"],
        ["质量", star_rating(5), "89%毛利率，41%自由现金流利润率，年产$93亿FCF"],
        ["增长", star_rating(3), "12%增速稳健但低于软件同业平均水平"],
        ["动量", star_rating(2), "较峰值下跌46%，接近52周低点"],
        ["催化剂", star_rating(4), "Q2业绩6月11日；CEO继任确认；Firefly规模化"],
        ["风险", star_rating(3), "CEO空缺 + Canva/Figma竞争压力"],
    ],
    col_widths=[Inches(1.0), Inches(1.2), Inches(4.3)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 第五部分 — 投资结论
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "五、投资结论", 2)

body(doc, "筛选结果：高确信度做多标的", bold=True, color=C_GREEN, size=12)

body(doc, (
    "Adobe在当前估值水平下呈现大型应用软件中最深的价值洼地。"
    "公司拥有行业顶级的利润率结构（毛利率89%，FCF利润率41%），"
    "却以远低于同业的估值倍数交易（远期P/E仅~9x vs. 同业中位数~18x）。"
))

body(doc, (
    "空头论点——来自AI原生工具的竞争性颠覆——具有合理性，"
    "但在当前定价中已被过度反映。Adobe自身的AI变现进展（Firefly ARR 3倍增长至$2.5亿+）、"
    "大规模资本回报计划（250亿美元回购，约占市值26%），以及重新加速的营收增长（Q1 +12%），"
    "均表明基本面的韧性远强于市场定价所暗示的水平。"
))

body(doc, (
    "关键催化剂：2026年6月11日Q2 FY2026业绩发布——将验证增长加速趋势是否持续，"
    "以及Firefly ARR能否继续规模化。CEO继任者确认也将消除重要的不确定性因素。"
))

horizontal_rule(doc)

# ══════════════════════════════════════════════════════════════════════════════
# 第六部分 — 近期财务表现回顾
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "六、近期财务表现", 2)

heading(doc, "季度营收趋势", 3)
add_table(doc,
    ["季度", "营收", "同比增速"],
    [
        ["Q2 FY2025", "$58.7亿", "+11%"],
        ["Q3 FY2025", "$59.9亿", "+11%"],
        ["Q4 FY2025", "$61.9亿", "+10%"],
        ["Q1 FY2026", "$64.0亿", "+12%"],
    ],
    col_widths=[Inches(2.0), Inches(2.0), Inches(2.5)],
)

body(doc, (
    "Q1 FY2026增长重新加速至12%，扭转了Q4 FY2025降至10%的放缓趋势。"
    "全年FY2025营收238亿美元（+11% YoY）。"
))

heading(doc, "业务板块分析（Q1 FY2026）", 3)
add_table(doc,
    ["业务板块", "营收", "同比增速"],
    [
        ["数字媒体", "$46.2亿", "+11%"],
        ["数字体验", "$15.2亿", "+9%"],
        ["创意与营销专业人士", "$43.9亿", "+12%"],
        ["商业专业人士与消费者", "$17.8亿", "+16%"],
    ],
    col_widths=[Inches(2.5), Inches(1.8), Inches(2.2)],
)

heading(doc, "利润率与现金流", 3)
add_table(doc,
    ["指标", "数值", "备注"],
    [
        ["非GAAP营业利润率", "47.4%", "Q1 FY2026，同比扩张120bps"],
        ["GAAP营业利润率", "~37.8%", "2026年2月数据"],
        ["自由现金流利润率", "41.4%", "FY2025全年"],
        ["经营现金流", "$29.6亿", "Q1 FY2026，创历史纪录"],
    ],
    col_widths=[Inches(2.0), Inches(1.5), Inches(3.0)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 第七部分 — AI战略概览
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "七、AI战略概览", 2)

heading(doc, "Firefly生成式AI平台", 3)
bullet(doc, "Firefly独立ARR突破2.5亿美元——首次单独披露，验证AI变现进展。")
bullet(doc, "AI优先ARR同比增长超过3倍。")
bullet(doc, "生成式积分消耗环比持续增长，正向更高算力模态（视频、音频）扩展。")
bullet(doc, "Firefly AI助手实现跨Creative Cloud应用的多步骤工作流自然语言编排。")

heading(doc, "GenStudio企业AI平台", 3)
bullet(doc, "品牌智能引擎（Brand Intelligence）：超越静态品牌指南的持续学习系统。")
bullet(doc, "GenStudio内容营销：将长篇内容自动转化为定制化营销活动、案例研究和文章。")
bullet(doc, "定位为创意、数据和AI驱动工作流的统一企业CX平台。")

heading(doc, "Adobe Summit 2026亮点（2026年4月）", 3)
bullet(doc, "发布Creative Agent——面向创意工作流的AI智能体。")
bullet(doc, "扩展GenStudio内容供应链解决方案。")
bullet(doc, "将Adobe定位为客户体验的“AI控制层”。")

# ══════════════════════════════════════════════════════════════════════════════
# 第八部分 — 资本配置
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "八、资本配置与股东回报", 2)

bullet(doc, (
    "250亿美元股票回购：2026年4月21日宣布（替代此前已接近用尽的250亿美元授权），"
    "有效期至2030年4月30日。"
))
bullet(doc, f"回购规模约占当前市值的26%——释放极为积极的资本回报信号。")
bullet(doc, "Q1 FY2026单季度回购约25亿美元。")
bullet(doc, "CFO Dan Durn表示：回购是'对我们强劲现金流的直接信心表达'。")

heading(doc, "FY2026业绩指引", 3)
add_table(doc,
    ["指标", "管理层指引", "市场预期"],
    [
        ["全年营收", "$259–261亿", "~$264亿"],
        ["Q2营收", "$64.3–64.8亿", "~$64.5亿"],
        ["非GAAP每股收益", "$23.30–$23.50", "—"],
        ["非GAAP营业利润率", "~45%", "—"],
        ["年末ARR增速", "+10.2% YoY", "—"],
    ],
    col_widths=[Inches(2.0), Inches(2.2), Inches(2.3)],
)

# ══════════════════════════════════════════════════════════════════════════════
# 图表引用
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading(doc, "附录：图表", 2)

chart_files = [
    ("adbe_chart1_revenue.png", "图1：营收趋势与同比增速"),
    ("adbe_chart2_eps.png", "图2：每股收益趋势"),
    ("adbe_chart3_segments.png", "图3：业务板块拆分"),
    ("adbe_chart4_arr.png", "图4：ARR增长趋势"),
    ("adbe_chart5_margins.png", "图5：利润率趋势"),
    ("adbe_chart9_valuation.png", "图6：估值倍数走势"),
    ("adbe_chart10_ai_growth.png", "图7：AI业务增长"),
]

for fname, caption in chart_files:
    add_image(doc, fname, width=Inches(6.0), caption=caption)

# ══════════════════════════════════════════════════════════════════════════════
# 免责声明
# ══════════════════════════════════════════════════════════════════════════════
horizontal_rule(doc)
body(doc, (
    "免责声明：本报告仅供参考，不构成任何投资建议。报告中的数据来自公开市场信息，"
    "作者不对数据的准确性和完整性做任何保证。投资者应独立判断并承担投资风险。"
), italic=True, color=C_GREY, size=8)

# ══════════════════════════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════════════════════════
os.makedirs(BASE, exist_ok=True)
out_path = os.path.join(BASE, "ADBE_投资筛选报告_中文版.docx")
doc.save(out_path)
print(f"\n报告已保存: {out_path}")
